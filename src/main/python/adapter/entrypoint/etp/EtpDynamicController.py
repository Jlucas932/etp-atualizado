import os
import json
import uuid
import tempfile
import logging
import traceback
from datetime import datetime
from pathlib import Path
from typing import Any, List, Optional, Tuple
from flask import Blueprint, request, jsonify, send_file, g
from flask_cors import cross_origin
import re
import unicodedata

from domain.interfaces.dataprovider.DatabaseConfig import db
from domain.dto.EtpOrm import EtpSession, EtpDocument
from domain.dto.EtpDto import DocumentAnalysis, KnowledgeBase, ChatSession, EtpTemplate
from domain.dto.ConversationModels import Conversation, Message
from domain.repositories.ConversationRepository import ConversationRepo, MessageRepo
from domain.usecase.etp.verify_federal import resolve_lexml, summarize_for_user, parse_legal_norm_string
from application.config.LimiterConfig import limiter
from rag.retrieval import search_requirements
from domain.services.etp_dynamic import init_etp_dynamic
from domain.usecase.etp.legal_norms_interpreter import parse_legal_norms
from domain.usecase.etp.price_research_interpreter import parse_price_research
from domain.usecase.etp.legal_basis_interpreter import parse_legal_basis
from domain.usecase.etp.document_composer import compose_etp_document
from domain.usecase.etp.html_renderer import render_etp_html
from application.ai.generator import (
    get_etp_generator,
    FallbackGenerator,
    dedupe_requirements,
    generate_text_with_model,
)
from application.ai.hybrid_models import OpenAIChatConsultive, OpenAIFinalWriter, OpenAIIntentParser
from application.nlu.intent_requirements import (
    ACCEPT as REQ_ACCEPT,
    EDIT as REQ_EDIT,
    UNCLEAR as REQ_UNCLEAR,
    detect_intent as detect_requirements_intent,
)
from application.ai import intents
from application.services.preview_builder import build_preview, build_etp_markdown
from domain.usecase.etp.state_machine import (
    is_user_confirmed, validate_state_transition, can_generate_etp,
    handle_other_intent, handle_http_error, validate_generator_exists,
    get_next_state_after_suggestion
)
from domain.usecase.etp import conversational_state_machine as csm
from io import BytesIO
from docx import Document as DocxDocument

etp_dynamic_bp = Blueprint('etp_dynamic', __name__)

# Module-level variables for lazy initialization
_etp_generator = None
_prompt_generator = None  
_rag_system = None
_initialized = False

# Simple generator for requirements (with fallback)
_simple_generator = None
_openai_client = None

# Hybrid model generators
chat_gen   = OpenAIChatConsultive()
final_gen  = OpenAIFinalWriter()
intent_par = OpenAIIntentParser()

# ===================== STAGE MACHINE CONFIGURATION =====================
# Deterministic FSM: 9-stage conversational flow without intermediate refinement
STAGE_ORDER = [
    "collect_need",
    "suggest_requirements",
    "solution_strategies",
    "pca",
    "legal_norms",
    "qty_value",
    "installment",
    "summary",
    "preview"
]

# Stage transition map
NEXT_STAGE = {stage: STAGE_ORDER[i+1] for i, stage in enumerate(STAGE_ORDER[:-1])}

logger = logging.getLogger(__name__)

_REQUIREMENTS_FOLLOWUP_TEXT: Optional[str] = None


def _requirements_followup_text() -> str:
    """Load follow-up prompt text from template file (cached)."""

    global _REQUIREMENTS_FOLLOWUP_TEXT
    if _REQUIREMENTS_FOLLOWUP_TEXT is not None:
        return _REQUIREMENTS_FOLLOWUP_TEXT

    template_path = (
        Path(__file__).resolve().parents[3]
        / "application"
        / "messages"
        / "templates"
        / "requirements_followup.pt-BR.txt"
    )

    try:
        _REQUIREMENTS_FOLLOWUP_TEXT = template_path.read_text(encoding="utf-8").strip()
    except FileNotFoundError:
        _REQUIREMENTS_FOLLOWUP_TEXT = (
            "Esses requisitos fazem sentido para sua necessidade? Quer ajustar algo agora?\n"
            "Posso trocar, remover ou adicionar itens, e também adaptar prazos, métricas e SLAs."
        )

    return _REQUIREMENTS_FOLLOWUP_TEXT


def _extract_requirement_text(item: Any) -> str:
    if isinstance(item, dict):
        return (
            item.get("text")
            or item.get("descricao")
            or item.get("requirement")
            or str(item)
        )
    return str(item)


def _format_requirements_list(items: List[Any]) -> str:
    lines: List[str] = []
    for idx, item in enumerate(items, 1):
        text = _extract_requirement_text(item).strip()
        if not text:
            continue
        lines.append(f"{idx}. {text}")
    return "\n".join(lines)


def handle_requirements_edit(user_message: str, current_requirements: List[str], context_need: str) -> List[str]:
    """Replace numbered requirements using the language model and clean suffixes."""

    indices = [int(n) for n in re.findall(r"\b\d+\b", user_message or "")]
    if not indices:
        return current_requirements

    updated_requirements: List[str] = []
    need_context = context_need or "a necessidade informada pelo usuário"

    for i, req in enumerate(current_requirements, start=1):
        req_text = _extract_requirement_text(req).strip()
        if i in indices:
            prompt = (
                "Substitua este requisito por outro mais adequado à necessidade: "
                f"{need_context}. "
                f"Requisito original: {req_text}. "
                "Não use marcações como (Obrigatório) ou (Desejável). "
                "Mantenha o estilo formal, objetivo e dentro do contexto da contratação pública."
            )
            new_req = generate_text_with_model(prompt).strip()
            new_req = re.sub(r"\s*\(.*?\)", "", new_req).strip()
            updated_requirements.append(new_req)
        else:
            clean_req = re.sub(r"\s*\(.*?\)", "", req_text).strip()
            updated_requirements.append(clean_req)

    return updated_requirements


def _update_dict_requirement(item: dict, new_text: str) -> dict:
    updated = dict(item)
    for key in ("text", "descricao", "requirement"):
        if key in updated and isinstance(updated[key], str) and updated[key]:
            updated[key] = new_text
            return updated
    updated["text"] = new_text
    return updated


def _apply_requirements_edit(requirements: List[Any], user_text: str) -> Tuple[Optional[List[Any]], str]:
    """Apply simple edit operations (remove, replace, add, SLA change)."""

    if not user_text:
        return None, (
            "Não consegui identificar o ajuste. Pode me dizer qual item quer alterar, "
            "remover ou o que deseja adicionar?"
        )

    lowered = user_text.lower()
    updated: List[Any] = list(requirements)

    # Remove by index (e.g., "remove o 2")
    remove_idx_match = re.search(r"remov\w*\s*(?:o\s+|a\s+)?(?:item\s+)?(\d+)", lowered)
    if remove_idx_match:
        idx = int(remove_idx_match.group(1)) - 1
        if 0 <= idx < len(updated):
            removed_item = updated.pop(idx)
            removed_text = _extract_requirement_text(removed_item).strip()
            return updated, f"Removi o item {idx + 1}: {removed_text}"

    # Replace by index ("troque o 3 por ...")
    replace_idx_match = re.search(
        r"(troq\w*|substitu\w*|altera\w*|muda\w*)\s*(?:o\s+|a\s+)?(?:item\s+)?(\d+)\s*(?:por|para)\s+(.+)$",
        user_text,
        flags=re.IGNORECASE,
    )
    if replace_idx_match:
        idx = int(replace_idx_match.group(2)) - 1
        new_value = replace_idx_match.group(3).strip()
        if 0 <= idx < len(updated) and new_value:
            old_item = updated[idx]
            old_text = _extract_requirement_text(old_item).strip()
            if isinstance(old_item, dict):
                updated[idx] = _update_dict_requirement(old_item, new_value)
            else:
                updated[idx] = new_value
            return updated, f"Troquei o item {idx + 1}: '{old_text}' por '{new_value}'"

    # Remove by partial text ("remove backup")
    if "remov" in lowered:
        remove_text_match = re.search(
            r"remov\w*\s*(?:o\s+|a\s+)?(?:item\s+)?(?:de\s+)?(.+)",
            user_text,
            flags=re.IGNORECASE,
        )
        if remove_text_match:
            snippet = remove_text_match.group(1).strip()
            if snippet:
                for idx, item in enumerate(updated):
                    if snippet.lower() in _extract_requirement_text(item).lower():
                        removed_item = updated.pop(idx)
                        removed_text = _extract_requirement_text(removed_item).strip()
                        return updated, f"Removi o item {idx + 1}: {removed_text}"

    # Replace by partial text ("troque o requisito de backup por ...")
    replace_text_match = re.search(
        r"(troq\w*|substitu\w*|altera\w*|muda\w*)\s+(?:o\s+|a\s+)?(.+?)\s+(?:por|para)\s+(.+)$",
        user_text,
        flags=re.IGNORECASE,
    )
    if replace_text_match:
        target = replace_text_match.group(2).strip()
        new_value = replace_text_match.group(3).strip()
        if target and new_value:
            for idx, item in enumerate(updated):
                if target.lower() in _extract_requirement_text(item).lower():
                    old_text = _extract_requirement_text(item).strip()
                    if isinstance(item, dict):
                        updated[idx] = _update_dict_requirement(item, new_value)
                    else:
                        updated[idx] = new_value
                    return updated, f"Troquei o item {idx + 1}: '{old_text}' por '{new_value}'"

    # Add new requirement ("adicione um requisito sobre ...", "faltou ...")
    add_match = re.search(
        r"(adicione|inclua|inclui|adiciona)\b.+?\b(sobre|de|com)\b\s+(.+)$",
        user_text,
        flags=re.IGNORECASE,
    )
    if add_match:
        new_req = add_match.group(3).strip()
        if new_req:
            updated.append(new_req)
            return updated, f"Adicionei um novo requisito: '{new_req}'"

    faltou_match = re.search(
        r"faltou\s+(?:um\s+)?(?:requisito\s+)?(.+)$",
        user_text,
        flags=re.IGNORECASE,
    )
    if faltou_match:
        new_req = faltou_match.group(1).strip()
        if new_req:
            updated.append(new_req)
            return updated, f"Adicionei um novo requisito: '{new_req}'"

    # SLA adjustments ("mude o SLA de 98 para 99")
    sla_match = re.search(r"sla.*?(\d{2,3}).*?para.*?(\d{2,3})", lowered)
    if sla_match:
        old_sla, new_sla = sla_match.group(1), sla_match.group(2)
        for idx, item in enumerate(updated):
            text = _extract_requirement_text(item)
            if re.search(r"\bSLA\b", text, flags=re.IGNORECASE) and old_sla in text:
                if isinstance(item, dict):
                    updated_item = dict(item)
                    replaced = False
                    for key in ("text", "descricao", "requirement"):
                        value = updated_item.get(key)
                        if isinstance(value, str) and old_sla in value:
                            updated_item[key] = re.sub(old_sla, new_sla, value, count=1)
                            replaced = True
                            break
                    if not replaced and isinstance(updated_item.get("sla"), dict):
                        updated_item["sla"] = dict(updated_item["sla"])
                        updated_item["sla"]["valor"] = new_sla
                    updated[idx] = updated_item
                else:
                    updated[idx] = re.sub(old_sla, new_sla, text, count=1)
                return updated, f"Ajustei o SLA de {old_sla} para {new_sla} no item {idx + 1}"

    return None, (
        "Não consegui identificar o ajuste. Pode me dizer qual item quer alterar, "
        "remover ou o que deseja adicionar?"
    )

def _get_openai_client():
    """Get or create OpenAI client"""
    global _openai_client
    if _openai_client is None:
        api_key = os.getenv('OPENAI_API_KEY')
        if api_key:
            import openai
            _openai_client = openai.OpenAI(api_key=api_key)
    return _openai_client

def get_llm_client():
    """Helper to ensure llm_client is always available (never None)"""
    client = _get_openai_client()
    if client is None:
        logger.warning("[CLIENT] OpenAI client not initialized, creating fallback")
        api_key = os.getenv('OPENAI_API_KEY')
        if api_key:
            import openai
            client = openai.OpenAI(api_key=api_key)
    return client

def get_model_name():
    """Get standardized model name from environment"""
    return os.getenv("OPENAI_MODEL", "gpt-4.1")

def _get_simple_generator():
    """Get or create the simple requirements generator with automatic fallback"""
    global _simple_generator
    if _simple_generator is None:
        openai_client = _get_openai_client()
        _simple_generator = get_etp_generator(openai_client=openai_client)
    return _simple_generator

def get_current_user_id():
    """Get current user ID from flask.g or request headers, fallback to 'anonymous'"""
    # Check if user_id is in flask.g (set by auth middleware)
    if hasattr(g, 'user_id') and g.user_id:
        return str(g.user_id)
    
    # Check X-User-Id header
    user_id = request.headers.get('X-User-Id')
    if user_id:
        return str(user_id)
    
    # Fallback to anonymous for development
    return 'anonymous'

# ===================== Helpers consultivos =====================
CONFIRM_RE = re.compile(r'\b(ok|ok\.?|pode seguir|segue|prosseguir|manter|aceito|concordo|fechou|confirmo|confirmar|fechado|vamos em frente)\b', re.I)
YES_RE = re.compile(r'\b(sim|s|claro|ok|certo|pode|pode sim)\b', re.I)
NO_RE = re.compile(r'\b(n[aã]o|nao|n)\b', re.I)
UNCERTAIN_RE = re.compile(r'\b(n[aã]o sei|nao sei|talvez|tanto faz|indefinido|a definir|a confirmar)\b', re.I)

def _normalize(txt: str) -> str:
    if not txt:
        return ''
    s = txt.strip()
    s = unicodedata.normalize('NFKD', s)
    s = ''.join(c for c in s if not unicodedata.combining(c))
    return s

def _is_confirm(txt: str) -> bool:
    return bool(CONFIRM_RE.search(txt or ''))

def _is_uncertain(txt: str) -> bool:
    t = _normalize(txt).lower()
    return not t or bool(UNCERTAIN_RE.search(t)) or t in {'nao sei','n sei','?','nada','-'}

def _yes(txt: str) -> bool:
    return bool(YES_RE.search(txt or ''))

def _no(txt: str) -> bool:
    # tolera typos tipo "naoi", "naoii"
    t = _normalize(txt).lower()
    return bool(NO_RE.search(t)) or t.startswith('nao')

def _pick_numbers(txt: str, max_n: int) -> list:
    """Extrai seleções do tipo '1 e 3', '2,4', '1-2'."""
    t = _normalize(txt).lower()
    picks = set()
    for a,b in re.findall(r'(\d+)\s*-\s*(\d+)', t):
        a,b = int(a), int(b)
        for k in range(min(a,b), max(a,b)+1):
            if 1 <= k <= max_n: picks.add(k)
    for n in re.findall(r'\b(\d+)\b', t):
        k = int(n)
        if 1 <= k <= max_n: picks.add(k)
    return sorted(picks)

def _any_in(txt: str, options: list) -> str:
    """Detecta uma opção por palavra-chave dentro do texto."""
    if not txt: return None
    t = _normalize(txt).lower()
    for op in options:
        k = _normalize(op).lower()
        if k in t:
            return op
    return None

def _is_free_confirm(text: str) -> bool:
    """
    Check if user input is a free confirmation (ok, pode seguir, etc.)
    Returns True if it's a simple confirmation without actual content.
    """
    if not text:
        return False
    t = text.strip().lower()
    return bool(re.match(r'^(ok(ay)?|pode(\s)?(seguir|continuar)|perfeito|segue|entendido|manda ver|toca ficha|vamos lá)\.?$', t, re.I))

def _strategy_selection(text: str, strategies: list[str]) -> int | None:
    """
    Returns the index (0-based) of the selected strategy, or None.
    Rules:
      - numeric "1" to "9" → index (n-1) if exists
      - text match: fuzzy (simple) >= 0.60 on title
    """
    if not text:
        return None
    t = text.strip().lower()

    # numeric selection
    m = re.match(r'^\s*([1-9])\s*$', t)
    if m:
        idx = int(m.group(1)) - 1
        return idx if 0 <= idx < len(strategies) else None

    # text selection with fuzzy matching
    def _score(a: str, b: str) -> float:
        # simple jaccard over tokens
        sa, sb = set(re.findall(r'\w+', a.lower())), set(re.findall(r'\w+', b.lower()))
        inter = len(sa & sb)
        uni = len(sa | sb) or 1
        score = inter / uni
        # boost if user input is substring of title
        if a in b:
            score = max(score, 0.70)

# ===================== Decision-awaiting helpers =====================
def ask_user_decision(session, prompt_text: str, proposal_text: str, stage: str) -> str:
    """
    Arms the session for decision-awaiting mode and returns a prompt with options.
    Does NOT advance stage until user makes a choice.
    
    Args:
        session: ETP session object
        prompt_text: Contextual prompt explaining the situation
        proposal_text: The proposed text/content for acceptance
        stage: Current stage name (for logging)
    
    Returns:
        String with prompt + options (1/2/3)
    """
    answers = session.get_answers()
    if 'state' not in answers:
        answers['state'] = {}
    
    answers['state']['awaiting_decision'] = True
    answers['state']['pending_proposal'] = proposal_text
    answers['state']['decision_stage'] = stage
    session.set_answers(answers)
    db.session.commit()
    
    logger.info(f"[{stage.upper()}] awaiting_decision (not advancing)")
    
    return (
        prompt_text + "\n\n"
        "**Opções:**\n"
        "1) **Aceitar** a sugestão inicial acima e seguir.\n"
        "2) **Registrar como Pendente** (com a justificativa proposta) e seguir.\n"
        "3) **Quero debater mais** (ficamos neste ponto e ajustamos juntos)."
    )

def try_consume_decision(session, user_text: str) -> dict | str | None:
    """
    Attempts to consume a pending decision from the user.
    
    Args:
        session: ETP session object
        user_text: User's message
    
    Returns:
        - None: No decision is pending
        - str: Confirmation prompt asking user to clarify their choice
        - dict: Decision object with 'action' and 'text' keys
            action: 'accept' | 'pendente' | 'debate'
            text: The proposal text (for accept/pendente actions)
    """
    answers = session.get_answers()
    state = answers.get('state', {})
    
    if not state.get('awaiting_decision'):
        return None  # No decision pending
    
    t = (user_text or "").strip().lower()
    
    # Check if it's a vague acknowledgment that should not be treated as a decision
    if intents.is_vague_ack(user_text):
        logger.info(f"[DECISION] Vague acknowledgment detected, asking for clarification: {user_text[:30]}")
        return (
            "Só para confirmar: prefere **1) aceitar**, **2) marcar como pendente**, "
            "ou **3) debater mais**?"
        )
    
    # Detect choice (removed "ok" and "sim" from accepted list - they should be explicit)
    accepted = t in {"1", "aceitar", "aceito", "1)"} or "aceitar" in t
    pendente = t in {"2", "pendente", "registre como pendente", "2)"} or "pendente" in t
    debate = t in {"3", "debater", "quero debater", "ajustar", "3)"} or "debater" in t
    
    if not (accepted or pendente or debate):
        # Ask for confirmation
        return (
            "Só para confirmar: prefere **1) aceitar**, **2) marcar como pendente**, "
            "ou **3) debater mais**?"
        )
    
    # Clear decision state
    state['awaiting_decision'] = False
    proposal = state.get('pending_proposal', '')
    decision_stage = state.get('decision_stage', '')
    state['pending_proposal'] = None
    state['decision_stage'] = None
    answers['state'] = state
    session.set_answers(answers)
    
    if accepted:
        logger.info(f"[{decision_stage.upper()}] decision=accept, advancing")
        return {"action": "accept", "text": proposal}
    elif pendente:
        logger.info(f"[{decision_stage.upper()}] decision=pendente, advancing")
        return {"action": "pendente", "text": proposal}
    else:
        logger.info(f"[{decision_stage.upper()}] decision=debate, staying at stage")
        return {"action": "debate", "text": ""}

def _strategy_selection_score(t: str, title: str) -> float:
    """Helper for _strategy_selection: calculates fuzzy score"""
    # simple jaccard over tokens
    sa, sb = set(re.findall(r'\w+', t.lower())), set(re.findall(r'\w+', title.lower()))
    inter = len(sa & sb)
    uni = len(sa | sb) or 1
    score = inter / uni
    # boost if user input is substring of title
    if t in title:
        score = max(score, 0.70)
    return score

def _strategy_selection_complete(text: str, strategies: list[str]) -> int | None:
    """Complete version of strategy selection"""
    if not text:
        return None
    t = text.strip().lower()
    
    # numeric selection
    m = re.match(r'^\s*([1-9])\s*$', t)
    if m:
        idx = int(m.group(1)) - 1
        return idx if 0 <= idx < len(strategies) else None
    
    best, bi = 0.0, None
    for i, title in enumerate(strategies):
        s = _score(t, title.lower())
        if s > best:
            best, bi = s, i
    return bi if best >= 0.60 else None

def _parse_strategy_selection(user_input: str, strategies: list) -> list:
    """
    Parse user selection of strategies.
    Supports:
    - Numeric selection: r'^\\s*([1-9])\\s*$' (1-9)
    - Text-based fuzzy matching with strategy titles (threshold ≥ 0.60)
    
    Returns list of selected strategies or None if no valid selection found.
    """
    from difflib import SequenceMatcher
    
    if not user_input or not strategies:
        return None
    
    user_input_stripped = user_input.strip()
    
    # Try numeric selection first
    numeric_pattern = r'^\s*([1-9])\s*$'
    match = re.match(numeric_pattern, user_input_stripped)
    if match:
        index = int(match.group(1)) - 1  # Convert to 0-based index
        if 0 <= index < len(strategies):
            logger.info(f"[STRATEGY_SELECT] Numeric selection: {index + 1}")
            return [strategies[index]]
        else:
            logger.info(f"[STRATEGY_SELECT] Numeric selection out of range: {index + 1}")
            return None
    
    # Try text-based fuzzy matching
    user_lower = user_input.lower().strip()
    best_match = None
    best_score = 0.0
    
    for strategy in strategies:
        titulo = strategy.get('titulo', '')
        if not titulo:
            continue
        
        titulo_lower = titulo.lower()
        
        # Calculate similarity
        similarity = SequenceMatcher(None, user_lower, titulo_lower).ratio()
        
        # Also check if user input is contained in title (partial match)
        if user_lower in titulo_lower:
            similarity = max(similarity, 0.70)  # Boost for substring match
        
        if similarity >= 0.60 and similarity > best_score:
            best_score = similarity
            best_match = strategy
    
    if best_match:
        logger.info(f"[STRATEGY_SELECT] Text match: '{user_input}' -> '{best_match.get('titulo')}' (score={best_score:.2f})")
        return [best_match]
    
    logger.info(f"[STRATEGY_SELECT] No match found for: '{user_input}'")
    return None

def _consult_path_explainer(necessity: str) -> tuple:
    """
    Gera alternativas contextualizadas a partir da necessidade usando LLM gpt-4.1.
    Retorna (texto_consultivo, lista_opcoes).
    """
    from application.ai.hybrid_models import OpenAIChatConsultive
    import logging
    import json
    import re
    
    logger = logging.getLogger(__name__)
    nec = (necessity or '').strip()
    
    prompt = f"""Você é um consultor de contratações públicas. Analise esta necessidade e sugira APENAS as alternativas realmente adequadas (não mais que 3):

Necessidade: "{nec}"

Para cada alternativa pertinente, forneça:
- Nome da modalidade (ex.: "Compra", "Serviço continuado", "Locação" apenas se fizer sentido)
- Prós e contras específicos para ESTA necessidade

IMPORTANTE:
- NÃO sugira locação para bens de consumo (alimentos, materiais descartáveis, insumos)
- NÃO sugira locação se o bem se exaure com uso
- Para gêneros alimentícios: considere apenas compra ou serviço de fornecimento continuado
- Para serviços: foque em terceirização ou execução direta
- Para equipamentos duráveis: pode incluir compra, locação ou serviço gerenciado

Retorne em formato JSON:
{{
  "alternativas": [
    {{"nome": "...", "pros": "...", "contras": "..."}},
    ...
  ]
}}"""

    try:
        llm = OpenAIChatConsultive()
        model_name = get_model_name()
        logger.info(f"[LM] model={model_name} temp=0.7 stage=solution_path_generation")
        raw = llm.generate(prompt)
        
        # Parse JSON do retorno
        json_match = re.search(r'\{.*\}', raw, re.DOTALL)
        if json_match:
            data = json.loads(json_match.group(0))
            alts = data.get('alternativas', [])
        else:
            alts = []
        
        if not alts:
            # Fallback simples contextual
            nec_l = _normalize(nec).lower()
            if any(k in nec_l for k in ['alimento', 'comida', 'gênero', 'merenda', 'refeição', 'alimentício']):
                alts = [
                    {"nome": "Compra direta", "pros": "Controle de qualidade e sazonalidade", "contras": "Gestão de estoque e validade"},
                    {"nome": "Serviço de fornecimento continuado", "pros": "Reposição regular e responsabilidade do fornecedor", "contras": "Menor controle sobre origem"}
                ]
            else:
                alts = [
                    {"nome": "Aquisição (compra)", "pros": "Controle do ativo", "contras": "CAPEX e manutenção própria"},
                    {"nome": "Serviço/terceirização", "pros": "Entrega por SLA", "contras": "Governança contratual intensa"}
                ]
        
        # Monta texto consultivo em prosa natural
        opcoes = []
        alt_texts = []
        for i, alt in enumerate(alts, start=1):
            nome = alt.get('nome', 'Opção')
            pros = alt.get('pros', '')
            contras = alt.get('contras', '')
            linha = f"{i}) {nome} — {pros}. Por outro lado, {contras.lower() if contras else 'requer atenção aos detalhes'}."
            alt_texts.append(linha)
            opcoes.append(nome)
        
        # Gerar parágrafo contextual usando LLM
        alternatives_text = "\n".join(alt_texts)
        context_prompt = f"""Para a necessidade "{nec}", apresente estas alternativas em 1-2 parágrafos de prosa natural, explicando por que cada uma é pertinente ao caso, sem usar headers ou comandos para o usuário:

{alternatives_text}

Escreva em tom consultivo e natural, finalizando de forma que o usuário se sinta confortável para escolher."""
        
        try:
            context_para = llm.generate(context_prompt)
            # Adicionar as opções numeradas após o contexto
            final_text = f"{context_para}\n\n{alternatives_text}"
        except Exception:
            # Fallback simples em prosa
            final_text = f'Para "{nec}", identifiquei estas alternativas que fazem sentido:\n\n{alternatives_text}\n\nCada uma tem suas vantagens específicas para seu contexto.'
        
        return final_text, opcoes
        
    except Exception as e:
        logger.error(f"Erro em _consult_path_explainer: {e}")
        # Fallback seguro
        return ("Precisamos escolher o caminho. Qual você prefere: compra ou serviço?", ["Compra", "Serviço"])

def _generate_prose_transition(stage: str, context: dict = None) -> str:
    """
    Gera transições em prosa natural usando LLM para cada etapa da conversa.
    Proibido usar: "Por que", "Como ajustar", "Você pode", cabeçalhos, menu rígido.
    """
    from application.ai.hybrid_models import OpenAIChatConsultive
    
    llm = OpenAIChatConsultive()
    context = context or {}
    
    prompts = {
        'confirm_to_solution_path': f"""O usuário acabou de confirmar os requisitos para a necessidade: "{context.get('necessity', '')}". 
        
Escreva 1-2 frases em prosa natural anunciando que agora vamos analisar as melhores alternativas para atender essa necessidade, sem usar comandos ou headers. Seja consultivo e direto.""",
        
        'solution_path_to_pca': f"""O usuário escolheu a alternativa "{context.get('chosen', '')}" para a contratação.
        
Escreva 1-2 frases em prosa natural perguntando sobre a previsão no PCA (Plano de Contratações Anual), aceitando incerteza. Não use formato de pergunta engessada.""",
        
        'pca_to_legal_norms': f"""O usuário respondeu sobre PCA: {context.get('pca_response', 'indefinido')}.
        
Escreva 1-2 frases em prosa natural fazendo a transição para normas legais, sugerindo opções comuns (Lei 14.133/21, regulamento local, etc.) mas aceitando que o usuário não saiba agora. Sem menu rígido.""",
        
        'legal_norms_to_qty': """O usuário informou as normas legais.
        
Escreva 1-2 frases em prosa natural perguntando sobre quantitativo e valor estimado, oferecendo ajuda caso não saiba (unidade, ordem de grandeza). Seja conversacional.""",
        
        'qty_to_installment': f"""O usuário informou: {context.get('qty_response', '')}.
        
Escreva 1-2 frases em prosa natural perguntando sobre parcelamento da contratação, explicando brevemente prós/contras sem formato didático. Seja natural."""
    }
    
    prompt = prompts.get(stage, "")
    if not prompt:
        return ""
    
    try:
        return llm.generate(prompt).strip()
    except Exception as e:
        logger.error(f"Erro ao gerar prosa para {stage}: {e}")
        # Fallback genérico
        return "Entendido. Vamos ao próximo ponto."

def _suggest_requirements(necessity: str) -> list:
    """Fallback de requisitos sugeridos quando o gerador não estiver disponível."""
    return [
        "Comprovação de experiência específica no escopo (projetos similares nos últimos 5 anos).",
        "Equipe dedicada com qualificações compatíveis e substituição sem prejuízo ao serviço.",
        "Conformidade com normas aplicáveis e regulatórias do setor.",
        "Plano de atendimento/manutenção com SLAs definidos.",
        "Integração com sistemas existentes e transferência de conhecimento/documentação.",
        "KPIs de desempenho e relatórios periódicos.",
        "Gestão de riscos e plano de contingência."
    ]

# ---------- HELPER: standardized response payload ----------
def _text_payload(session, text: str, extra: dict = None) -> dict:
    """
    Standardized response payload for all conversation endpoints.
    Always includes: success, kind, ai_response, message, conversation_stage, session_id
    """
    base = {
        'success': True,
        'kind': 'text',
        'ai_response': text,
        'message': text,
        'conversation_stage': getattr(session, 'conversation_stage', None),
        'session_id': getattr(session, 'session_id', None)
    }
    if extra:
        base.update(extra)
    return base

# ---------- HELPER: ensure session exists ----------
def ensure_session(sid: str = None, title: str = None) -> EtpSession:
    """
    Ensures that a session exists and its ID matches exactly what the client provided.
    Rules:
    - If sid is empty/None: generate a new sid and create the session.
    - If sid exists in DB: return that session.
    - If sid doesn't exist in DB: create session with exactly that sid.
    """
    if sid:
        s = EtpSession.query.filter_by(session_id=sid).first()
        if s:
            return s
        # Create new session with the exact sid provided by client
        s = EtpSession(
            user_id=1,
            session_id=sid,
            status='active',
            title=title or "Novo Estudo Técnico Preliminar",
            answers={},
            conversation_stage='collect_need',
            created_at=datetime.utcnow()
        )
        db.session.add(s)
        db.session.commit()
        return s
    # No sid provided → create new with generated UUID
    new_sid = str(uuid.uuid4())
    s = EtpSession(
        user_id=1,
        session_id=new_sid,
        status='active',
        title=title or "Novo Estudo Técnico Preliminar",
        answers={},
        conversation_stage='collect_need',
        created_at=datetime.utcnow()
    )
    db.session.add(s)
    db.session.commit()
    return s

def parse_update_command(user_message, current_requirements):
    """
    Parse user commands for requirement updates
    Returns dict with:
    - action: 'remove', 'edit', 'keep_only', 'add', 'unclear'
    - items: list of requirement IDs or content
    - message: explanation of what was done
    """
    import re
    
    message_lower = user_message.lower()
    
    # Check for explicit necessity restart keywords
    restart_keywords = [
        "nova necessidade", "trocar a necessidade", "na verdade a necessidade é",
        "mudou a necessidade", "preciso trocar a necessidade"
    ]
    
    for keyword in restart_keywords:
        if keyword in message_lower:
            return {
                'action': 'restart_necessity',
                'items': [],
                'message': 'Detectada solicitação para reiniciar necessidade'
            }
    
    # Extract requirement numbers (R1, R2, etc. or just numbers)
    req_numbers = []
    
    # Look for patterns like "R1", "R2", "requisito 1", "primeiro", "último", etc.
    r_patterns = re.findall(r'[rR](\d+)', user_message)
    req_numbers.extend([f"R{n}" for n in r_patterns])
    
    # Look for standalone numbers that might refer to requirements
    number_patterns = re.findall(r'\b(\d+)\b', user_message)
    for n in number_patterns:
        if int(n) <= len(current_requirements):
            req_id = f"R{n}"
            if req_id not in req_numbers:
                req_numbers.append(req_id)
    
    # Handle positional references
    if 'último' in message_lower or 'ultima' in message_lower:
        if current_requirements:
            req_numbers.append(current_requirements[-1].get('id', f"R{len(current_requirements)}"))
    
    if 'primeiro' in message_lower or 'primeira' in message_lower:
        if current_requirements:
            req_numbers.append(current_requirements[0].get('id', 'R1'))
    
    if 'penúltimo' in message_lower or 'penultima' in message_lower:
        if len(current_requirements) > 1:
            req_numbers.append(current_requirements[-2].get('id', f"R{len(current_requirements)-1}"))
    
    # Determine action type
    if any(word in message_lower for word in ['remover', 'tirar', 'excluir', 'deletar', 'retirar']):
        if req_numbers:
            return {
                'action': 'remove',
                'items': req_numbers,
                'message': f'Removidos requisitos: {", ".join(req_numbers)}'
            }
        else:
            return {'action': 'unclear', 'items': [], 'message': 'Não foi possível identificar quais requisitos remover'}
    
    if any(word in message_lower for word in ['manter apenas', 'só manter', 'manter só', 'manter somente']):
        if req_numbers:
            return {
                'action': 'keep_only',
                'items': req_numbers,
                'message': f'Mantidos apenas requisitos: {", ".join(req_numbers)}'
            }
        else:
            return {'action': 'unclear', 'items': [], 'message': 'Não foi possível identificar quais requisitos manter'}
    
    if any(word in message_lower for word in ['alterar', 'modificar', 'trocar', 'mudar', 'editar']):
        if req_numbers:
            return {
                'action': 'edit',
                'items': req_numbers,
                'message': f'Requisitos para edição: {", ".join(req_numbers)}'
            }
        else:
            return {'action': 'unclear', 'items': [], 'message': 'Não foi possível identificar quais requisitos alterar'}
    
    if any(word in message_lower for word in ['adicionar', 'incluir', 'acrescentar', 'novo requisito']):
        # Extract the content after the add command
        add_content = user_message
        for word in ['adicionar', 'incluir', 'acrescentar']:
            if word in message_lower:
                parts = user_message.lower().split(word, 1)
                if len(parts) > 1:
                    add_content = parts[1].strip()
                break
        
        return {
            'action': 'add',
            'items': [add_content],
            'message': f'Novo requisito adicionado: {add_content}'
        }
    
    # Check for confirmation words
    confirm_words = [
        'confirmar', 'confirmo', 'manter', 'ok', 'está bom', 'perfeito',
        'concordo', 'aceito', 'pode ser', 'sim'
    ]
    
    if any(word in message_lower for word in confirm_words):
        return {
            'action': 'confirm',
            'items': [],
            'message': 'Requisitos confirmados'
        }
    
    # If nothing clear was detected
    return {
        'action': 'unclear',
        'items': [],
        'message': 'Comando não reconhecido'
    }

def _get_etp_components():
    """Lazy initialization of ETP components to avoid circular imports"""
    global _etp_generator, _prompt_generator, _rag_system, _initialized
    
    if not _initialized:
        _etp_generator, _prompt_generator, _rag_system = init_etp_dynamic()
        _initialized = True
    
    return _etp_generator, _prompt_generator, _rag_system

# Initialize components for backward compatibility
def _ensure_initialized():
    """Ensure components are initialized"""
    global etp_generator, prompt_generator, rag_system
    etp_generator, prompt_generator, rag_system = _get_etp_components()

# Configure logging
logger = logging.getLogger(__name__)

# Perguntas do ETP conforme especificado
ETP_QUESTIONS = [
    {
        "id": 1,
        "question": "Qual a descrição da necessidade da contratação?",
        "type": "text",
        "required": True,
        "section": "OBJETO DO ESTUDO E ESPECIFICAÇÕES GERAIS"
    },
    {
        "id": 2,
        "question": "Você gostaria de manter esses requisitos, ajustar algum deles ou incluir outros?",
        "type": "text",
        "required": True,
        "section": "DESCRIÇÃO DOS REQUISITOS DA CONTRATAÇÃO"
    },
    {
        "id": 3,
        "question": "Possui demonstrativo de previsão no PCA?",
        "type": "boolean",
        "required": True,
        "section": "OBJETO DO ESTUDO E ESPECIFICAÇÕES GERAIS"
    },
    {
        "id": 4,
        "question": "Quais normas legais pretende utilizar?",
        "type": "text",
        "required": True,
        "section": "DESCRIÇÃO DOS REQUISITOS DA CONTRATAÇÃO"
    },
    {
        "id": 5,
        "question": "Qual o quantitativo e valor estimado?",
        "type": "text",
        "required": True,
        "section": "ESTIMATIVA DAS QUANTIDADES E VALORES"
    },
    {
        "id": 6,
        "question": "Haverá parcelamento da contratação?",
        "type": "boolean",
        "required": True,
        "section": "JUSTIFICATIVA PARA O PARCELAMENTO"
    }
]

def _text_payload(session, text, extra=None):
    """Helper para criar payload de resposta de texto padronizado"""
    payload = {
        'success': True,
        'kind': 'text',
        'ai_response': text,
        'message': text,
        'session_id': session.session_id,
        'conversation_stage': session.conversation_stage
    }
    if extra:
        payload.update(extra)
    return payload

def _parse_strategy_selection(user_text: str, strategies: list) -> list:
    """
    Interpreta a seleção de estratégia pelo usuário.
    Aceita: número puro ("2"), título parcial ("outsourcing"), ou múltiplos ("1 e 3", "2,3")
    
    Returns:
        Lista de estratégias selecionadas ou None se não houver correspondência
    """
    if not strategies or not user_text:
        return None
    
    text = (user_text or "").strip().lower()
    
    # 1) Número puro (ex: "2")
    if text.isdigit():
        idx = int(text) - 1
        if 0 <= idx < len(strategies):
            logger.info(f"[STRATEGY_SELECT] Selected by number: {idx+1}")
            return [strategies[idx]]
    
    # 2) Busca por título aproximado
    for s in strategies:
        title = (s.get("titulo") or s.get("title") or "").lower()
        if title:
            # Match por substring inicial ou palavras-chave
            if title[:20] in text or any(word in text for word in title.split()[:3] if len(word) > 3):
                logger.info(f"[STRATEGY_SELECT] Selected by title match: {s.get('titulo')}")
                return [s]
    
    # 3) Múltiplas seleções (ex.: "1 e 3", "2,3", "1 2")
    import re
    nums = [int(n)-1 for n in re.findall(r"\b\d+\b", text)]
    if nums:
        picked = [strategies[i] for i in nums if 0 <= i < len(strategies)]
        if picked:
            logger.info(f"[STRATEGY_SELECT] Selected multiple: {[s.get('titulo') for s in picked]}")
            return picked
    
    return None

@etp_dynamic_bp.route('/health', methods=['GET'])
@cross_origin()
def health_check():
    """Verificação de saúde da API dinâmica e do gerador de requisitos"""
    try:
        _ensure_initialized()
        # Verificar se os geradores estão funcionando
        kb_info = etp_generator.get_knowledge_base_info() if etp_generator else {}
        
        # Check simple generator status
        try:
            simple_gen = _get_simple_generator()
            generator_name = simple_gen.__class__.__name__
            generator_ready = True
        except Exception as gen_error:
            generator_name = f"Error: {str(gen_error)}"
            generator_ready = False

        return jsonify({
            'status': 'healthy',
            'version': '3.0.0-dynamic',
            'openai_configured': bool(os.getenv('OPENAI_API_KEY')),
            'etp_generator_ready': etp_generator is not None,
            'simple_generator': generator_name,
            'simple_generator_ready': generator_ready,
            'knowledge_base': {
                'documents_loaded': kb_info.get('total_documents', 0),
                'common_sections': len(kb_info.get('common_sections', []))
            },
            'timestamp': datetime.now().isoformat()
        })
    except Exception as e:
        return jsonify({
            'status': 'error',
            'error': str(e),
            'timestamp': datetime.now().isoformat()
        }), 500

@etp_dynamic_bp.route('/new', methods=['POST'])
@cross_origin()
def create_new_conversation():
    """Create a new conversation with persistent storage and fresh EtpSession"""
    try:
        user_id = get_current_user_id()
        
        # Generate sequential title with timestamp
        from datetime import datetime
        title = f"Novo Documento {datetime.now().strftime('%d/%m %H:%M')}"
        
        # Create conversation in database
        conv = ConversationRepo.create(user_id=user_id, title=title)
        
        # Create a new EtpSession with fresh state using the conversation_id as session_id
        # This ensures state is reset for the new document
        new_session_id = str(conv.id)  # Use conversation ID as session ID
        new_session = ensure_session(sid=new_session_id, title=title)
        
        # Ensure the session starts fresh
        new_session.necessity = None
        new_session.set_requirements([])
        new_session.set_answers({})
        new_session.conversation_stage = 'collect_need'
        new_session.updated_at = datetime.utcnow()
        
        db.session.commit()
        
        logger.info(f"[NEW_DOCUMENT] Created conversation {conv.id} and session {new_session_id} for user: {user_id}")
        
        return jsonify({
            'success': True,
            'conversation_id': conv.id,
            'session_id': new_session_id,
            'title': conv.title,
            'conversation_stage': 'collect_need',
            'created_at': conv.created_at.isoformat(),
            'updated_at': conv.updated_at.isoformat()
        }), 201
        
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error creating new conversation: {e}")
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': f'Erro ao criar nova conversa: {str(e)}'
        }), 500

@etp_dynamic_bp.route('/list', methods=['GET'])
@cross_origin()
def list_conversations():
    """List all conversations for the current user with message preview"""
    try:
        user_id = get_current_user_id()
        
        # Get conversations ordered by updated_at DESC
        conversations_list = ConversationRepo.list_by_user(user_id=user_id, limit=100)
        
        result = []
        for conv in conversations_list:
            # Get last message for preview
            last_msg = MessageRepo.get_last_message(conv.id)
            preview = ""
            if last_msg:
                preview = last_msg.content[:120] + "..." if len(last_msg.content) > 120 else last_msg.content
            
            result.append({
                'id': conv.id,
                'title': conv.title,
                'preview': preview,
                'updated_at': conv.updated_at.isoformat(),
                'created_at': conv.created_at.isoformat()
            })
        
        logger.info(f"[CONVERSATION] Listed {len(result)} conversations for user: {user_id}")
        
        return jsonify({
            'success': True,
            'conversations': result
        }), 200
        
    except Exception as e:
        logger.error(f"Error listing conversations: {e}")
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': f'Erro ao listar conversas: {str(e)}'
        }), 500

@etp_dynamic_bp.route('/chat-stage', methods=['POST'])
@cross_origin()
def chat_stage_based():
    """
    Stage-based chat endpoint with deterministic FSM and RAG integration.
    Implements the 9-stage flow: collect_need → suggest_requirements → solution_path → ... → preview
    No intermediate refine_requirements stage - edits happen within suggest_requirements.
    """
    try:
        data = request.get_json()
        conversation_id = data.get('conversation_id')
        user_message = data.get('message', '').strip()
        
        if not conversation_id:
            return jsonify({'success': False, 'error': 'conversation_id é obrigatório'}), 400
        
        if not user_message:
            return jsonify({'success': False, 'error': 'Mensagem é obrigatória'}), 400
        
        # Load conversation and session
        conv = ConversationRepo.get(conversation_id)
        if not conv:
            return jsonify({'success': False, 'error': 'Conversa não encontrada'}), 404
        
        session = EtpSession.query.filter_by(session_id=str(conversation_id)).first()
        if not session:
            session = ensure_session(sid=str(conversation_id))
        
        current_stage = session.conversation_stage or 'collect_need'
        necessity = session.necessity or ''
        requirements = session.get_requirements() or []
        answers = session.get_answers() or {}
        
        logger.info(f"[STAGE_CHAT] conversation={conversation_id}, stage={current_stage}, message={user_message[:50]}")
        
        # Save user message
        MessageRepo.add(
            conversation_id=conversation_id,
            role='user',
            content=user_message,
            stage=current_stage
        )
        
        # Process based on stage
        next_stage = current_stage
        ai_response = ""
        generator_output = {}
        
        if current_stage == 'collect_need':
            # User provided necessity - generate requirements immediately
            session.necessity = user_message
            necessity = user_message
            next_stage = 'suggest_requirements'
            
            # Generate requirements with RAG using new generate_answer
            from rag.retrieval import retrieve_for_stage
            from application.ai.generator import generate_answer
            
            context_chunks = retrieve_for_stage(necessity, 'suggest_requirements', k=12)
            
            # Build RAG context
            rag_context = {
                'chunks': context_chunks,
                'necessity': necessity
            }
            
            # Call generate_answer for collect_need stage
            result = generate_answer('collect_need', [], user_message, rag_context)
            
            # Extract response components
            intro = result.get('intro', '')
            requirements = result.get('requirements', [])
            justification = result.get('justification', '')
            
            # Deduplicate requirements before storing
            requirements = dedupe_requirements(requirements)

            # Clean suffix markers such as (Obrigatório) and normalize spacing
            requirements = [
                re.sub(r"\s*\(.*?\)", "", _extract_requirement_text(req)).strip()
                for req in requirements
            ]

            # Store requirements
            session.set_requirements(requirements)
            
            # Build response without hardcoded text
            response_parts = []
            if intro:
                response_parts.append(intro)
            if requirements:
                response_parts.append("\n".join(requirements))
            if justification:
                response_parts.append(f"\n{justification}")
            
            ai_response = "\n\n".join(response_parts) if response_parts else ""

            followup_text = _requirements_followup_text()
            if followup_text:
                ai_response = f"{ai_response}\n\n{followup_text}" if ai_response else followup_text

            answers = session.get_answers() or {}
            state = answers.get('state', {})
            state['awaiting_requirements_feedback'] = True
            state['awaiting_requirements_edit'] = False
            state['requirements_current'] = list(requirements)
            answers['state'] = state
            session.set_answers(answers)
            logger.info("[REQUIREMENTS] Follow-up asked")

            # Store in generator_output for compatibility
            generator_output = result
        
        elif current_stage == 'suggest_requirements':
            answers = session.get_answers() or {}
            state = answers.get('state', {})
            awaiting_feedback = bool(state.get('awaiting_requirements_feedback'))
            awaiting_edit = bool(state.get('awaiting_requirements_edit'))
            need_context = session.necessity or state.get('need_description') or ''

            intent = detect_requirements_intent(user_message)
            logger.info(f"[REQUIREMENTS] Intent={intent}")

            if awaiting_edit:
                updated_model_list = handle_requirements_edit(user_message, requirements, need_context)
                if updated_model_list != requirements:
                    session.set_requirements(updated_model_list)
                    requirements = updated_model_list
                    answers['requirements'] = list(updated_model_list)
                    state['awaiting_requirements_edit'] = False
                    state['awaiting_requirements_feedback'] = True
                    state['requirements_current'] = list(updated_model_list)
                    answers['state'] = state
                    session.set_answers(answers)
                    formatted = _format_requirements_list(updated_model_list)
                    ai_response = (
                        "Aqui está a lista atualizada de requisitos conforme solicitado:\n\n"
                        f"{formatted}\n\n"
                        "Deseja ajustar mais algum item ou posso seguir para as estratégias de solução?"
                    )
                    next_stage = 'suggest_requirements'
                    logger.info("[REQUIREMENTS] Model-based edit applied during awaiting_edit state")
                elif intent == REQ_ACCEPT:
                    state['awaiting_requirements_edit'] = False
                    state['awaiting_requirements_feedback'] = False
                    answers['state'] = state
                    session.set_answers(answers)
                    next_stage = 'solution_strategies'
                    ai_response = "Perfeito. Vou seguir para a próxima etapa."
                else:
                    updated_list, edit_message = _apply_requirements_edit(requirements, user_message)
                    if updated_list is None:
                        ai_response = edit_message
                        next_stage = 'suggest_requirements'
                    else:
                        session.set_requirements(updated_list)
                        requirements = updated_list
                        answers['requirements'] = list(updated_list)
                        state['awaiting_requirements_edit'] = False
                        state['awaiting_requirements_feedback'] = True
                        state['requirements_current'] = list(updated_list)
                        answers['state'] = state
                        session.set_answers(answers)
                        formatted = _format_requirements_list(updated_list)
                        ai_response = (
                            f"{edit_message}\n\n{formatted}\n\n"
                            "Queremos deixar isso com a sua cara. Ficou bom assim ou prefere ajustar mais algum item?"
                        )
                        logger.info(f"[REQUIREMENTS] Edit operation={edit_message}")
                        logger.info("[REQUIREMENTS] Follow-up asked")
                        next_stage = 'suggest_requirements'
            else:
                updated_model_list = handle_requirements_edit(user_message, requirements, need_context)
                if updated_model_list != requirements:
                    session.set_requirements(updated_model_list)
                    requirements = updated_model_list
                    answers['requirements'] = list(updated_model_list)
                    state['awaiting_requirements_feedback'] = True
                    state['awaiting_requirements_edit'] = False
                    state['requirements_current'] = list(updated_model_list)
                    answers['state'] = state
                    session.set_answers(answers)
                    formatted = _format_requirements_list(updated_model_list)
                    ai_response = (
                        "Aqui está a lista atualizada de requisitos conforme solicitado:\n\n"
                        f"{formatted}\n\n"
                        "Deseja ajustar mais algum item ou posso seguir para as estratégias de solução?"
                    )
                    logger.info("[REQUIREMENTS] Model-based edit applied")
                    next_stage = 'suggest_requirements'
                elif intent == REQ_ACCEPT:
                    state['awaiting_requirements_feedback'] = False
                    state['awaiting_requirements_edit'] = False
                    answers['state'] = state
                    session.set_answers(answers)
                    next_stage = 'solution_strategies'
                    ai_response = "Perfeito. Vou seguir para a próxima etapa."
                elif intent == REQ_EDIT:
                    updated_list, edit_message = _apply_requirements_edit(requirements, user_message)
                    if updated_list is None:
                        state['awaiting_requirements_feedback'] = False
                        state['awaiting_requirements_edit'] = True
                        answers['requirements'] = list(requirements)
                        state['requirements_current'] = list(requirements)
                        answers['state'] = state
                        session.set_answers(answers)
                        ai_response = (
                            "Certo. Me diga o que você quer ajustar. Posso trocar um item pelo número, "
                            "remover um específico ou adicionar novos requisitos."
                        )
                        next_stage = 'suggest_requirements'
                    else:
                        session.set_requirements(updated_list)
                        requirements = updated_list
                        answers['requirements'] = list(updated_list)
                        state['awaiting_requirements_feedback'] = True
                        state['awaiting_requirements_edit'] = False
                        state['requirements_current'] = list(updated_list)
                        answers['state'] = state
                        session.set_answers(answers)
                        formatted = _format_requirements_list(updated_list)
                        ai_response = (
                            f"{edit_message}\n\n{formatted}\n\n"
                            "Queremos deixar isso com a sua cara. Ficou bom assim ou prefere ajustar mais algum item?"
                        )
                        logger.info(f"[REQUIREMENTS] Edit operation={edit_message}")
                        logger.info("[REQUIREMENTS] Follow-up asked")
                        next_stage = 'suggest_requirements'
                else:
                    if not awaiting_feedback:
                        state['awaiting_requirements_feedback'] = True
                        state['requirements_current'] = list(requirements)
                        answers['requirements'] = list(requirements)
                        answers['state'] = state
                        session.set_answers(answers)
                        logger.info("[REQUIREMENTS] Follow-up asked")
                    ai_response = (
                        "Você quer manter como está ou prefere ajustar algum item? "
                        "Se quiser, posso sugerir alternativas."
                    )
                    next_stage = 'suggest_requirements'
        
        elif current_stage == 'refine_requirements_assist':
            # Handle refinement commands
            user_lower = user_message.lower()
            
            if 'refinar' in user_lower or 'refinar automaticamente' in user_lower:
                # User requested automatic refinement
                from rag.retrieval import retrieve_for_stage
                context = retrieve_for_stage(necessity, 'refine_requirements_assist', k=12)
                
                generator = _get_simple_generator()
                if hasattr(generator, 'refine_requirements'):
                    refinement_result = generator.refine_requirements(requirements, necessity, context)
                    
                    if refinement_result.get('success', False):
                        refined_reqs = refinement_result.get('requirements', [])
                        mudancas = refinement_result.get('mudancas', [])
                        
                        session.set_requirements(refined_reqs)
                        requirements = refined_reqs
                        
                        mudancas_text = "\n".join([f"- {m}" for m in mudancas[:5]]) if mudancas else "Nenhuma mudança necessária"
                        
                        # Re-validate
                        generator_output = generator.generate('refine_requirements', necessity, context, {'requirements': refined_reqs})
                        is_blocked = generator_output.get('blocked', False)
                        
                        if not is_blocked:
                            next_stage = 'solution_path'
                            
                            # Generate solution path
                            context_sol = retrieve_for_stage(necessity, 'solution_path', k=12)
                            sol_output = generator.generate('solution_path', necessity, context_sol, {'requirements': refined_reqs})
                            steps = sol_output.get('steps', [])
                            answers['solution_path'] = steps
                            session.set_answers(answers)
                            
                            steps_text = "\n".join([f"{i+1}. {s}" for i, s in enumerate(steps)])
                            ai_response = f"✅ Requisitos refinados com sucesso!\n\nMudanças aplicadas:\n{mudancas_text}\n\n**Caminho da solução:**\n\n{steps_text}\n\nConfirme para prosseguir."
                        else:
                            ai_response = f"Requisitos refinados, mas ainda há problemas:\n\n{generator_output.get('message', 'Verifique os requisitos.')}\n\nMudanças aplicadas:\n{mudancas_text}"
                            next_stage = 'refine_requirements_assist'
                    else:
                        ai_response = f"Erro ao refinar requisitos: {refinement_result.get('error', 'Erro desconhecido')}"
                        next_stage = 'refine_requirements_assist'
                else:
                    ai_response = "Funcionalidade de refinamento não disponível."
                    next_stage = 'refine_requirements_assist'
            
            elif CONFIRM_RE.search(user_message) or 'pode seguir' in user_lower:
                # User confirmed after validation passed - move to solution_path
                next_stage = 'solution_path'
                
                # Generate solution path
                from rag.retrieval import retrieve_for_stage
                context = retrieve_for_stage(necessity, 'solution_path', k=12)
                
                generator = _get_simple_generator()
                if hasattr(generator, 'generate'):
                    generator_output = generator.generate('solution_path', necessity, context, {'requirements': requirements})
                    steps = generator_output.get('steps', [])
                    answers['solution_path'] = steps
                    session.set_answers(answers)
                    
                    steps_text = "\n".join([f"{i+1}. {s}" for i, s in enumerate(steps)])
                    ai_response = f"Caminho da solução:\n\n{steps_text}\n\nConfirme para prosseguir."
                else:
                    ai_response = "Definindo caminho da solução..."
            
            elif 'editar:' in user_lower or 'adicionar:' in user_lower or 'remover:' in user_lower:
                # User wants to manually edit requirements - go back to suggest_requirements stage
                next_stage = 'suggest_requirements'
                ai_response = "Ok, voltando para edição de requisitos. Use os comandos: adicionar:, remover:, editar:"
            
            else:
                # Re-show the validation message
                from rag.retrieval import retrieve_for_stage
                context = retrieve_for_stage(necessity, 'refine_requirements_assist', k=12)
                
                generator = _get_simple_generator()
                if hasattr(generator, 'generate'):
                    generator_output = generator.generate('refine_requirements', necessity, context, {'requirements': requirements})
                    ai_response = generator_output.get('message', 'Posso refinar os requisitos automaticamente ou você pode editá-los manualmente.')
                else:
                    ai_response = "Posso refinar os requisitos automaticamente ou você pode editá-los manualmente. Como prefere?"
                next_stage = 'refine_requirements_assist'
        
        elif current_stage == 'solution_strategies':
            # Generate contracting strategies (not ETP steps)
            from rag.retrieval import retrieve_for_stage
            
            # Check if we already have strategies stored
            stored_strategies = answers.get('strategies', [])
            selected_strategies = answers.get('selected_strategies', [])
            
            # If strategies already exist, try to parse selection
            if stored_strategies:
                logger.info(f"[STRATEGY] Already have {len(stored_strategies)} strategies, checking for selection")
                
                # Check for vague acknowledgment first
                if intents.is_vague_ack(user_message):
                    logger.info(f"[STRATEGY] Vague acknowledgment detected, staying at stage: {user_message[:30]}")
                    ai_response = "Qual dessas estratégias faz mais sentido para o seu caso? Pode me dizer o número ou o nome da que preferir."
                    next_stage = 'solution_strategies'
                else:
                    # Extract strategy titles for selection matching
                    strategy_titles = [s.get('titulo', '') for s in stored_strategies]
                    
                    # Try to parse user selection using new helper
                    sel_idx = _strategy_selection(user_message, strategy_titles)
                    
                    if sel_idx is not None:
                        # User selected a strategy - save and advance
                        selected_strategy = stored_strategies[sel_idx]
                        answers['selected_strategies'] = [selected_strategy]
                        answers['strategy_selected'] = sel_idx
                        session.set_answers(answers)
                        db.session.commit()
                        
                        selected_title = selected_strategy.get('titulo', 'Estratégia')
                        ai_response = f"Perfeito. Seguiremos com: {selected_title}. Agora vamos tratar do PCA (Plano de Contratações Anual)."
                        next_stage = 'pca'
                        logger.info(f"[STRATEGY] Selection detected (index={sel_idx}), advancing to {next_stage}")
                    else:
                        # No valid selection detected
                        # Do NOT advance on free confirmation - stay in stage
                        ai_response = "Não consegui identificar qual estratégia você escolheu. Pode me indicar qual prefere? Use o número ou o nome dela."
                        next_stage = 'solution_strategies'
                        logger.info(f"[STRATEGY] No valid selection, staying in stage")
            else:
                # First time in this stage - generate strategies
                context = retrieve_for_stage(necessity, 'solution_strategies', k=8)
                logger.info(
                    "[RAG:USED n=%s] stage=solution_strategies necessity='%s'",
                    len(context),
                    (necessity or '')[:120]
                )
                
                # Build history
                history = []
                try:
                    messages = MessageRepo.list_for_conversation(conversation_id) or []
                except Exception:
                    messages = []
                for msg in messages[-10:]:
                    history.append({'role': msg.role, 'content': msg.content})
                
                # Call new generate_answer for solution_strategies
                from application.ai.generator import generate_answer
                
                # Format requirements as simple strings
                req_strings = []
                for req in requirements:
                    if isinstance(req, dict):
                        req_strings.append(req.get('text') or req.get('descricao') or str(req))
                    else:
                        req_strings.append(str(req))
                
                rag_context = {
                    'chunks': context,
                    'necessity': necessity,
                    'requirements': req_strings
                }
                
                result = generate_answer('solution_strategies', history, user_message, rag_context)

                diagnostics = result.get('diagnostics', {})
                fallback_used = diagnostics.get('fallback_used')
                context_chunks = diagnostics.get('context_chunks', len(context))
                sectors = diagnostics.get('sector_tags') or []
                object_nature = diagnostics.get('object_nature') or 'aquisição ou contratação'
                if fallback_used:
                    logger.info(
                        "[STRATEGY:DEFAULT_FALLBACK_USED] stage=solution_strategies context_chunks=%s necessity='%s'",
                        context_chunks,
                        (necessity or '')[:120]
                    )
                else:
                    logger.info(
                        "[STRATEGY:CONTEXTUAL_GENERATION] stage=solution_strategies context_chunks=%s nature=%s sectors=%s necessity='%s'",
                        context_chunks,
                        object_nature,
                        ', '.join(sectors) if sectors else 'não identificado',
                        (necessity or '')[:120]
                    )

                # Diagnostics are only for logging
                result.pop('diagnostics', None)

                # Extract strategies
                strategies = result.get('strategies', [])
                intro = result.get('intro', '')
                
                # Format strategies for display
                if strategies:
                    strat_lines = []
                    for i, strat in enumerate(strategies, 1):
                        titulo = strat.get('titulo', f'Estratégia {i}')
                        quando = strat.get('quando_indicado', '')
                        vantagens = strat.get('vantagens', [])
                        riscos = strat.get('riscos', [])
                        
                        strat_text = f"\n**{i}. {titulo}**"
                        if quando:
                            strat_text += f"\n   Quando indicado: {quando}"
                        if vantagens:
                            strat_text += f"\n   Vantagens: {', '.join(vantagens)}"
                        if riscos:
                            strat_text += f"\n   Riscos/Cuidados: {', '.join(riscos)}"
                        
                        strat_lines.append(strat_text)
                    
                    strat_text_full = "\n".join(strat_lines)
                    if intro:
                        ai_response = f"{intro}\n{strat_text_full}"
                    else:
                        ai_response = strat_text_full
                else:
                    ai_response = intro if intro else "Estratégias de contratação geradas."
                
                # Save strategies to session for next interaction
                answers['strategies'] = strategies
                session.set_answers(answers)
                db.session.commit()
                
                # Stay in same stage to wait for selection
                next_stage = 'solution_strategies'
                logger.info(f"[STRATEGY] Generated {len(strategies)} strategies, waiting for selection")
        
        elif current_stage == 'pca':
            # First, try to consume pending decision
            decision = try_consume_decision(session, user_message)
            if isinstance(decision, str):
                # Still asking for confirmation
                ai_response = decision
                next_stage = current_stage  # Stay in same stage
            elif isinstance(decision, dict):
                # Decision made
                if decision["action"] == "accept":
                    answers['pca'] = f'Sugestão aceita: {decision["text"]}'
                    session.set_answers(answers)
                    next_stage = 'legal_norms'
                    ai_response = "Perfeito. Agora sobre normas e base legal: posso te sugerir um pacote inicial típico do setor e você me diz se mantemos, ajustamos ou deixamos como rascunho?"
                elif decision["action"] == "pendente":
                    answers['pca'] = f'Pendente (com justificativa): {decision["text"]}'
                    session.set_answers(answers)
                    next_stage = 'legal_norms'
                    ai_response = "Ok, registrado como Pendente (com justificativa). Agora sobre normas e base legal: posso te sugerir um pacote inicial típico do setor e você me diz se mantemos, ajustamos ou deixamos como rascunho?"
                elif decision["action"] == "debate":
                    next_stage = current_stage  # Stay in same stage
                    ai_response = "Vamos debater o PCA. A necessidade é emergente (inclusão extraordinária) ou pode aguardar o próximo ciclo anual?"
            else:
                # No decision pending, process normally
                # Check for vague acknowledgment first
                if intents.is_vague_ack(user_message):
                    logger.info(f"[PCA] Vague acknowledgment detected, staying at stage: {user_message[:30]}")
                    ai_response = "Sobre o PCA (Plano de Contratações Anual): essa demanda já aparece no seu planejamento deste ano? Se não tiver certeza, eu explico rapidamente e te proponho dois caminhos simples."
                    next_stage = current_stage  # Stay in same stage
                
                # Check for uncertainty using robust intent detection
                elif intents.is_uncertain_value(user_message):
                    logger.info(f"[PCA] Uncertainty detected, offering help: {user_message[:30]}")
                    # Build PCA suggestion with three-option pattern
                    proposal_text = """Sugestão: Inclusão extraordinária no PCA - necessidade emergente não prevista.
Justificativa: permite prosseguir sem aguardar próximo ciclo anual, conforme urgência da demanda."""
                    
                    prompt_text = """Você sinalizou que não tem certeza sobre o PCA (Plano de Contratações Anual). O PCA prevê as contratações do órgão para o ano.

**Opções comuns:**
1. **Inclusão extraordinária** - se a necessidade é emergente e não estava prevista
2. **Agendar para próximo ciclo** - se pode aguardar o próximo PCA
3. **Condicionar a PT específico** - se depende de aprovação orçamentária

Posso sugerir **inclusão extraordinária** como padrão. Preferir **aceitar**, **pendente** ou **debater**?"""
                    
                    ai_response = ask_user_decision(session, prompt_text, proposal_text, 'pca')
                    next_stage = current_stage  # Don't advance yet - wait for decision
                
                # Check for explicit pending request
                elif intents.is_explicit_pending_request(user_message):
                    logger.info(f"[PCA] Explicit pending request: {user_message[:30]}")
                    answers['pca'] = 'Pendente (a definir posteriormente)'
                    session.set_answers(answers)
                    next_stage = 'legal_norms'
                    ai_response = "Ok, registrado como Pendente. Agora sobre normas e base legal: posso te sugerir um pacote inicial típico do setor e você me diz se mantemos, ajustamos ou deixamos como rascunho?"
                
                else:
                    # Valid substantive input
                    if user_message.strip():
                        logger.info(f"[PCA] Storing substantive input: {user_message[:30]}")
                        answers['pca'] = user_message
                        session.set_answers(answers)
                        next_stage = 'legal_norms'
                        ai_response = "Normas e base legal: posso te sugerir um pacote inicial típico do setor (Lei 14.133/2021 + regulatórias aplicáveis) e você me diz se mantemos, ajustamos ou deixamos como rascunho?"
                    else:
                        # Empty input
                        logger.warning(f"[PCA] Empty input received, staying at stage")
                        ai_response = "Para prosseguir, preciso de informações sobre o PCA. Diga 'não sei' se precisar de ajuda."
                        next_stage = current_stage
        
        elif current_stage == 'legal_norms':
            # First, try to consume pending decision
            decision = try_consume_decision(session, user_message)
            if isinstance(decision, str):
                # Still asking for confirmation
                ai_response = decision
                next_stage = current_stage  # Stay in same stage
            elif isinstance(decision, dict):
                # Decision made
                if decision["action"] == "accept":
                    answers['legal_norms'] = f'Sugestão aceita: {decision["text"]}'
                    session.set_answers(answers)
                    next_stage = 'qty_value'
                    ai_response = "Perfeito, sigo com as normas sugeridas. Vamos estimar quantitativo/valor..."
                elif decision["action"] == "pendente":
                    answers['legal_norms'] = f'Pendente (com justificativa): {decision["text"]}'
                    session.set_answers(answers)
                    next_stage = 'qty_value'
                    ai_response = "Registro feito como Pendente (com justificativa). Vamos estimar quantitativo/valor..."
                elif decision["action"] == "debate":
                    next_stage = current_stage  # Stay in same stage
                    ai_response = "Sem problemas — seguimos debatendo as normas. Quer iniciar por base legal federal (Lei 14.133/2021 + Decreto 11.462/2023) e regulatório setorial?"
            else:
                # No decision pending, process normally
                # Check for vague acknowledgment first
                if intents.is_vague_ack(user_message):
                    logger.info(f"[LEGAL_NORMS] Vague acknowledgment detected, staying at stage: {user_message[:30]}")
                    ai_response = "Normas e base legal: posso te sugerir um pacote inicial típico do setor (Lei 14.133/2021 + regulatórias aplicáveis) e você me diz se mantemos, ajustamos ou deixamos como rascunho?"
                    next_stage = current_stage  # Stay in same stage
                
                # Check for uncertainty using robust intent detection
                elif intents.is_uncertain_value(user_message):
                    logger.info(f"[LEGAL_NORMS] Uncertainty detected, offering help: {user_message[:30]}")
                    # Build sector-specific norm suggestions
                    necessity_lower = necessity.lower() if necessity else ""
                    
                    sector_norms = []
                    if any(word in necessity_lower for word in ['aeronave', 'avião', 'aviação', 'helicóptero']):
                        sector_norms = ["Regulamentos ANAC (RBAC aplicáveis)", "Normas de segurança aeronáutica"]
                    elif any(word in necessity_lower for word in ['telecom', 'telecomunicação', 'internet', 'rede']):
                        sector_norms = ["Regulamentos ANATEL", "Normas INMETRO para equipamentos de telecomunicações"]
                    elif any(word in necessity_lower for word in ['software', 'sistema', 'ti', 'tecnologia']):
                        sector_norms = ["LGPD (Lei 13.709/2018) se houver dados pessoais", "Normas ABNT NBR ISO/IEC para segurança da informação"]
                    elif any(word in necessity_lower for word in ['construção', 'obra', 'reforma', 'edificação']):
                        sector_norms = ["ABNT NBR aplicáveis", "Código de Obras local", "NRs de segurança do trabalho"]
                    
                    sector_text = "\n".join([f"- {norm}" for norm in sector_norms]) if sector_norms else "- Normas técnicas ABNT específicas do setor"
                    
                    proposal_text = f"""Pacote-base sugerido:
Obrigatórias: Lei 14.133/2021; Decreto 11.462/2023
Setoriais: {', '.join(sector_norms) if sector_norms else 'Normas técnicas ABNT específicas do setor'}
Justificativa: asseguram conformidade licitatória e regulatória mínima."""
                    
                    prompt_text = f"""Você indicou que não sabe quais normas aplicar. Segue uma **sugestão inicial**:

**Obrigatórias:**
- Lei 14.133/2021 (Nova Lei de Licitações)
- Decreto 11.462/2023 (Regulamenta a Lei 14.133/2021)

**De referência/setoriais para seu caso:**
{sector_text}

Podemos **aceitar** essa sugestão, **marcar como Pendente** (com justificativa), ou **debater** mais."""
                    
                    ai_response = ask_user_decision(session, prompt_text, proposal_text, 'legal_norms')
                    next_stage = current_stage  # Don't advance yet
                
                # Check for explicit pending request
                elif intents.is_explicit_pending_request(user_message):
                    logger.info(f"[LEGAL_NORMS] Explicit pending request: {user_message[:30]}")
                    answers['legal_norms'] = 'Pendente (a definir posteriormente)'
                    session.set_answers(answers)
                    next_stage = 'qty_value'
                    ai_response = "Ok, registrado como Pendente. Sobre quantitativo e valor: dá para chutar uma ordem de grandeza? Se estiver nebuloso, eu te mostro duas formas rápidas de chegar a um número defensável e já deixo uma faixa inicial para você aprovar."
                
                elif 'adotar todas' in user_message.lower() or 'adotar todos' in user_message.lower():
                    # User wants to adopt all suggested norms
                    logger.info(f"[LEGAL_NORMS] Adopting all suggested norms")
                    answers['legal_norms'] = user_message + " (todas as normas sugeridas adotadas)"
                    session.set_answers(answers)
                    next_stage = 'qty_value'
                    ai_response = "Normas adotadas. Sobre quantitativo e valor: dá para chutar uma ordem de grandeza? Se estiver nebuloso, eu te mostro duas formas rápidas de chegar a um número defensável e já deixo uma faixa inicial para você aprovar."
                
                elif 'adotar' in user_message.lower() and any(char.isdigit() for char in user_message):
                    # User selecting specific norms by number (e.g., "adotar 1 e 3")
                    logger.info(f"[LEGAL_NORMS] Adopting specific norms: {user_message[:30]}")
                    answers['legal_norms'] = user_message
                    session.set_answers(answers)
                    next_stage = 'qty_value'
                    ai_response = "Normas selecionadas adotadas. Sobre quantitativo e valor: dá para chutar uma ordem de grandeza? Se estiver nebuloso, eu te mostro duas formas rápidas de chegar a um número defensável e já deixo uma faixa inicial para você aprovar."
                
                elif 'pular' in user_message.lower():
                    # User wants to skip this stage
                    logger.info(f"[LEGAL_NORMS] User chose to skip")
                    answers['legal_norms'] = 'Pulado - a definir posteriormente'
                    session.set_answers(answers)
                    next_stage = 'qty_value'
                    ai_response = "Ok, pulando normas. Sobre quantitativo e valor: dá para chutar uma ordem de grandeza? Se estiver nebuloso, eu te mostro duas formas rápidas de chegar a um número defensável e já deixo uma faixa inicial para você aprovar."
                
                else:
                    # Valid substantive input
                    if user_message.strip():
                        logger.info(f"[LEGAL_NORMS] Storing substantive input: {user_message[:30]}")
                        answers['legal_norms'] = user_message
                        session.set_answers(answers)
                        next_stage = 'qty_value'
                        ai_response = "Sobre quantitativo e valor: dá para chutar uma ordem de grandeza? Se estiver nebuloso, eu te mostro duas formas rápidas de chegar a um número defensável e já deixo uma faixa inicial para você aprovar."
                    else:
                        # Empty input
                        logger.warning(f"[LEGAL_NORMS] Empty input received, staying at stage")
                        ai_response = "Para prosseguir, preciso saber as normas aplicáveis. Diga 'não sei' se precisar de ajuda."
                        next_stage = current_stage
        
        elif current_stage == 'qty_value':
            # First, try to consume pending decision
            decision = try_consume_decision(session, user_message)
            if isinstance(decision, str):
                # Still asking for confirmation
                ai_response = decision
                next_stage = current_stage  # Stay in same stage
            elif isinstance(decision, dict):
                # Decision made
                if decision["action"] == "accept":
                    answers['qty_value'] = f'Estimativa aceita: {decision["text"]}'
                    session.set_answers(answers)
                    next_stage = 'installment'
                    ai_response = "Certo, adotei a estimativa inicial. Vamos falar de parcelamento…"
                elif decision["action"] == "pendente":
                    answers['qty_value'] = f'Pendente (com justificativa): {decision["text"]}'
                    session.set_answers(answers)
                    next_stage = 'installment'
                    ai_response = "Ok, registrei como Pendente (com justificativa). Agora, parcelamento…"
                elif decision["action"] == "debate":
                    next_stage = current_stage  # Stay in same stage
                    ai_response = "Vamos afinar a estimativa. Prefere estimar por **faixa de preço** por unidade/serviço ou por **benchmark** de contratos similares?"
            else:
                # No decision pending, process normally
                # Check for vague acknowledgment first
                if intents.is_vague_ack(user_message):
                    logger.info(f"[QTY_VALUE] Vague acknowledgment detected, staying at stage: {user_message[:30]}")
                    ai_response = "Sobre quantitativo e valor: dá para chutar uma ordem de grandeza? Se estiver nebuloso, eu te mostro duas formas rápidas de chegar a um número defensável e já deixo uma faixa inicial para você aprovar."
                    next_stage = current_stage  # Stay in same stage
                
                # Check for uncertainty using robust intent detection
                elif intents.is_uncertain_value(user_message):
                    logger.info(f"[QTY_VALUE] Uncertainty detected, offering help: {user_message[:30]}")
                    # Build proposal with estimation methods
                    proposal_text = """Método sugerido: benchmark no PNCP + 3 cotações preliminares.
Faixa inicial (exemplo): estimativa baseada em parâmetros conhecidos ou ordem de grandeza.
Justificativa: base prática para revisão posterior sem travar o ETP."""
                    
                    prompt_text = """Você sinalizou que não tem a estimativa agora. Aqui estão alguns métodos:

**Métodos de estimativa:**
1. **Benchmark de mercado** - Pesquisar contratos similares (Portal Nacional de Contratações Públicas)
2. **Histórico interno** - Consultar contratos anteriores do seu órgão
3. **Pesquisa de preços** - Coletar cotações de 3+ fornecedores
4. **Faixas estimadas** - Definir mínimo-máximo baseado em parâmetros conhecidos

Posso sugerir uma **faixa inicial** e metodologia de validação. Preferir **aceitar**, **pendente** ou **debater**?"""
                    
                    ai_response = ask_user_decision(session, prompt_text, proposal_text, 'qty_value')
                    next_stage = current_stage  # Don't advance yet
                
                # Check for explicit pending request
                elif intents.is_explicit_pending_request(user_message):
                    logger.info(f"[QTY_VALUE] Explicit pending request: {user_message[:30]}")
                    answers['qty_value'] = 'Pendente (a definir posteriormente)'
                    session.set_answers(answers)
                    next_stage = 'installment'
                    ai_response = "Ok, registrado como Pendente. Sobre parcelamento: você acha que faz sentido dividir em lotes ou fases? Se não tiver certeza, eu explico os prós e contras rapidamente e te ajudo a escolher o melhor para o seu caso."
                
                elif 'adotar estimativa' in user_message.lower() or 'adotar a estimativa' in user_message.lower():
                    # User wants to adopt the suggested estimate
                    logger.info(f"[QTY_VALUE] Adopting suggested estimate")
                    answers['qty_value'] = user_message + " (estimativa do sistema adotada)"
                    session.set_answers(answers)
                    next_stage = 'installment'
                    ai_response = "Estimativa adotada. Sobre parcelamento: você acha que faz sentido dividir em lotes ou fases? Se não tiver certeza, eu explico os prós e contras rapidamente e te ajudo a escolher o melhor para o seu caso."
                
                else:
                    # Valid substantive input
                    if user_message.strip():
                        logger.info(f"[QTY_VALUE] Storing substantive input: {user_message[:30]}")
                        answers['qty_value'] = user_message
                        session.set_answers(answers)
                        next_stage = 'installment'
                        ai_response = "Sobre parcelamento: você acha que faz sentido dividir em lotes ou fases? Se não tiver certeza, eu explico os prós e contras rapidamente e te ajudo a escolher o melhor para o seu caso."
                    else:
                        # Empty input
                        logger.warning(f"[QTY_VALUE] Empty input received, staying at stage")
                        ai_response = "Para prosseguir, preciso saber o quantitativo e valor. Diga 'não sei' se precisar de ajuda."
                        next_stage = current_stage
        
        elif current_stage == 'installment':
            # First, try to consume pending decision
            decision = try_consume_decision(session, user_message)
            if isinstance(decision, str):
                # Still asking for confirmation
                ai_response = decision
                next_stage = current_stage  # Stay in same stage
            elif isinstance(decision, dict):
                # Decision made
                if decision["action"] == "accept":
                    answers['installment'] = f'Diretriz aceita: {decision["text"]}'
                    session.set_answers(answers)
                    next_stage = 'summary'
                    # Generate compact summary
                    req_count = len(requirements)
                    selected_strategies = answers.get('selected_strategies', [])
                    strategy_titles = ', '.join([s.get('titulo', 'N/A') for s in selected_strategies]) if selected_strategies else 'não informado'
                    
                    ai_response = f"""Perfeito! Gerando resumo do ETP...

Resumo do ETP:

Necessidade: {necessity}

Requisitos: {req_count} itens

Estratégia(s) de Contratação: {strategy_titles}

PCA: {answers.get('pca', 'não informado')}

Normas: {answers.get('legal_norms', 'não informado')}

Quantitativo/Valor: {answers.get('qty_value', 'não informado')}

Parcelamento: {answers.get('installment', 'não informado')}

Confirme para gerar a prévia do documento."""
                elif decision["action"] == "pendente":
                    answers['installment'] = f'Pendente (com justificativa): {decision["text"]}'
                    session.set_answers(answers)
                    next_stage = 'summary'
                    # Generate compact summary
                    req_count = len(requirements)
                    selected_strategies = answers.get('selected_strategies', [])
                    strategy_titles = ', '.join([s.get('titulo', 'N/A') for s in selected_strategies]) if selected_strategies else 'não informado'
                    
                    ai_response = f"""Ok, registrado como Pendente. Gerando resumo do ETP...

Resumo do ETP:

Necessidade: {necessity}

Requisitos: {req_count} itens

Estratégia(s) de Contratação: {strategy_titles}

PCA: {answers.get('pca', 'não informado')}

Normas: {answers.get('legal_norms', 'não informado')}

Quantitativo/Valor: {answers.get('qty_value', 'não informado')}

Parcelamento: {answers.get('installment', 'não informado')}

Confirme para gerar a prévia do documento."""
                elif decision["action"] == "debate":
                    next_stage = current_stage  # Stay in same stage
                    ai_response = "Sem problemas. Quer discutir parcelamento por **lotes/itens**, por **especialidade técnica**, ou por **localidade geográfica**?"
            else:
                # No decision pending, process normally
                user_lower = user_message.lower()
                nao_sei_patterns = ['não sei', 'nao sei', 'n sei', 'ns', 'desconheço', 'incerto', 'em dúvida']
                
                if any(pattern in user_lower for pattern in nao_sei_patterns):
                    # Build strategy-based recommendation
                    selected_strategies = answers.get('selected_strategies', [])
                    strategy_title = selected_strategies[0].get('titulo', '') if selected_strategies else ''
                    
                    recommendation = "Avaliar com base na estratégia escolhida"
                    if 'arp' in strategy_title.lower() or 'registro de preços' in strategy_title.lower():
                        recommendation = "Para ARP, o parcelamento é comum para ampliar competitividade"
                    elif 'leasing' in strategy_title.lower() or 'locação' in strategy_title.lower():
                        recommendation = "Para locação/leasing, geralmente não há parcelamento (contrato único)"
                    elif 'outsourcing' in strategy_title.lower():
                        recommendation = "Para outsourcing, parcelamento pode ser feito por especialidade/localidade"
                    
                    proposal_text = f"""Diretriz sugerida: parcelar por etapas/lotação (implantação; operação/manutenção; relatórios), com salvaguardas de desempenho (SLA/KPIs) por parcela.
Justificativa: equilibra competitividade e controle de entrega. {recommendation}"""
                    
                    prompt_text = f"""Parcelamento é a divisão da contratação em lotes.

**Prós:**
- Amplia competitividade (permite ME/EPP participarem)
- Reduz risco de desabastecimento total
- Facilita gestão por áreas/especialidades

**Contras:**
- Pode aumentar custo total (perde economia de escala)
- Aumenta complexidade de gestão contratual
- Exige mais recursos para fiscalização

**Recomendação para seu caso:** {recommendation}

Posso propor uma diretriz de parcelamento e você decide: **aceitar**, **marcar pendente** ou **debater**?"""
                    
                    ai_response = ask_user_decision(session, prompt_text, proposal_text, 'installment')
                    next_stage = current_stage  # Don't advance yet
                elif 'inst_opt' in user_lower or any(opt in user_lower for opt in ['sem parcelamento', 'com parcelamento', 'híbrido']):
                    # User selected an option
                    if 'opt1' in user_lower or 'sem parcelamento' in user_lower:
                        answers['installment'] = 'Não haverá parcelamento. Justificativa: busca de economia de escala e uniformidade da solução.'
                    elif 'opt2' in user_lower or 'com parcelamento' in user_lower:
                        answers['installment'] = 'Haverá parcelamento em lotes. Justificativa: ampliar competitividade e permitir participação de ME/EPP.'
                    elif 'opt3' in user_lower or 'híbrido' in user_lower:
                        answers['installment'] = 'Haverá parcelamento híbrido. Justificativa: equilibrar economia de escala com competitividade.'
                    else:
                        answers['installment'] = user_message
                    
                    session.set_answers(answers)
                    next_stage = 'summary'
                    
                    # Generate compact summary
                    req_count = len(requirements)
                    selected_strategies = answers.get('selected_strategies', [])
                    strategy_titles = ', '.join([s.get('titulo', 'N/A') for s in selected_strategies]) if selected_strategies else 'não informado'
                    
                    summary = f"""Resumo do ETP:

Necessidade: {necessity}

Requisitos: {req_count} itens

Estratégia(s) de Contratação: {strategy_titles}

PCA: {answers.get('pca', 'não informado')}

Normas: {answers.get('legal_norms', 'não informado')}

Quantitativo/Valor: {answers.get('qty_value', 'não informado')}

Parcelamento: {answers.get('installment', 'não informado')}

Confirme para gerar a prévia do documento."""
                    
                    ai_response = summary
                else:
                    # Normal input
                    answers['installment'] = user_message
                    session.set_answers(answers)
                    next_stage = 'summary'
                    
                    # Generate compact summary
                    req_count = len(requirements)
                    selected_strategies = answers.get('selected_strategies', [])
                    strategy_titles = ', '.join([s.get('titulo', 'N/A') for s in selected_strategies]) if selected_strategies else 'não informado'
                    
                    summary = f"""Resumo do ETP:

Necessidade: {necessity}

Requisitos: {req_count} itens

Estratégia(s) de Contratação: {strategy_titles}

PCA: {answers.get('pca', 'não informado')}

Normas: {answers.get('legal_norms', 'não informado')}

Quantitativo/Valor: {answers.get('qty_value', 'não informado')}

Parcelamento: {answers.get('installment', 'não informado')}

Confirme para gerar a prévia do documento."""
                    
                    ai_response = summary
        
        elif current_stage == 'summary':
            if CONFIRM_RE.search(user_message):
                next_stage = 'preview'
                
                # Generate markdown preview for display in chat
                try:
                    # Prepare context for markdown builder
                    context_data = {
                        'necessity': necessity,
                        'requirements': requirements,
                        'answers': answers
                    }
                    
                    # Generate markdown document with multipass (returns etp-preview code block)
                    markdown_content = build_etp_markdown(context_data)
                    
                    # Also generate HTML/PDF files for download (optional backup)
                    preview_meta = build_preview(conversation_id, context_data)
                    
                    # Store in generator_output for response
                    generator_output = {
                        'preview_ready': True,
                        'html_path': preview_meta.get('html_path'),
                        'pdf_path': preview_meta.get('pdf_path'),
                        'file_path': preview_meta.get('html_path'),
                        'filename': preview_meta.get('filename'),
                        'markdown_content': markdown_content
                    }
                    
                    # Send etp-preview directly (already wrapped by build_etp_markdown)
                    # Frontend will detect etp-preview block and add download button
                    ai_response = f"""Perfeito! Aqui está a prévia completa do seu Estudo Técnico Preliminar:

{markdown_content}

O documento está pronto para revisão. Use o botão de download no canto direito do bloco acima para baixar o arquivo."""
                    
                    logger.info(f"[PREVIEW] Generated markdown preview for conversation {conversation_id}")
                    
                except Exception as e:
                    logger.error(f"[PREVIEW] Error generating preview: {e}")
                    logger.error(traceback.format_exc())
                    generator_output = {}
                    ai_response = f"Erro ao gerar prévia: {str(e)}"
            else:
                ai_response = "Confirme para gerar a prévia do documento."
        
        elif current_stage == 'preview':
            ai_response = "Documento finalizado. Você pode iniciar um novo documento."
        
        # Update session
        session.conversation_stage = next_stage
        session.updated_at = datetime.utcnow()
        db.session.commit()
        
        # Save AI response
        MessageRepo.add(
            conversation_id=conversation_id,
            role='assistant',
            content=ai_response,
            stage=next_stage
        )
        
        # Update conversation timestamp
        conv.updated_at = datetime.utcnow()
        db.session.commit()
        
        logger.info(f"[STAGE_CHAT] Transition: {current_stage} → {next_stage}")
        
        return jsonify({
            'success': True,
            'ai_response': ai_response,
            'stage': next_stage,
            'requirements': requirements if next_stage == 'suggest_requirements' else None,
            'preview_ready': generator_output.get('preview_ready', False),
            'file_path': generator_output.get('file_path'),
            'html_path': generator_output.get('html_path'),
            'pdf_path': generator_output.get('pdf_path'),
            'filename': generator_output.get('filename')
        }), 200
        
    except Exception as e:
        logger.error(f"[STAGE_CHAT] Error: {e}")
        logger.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'code': 'STAGE_CHAT_ERROR',
            'message': f'Erro ao processar mensagem: {str(e)}'
        }), 500

@etp_dynamic_bp.route('/open/<conversation_id>', methods=['GET'])
@cross_origin()
def open_conversation(conversation_id):
    """Open an existing conversation and return metadata with all messages"""
    try:
        # Get conversation
        conv = ConversationRepo.get(conversation_id)
        
        if not conv:
            return jsonify({
                'success': False,
                'error': 'Conversa não encontrada'
            }), 404
        
        # Get all messages for the conversation
        messages = MessageRepo.list_for_conversation(conversation_id)
        
        messages_data = []
        for msg in messages:
            messages_data.append({
                'id': msg.id,
                'role': msg.role,
                'content': msg.content,
                'stage': msg.stage,
                'payload': msg.payload,
                'created_at': msg.created_at.isoformat()
            })
        
        logger.info(f"[CONVERSATION] Opened conversation {conversation_id} with {len(messages_data)} messages")
        
        return jsonify({
            'success': True,
            'id': conv.id,
            'title': conv.title,
            'messages': messages_data,
            'created_at': conv.created_at.isoformat(),
            'updated_at': conv.updated_at.isoformat()
        }), 200
        
    except Exception as e:
        logger.error(f"Error opening conversation {conversation_id}: {e}")
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': f'Erro ao abrir conversa: {str(e)}'
        }), 500

@etp_dynamic_bp.route('/rename', methods=['PATCH'])
@cross_origin()
def rename_conversation():
    """Rename a conversation"""
    try:
        data = request.get_json(force=True)
        conversation_id = data.get('conversation_id')
        title = data.get('title')
        
        if not conversation_id or not title:
            return jsonify({
                'success': False,
                'error': 'conversation_id e title são obrigatórios'
            }), 400
        
        # Rename conversation
        conv = ConversationRepo.rename(conversation_id, title.strip())
        
        if not conv:
            return jsonify({
                'success': False,
                'error': 'Conversa não encontrada'
            }), 404
        
        db.session.commit()
        
        logger.info(f"[CONVERSATION] Renamed conversation {conversation_id} to '{title}'")
        
        return jsonify({
            'success': True,
            'id': conv.id,
            'title': conv.title,
            'updated_at': conv.updated_at.isoformat()
        }), 200
        
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error renaming conversation: {e}")
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': f'Erro ao renomear conversa: {str(e)}'
        }), 500

@etp_dynamic_bp.route('/<session_id>/title', methods=['PATCH'])
@cross_origin()
def update_conversation_title(session_id):
    """Update conversation title with updated_at timestamp"""
    try:
        data = request.get_json(force=True)
        title = data.get('title')
        
        if not title:
            return jsonify({
                'success': False,
                'error': 'title é obrigatório'
            }), 400
        
        # Get conversation
        conv = ConversationRepo.get(session_id)
        
        if not conv:
            return jsonify({
                'success': False,
                'error': 'Conversa não encontrada'
            }), 404
        
        # Update title and updated_at
        conv.title = title.strip()
        conv.updated_at = datetime.utcnow()
        db.session.commit()
        
        logger.info(f"[CONVERSATION] Updated title for {session_id} to '{title}'")
        
        return jsonify({
            'success': True,
            'id': conv.id,
            'title': conv.title,
            'updated_at': conv.updated_at.isoformat()
        }), 200
        
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error updating conversation title: {e}")
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': f'Erro ao atualizar título: {str(e)}'
        }), 500

@etp_dynamic_bp.route('/preview/<session_id>', methods=['GET'])
@cross_origin()
def preview_etp_document(session_id):
    """Generate preview of ETP document based on current session state"""
    try:
        # Initialize LLM client if needed
        global _openai_client, _simple_generator
        if _openai_client is None:
            api_key = os.getenv('OPENAI_API_KEY')
            if api_key:
                import openai
                _openai_client = openai.OpenAI(api_key=api_key)
                logger.info("[PREVIEW] OpenAI client initialized")
        
        if _simple_generator is None:
            _simple_generator = get_etp_generator(openai_client=_openai_client)
            logger.info("[PREVIEW] Generator initialized")
        
        # Get session from database
        session = db.session.query(EtpSession).filter_by(session_id=session_id).first()
        
        if not session:
            return jsonify({
                'success': False,
                'error': 'Sessão não encontrada'
            }), 404
        
        # Get answers
        answers = session.get_answers() or {}
        requirements = session.get_requirements() or []
        
        # Build document sections
        sections = {
            'titulo': f"Estudo Técnico Preliminar - {session.necessity or 'Contratação'}",
            'necessidade': session.necessity or "Não informada",
            'requisitos': requirements,
            'caminho_solucao': answers.get('solution_path', 'Não definido'),
            'pca': answers.get('pca', 'Não informado'),
            'normas_legais': answers.get('legal_norms', 'Lei 14.133/2021'),
            'quantidade_valor': answers.get('qty_value', 'A definir'),
            'parcelamento': answers.get('installment', 'Não definido'),
            'sumario': answers.get('summary', 'Aguardando finalização do fluxo')
        }
        
        logger.info(f"[PREVIEW] Generated preview for session {session_id}")
        
        return jsonify({
            'success': True,
            'session_id': session_id,
            'preview': sections,
            'stage': session.conversation_stage,
            'created_at': session.created_at.isoformat() if session.created_at else None,
            'updated_at': session.updated_at.isoformat() if session.updated_at else None
        }), 200
        
    except Exception as e:
        logger.error(f"[PREVIEW] Error generating preview for {session_id}: {e}")
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': f'Erro ao gerar prévia: {str(e)}'
        }), 500

@etp_dynamic_bp.route('/knowledge-base/info', methods=['GET'])
@cross_origin()
def get_knowledge_base_info():
    """Retorna informações sobre a base de conhecimento"""
    try:
        _ensure_initialized()
        if not etp_generator:
            return jsonify({'error': 'Gerador ETP não configurado'}), 500

        kb_info = etp_generator.get_knowledge_base_info()
        return jsonify({
            'success': True,
            'knowledge_base': kb_info,
            'timestamp': datetime.now().isoformat()
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e),
            'timestamp': datetime.now().isoformat()
        }), 500

@etp_dynamic_bp.route('/knowledge-base/refresh', methods=['POST'])
@cross_origin()
def refresh_knowledge_base():
    """Recarrega a base de conhecimento"""
    try:
        _ensure_initialized()
        if not etp_generator:
            return jsonify({'error': 'Gerador ETP não configurado'}), 500

        kb_info = etp_generator.refresh_knowledge_base()
        return jsonify({
            'success': True,
            'message': 'Base de conhecimento recarregada com sucesso',
            'knowledge_base': kb_info,
            'timestamp': datetime.now().isoformat()
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e),
            'timestamp': datetime.now().isoformat()
        }), 500

@etp_dynamic_bp.route('/questions', methods=['GET'])
@cross_origin()
def get_questions():
    """Retorna as perguntas do ETP"""
    _ensure_initialized()
    return jsonify({
        'success': True,
        'questions': ETP_QUESTIONS,
        'total': len(ETP_QUESTIONS)
    })

@etp_dynamic_bp.route('/session/start', methods=['POST'])
@cross_origin()
def start_session():
    """Inicia uma nova sessão de ETP"""
    try:
        _ensure_initialized()
        data = request.get_json()
        user_id = data.get('user_id', 1)  # Default user_id se não fornecido

        # Criar nova sessão
        session = EtpSession(
            user_id=user_id,
            session_id=str(uuid.uuid4()),
            status='active',
            answers=json.dumps({}),
            created_at=datetime.utcnow()
        )

        db.session.add(session)
        db.session.commit()

        return jsonify({
            'success': True,
            'session_id': session.session_id,
            'questions': ETP_QUESTIONS,
            'message': 'Sessão iniciada com sucesso'
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@etp_dynamic_bp.route('/session/<session_id>/answer', methods=['POST'])
@cross_origin()
def save_answer(session_id):
    """Salva resposta de uma pergunta"""
    try:
        _ensure_initialized()
        data = request.get_json()
        question_id = data.get('question_id')
        answer = data.get('answer')

        if not question_id or answer is None:
            return jsonify({'error': 'question_id e answer são obrigatórios'}), 400

        # Buscar sessão
        session = EtpSession.query.filter_by(session_id=session_id).first()
        if not session:
            return jsonify({'error': 'Sessão não encontrada'}), 404

        # Atualizar resposta
        answers = session.get_answers()
        answers[str(question_id)] = answer
        session.set_answers(answers)
        session.updated_at = datetime.utcnow()

        db.session.commit()

        return jsonify({
            'success': True,
            'message': 'Resposta salva com sucesso',
            'answers_count': len(answers)
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@etp_dynamic_bp.route('/session/<session_id>/generate', methods=['POST'])
@limiter.limit("10 per minute")
@cross_origin()
def generate_etp(session_id):
    """Gera ETP completo usando prompts dinâmicos"""
    try:
        _ensure_initialized()
        
        # GUARDRAIL: Validate generator exists before proceeding
        is_valid, error_msg = validate_generator_exists(etp_generator)
        if not is_valid:
            # Return controlled error message without calling generation
            error_response = handle_http_error(500, error_msg, 'unknown')
            return jsonify({
                'success': False,
                'error': error_msg,
                'message': error_response['ai_response']
            }), 500

        # Buscar sessão
        session = EtpSession.query.filter_by(session_id=session_id).first()
        if not session:
            return jsonify({'error': 'Sessão não encontrada'}), 404

        # GUARDRAIL: Check if ETP generation is allowed (state = confirm_requirements and user_confirmed = true)
        answers = session.get_answers() or {}
        user_confirmed = answers.get('user_confirmed_requirements', False)
        current_state = session.conversation_stage or 'collect_need'
        
        can_generate, error_msg = can_generate_etp(current_state, user_confirmed)
        
        if not can_generate:
            print(f"🔸 [STATE_MACHINE] Tentativa de gerar ETP bloqueada: {error_msg}")
            # Return controlled error without changing state
            error_response = handle_http_error(400, error_msg, current_state)
            return jsonify({
                'success': False,
                'error': error_msg,
                'message': error_response['ai_response'],
                'conversation_stage': current_state
            }), 400

        # Verificar se há respostas suficientes
        if not answers or len(answers) < 3:
            error_response = handle_http_error(400, 'Respostas insuficientes. Mínimo 3 respostas necessárias.', current_state)
            return jsonify({
                'success': False,
                'error': 'Respostas insuficientes. Mínimo 3 respostas necessárias.',
                'message': error_response['ai_response'],
                'conversation_stage': current_state
            }), 400

        # Preparar dados da sessão
        session_data = {
            'session_id': session_id,
            'answers': answers,
            'user_id': session.user_id
        }

        # Gerar ETP usando sistema dinâmico
        try:
            etp_content = etp_generator.generate_complete_etp(
                session_data=session_data,
                context_data=None,
                is_preview=False
            )
        except Exception as gen_error:
            # GUARDRAIL: Handle HTTP error (>=400) without changing state
            print(f"🔸 [GENERATOR ERROR] {gen_error}")
            error_response = handle_http_error(500, str(gen_error), current_state)
            return jsonify({
                'success': False,
                'error': str(gen_error),
                'message': error_response['ai_response'],
                'conversation_stage': current_state,
                'generation_method': 'dynamic_prompts'
            }), 500

        # GUARDRAIL: Transition to preview state after successful generation
        session.generated_etp = etp_content
        session.conversation_stage = 'preview'  # Transition to preview stage
        session.status = 'completed'
        session.updated_at = datetime.utcnow()

        db.session.commit()

        return jsonify({
            'success': True,
            'etp_content': etp_content,
            'message': 'ETP gerado com sucesso usando sistema dinâmico',
            'generation_method': 'dynamic_prompts',
            'knowledge_base_used': True,
            'conversation_stage': 'preview'
        })

    except Exception as e:
        # GUARDRAIL: Handle any unexpected error without changing state
        print(f"🔸 [UNEXPECTED ERROR] {e}")
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e),
            'generation_method': 'dynamic_prompts'
        }), 500

@etp_dynamic_bp.route('/session/<session_id>/preview', methods=['POST'])
@cross_origin()
def generate_preview(session_id):
    """Gera preview do ETP usando prompts dinâmicos"""
    try:
        _ensure_initialized()
        if not etp_generator:
            return jsonify({'error': 'Gerador ETP não configurado'}), 500

        # Buscar sessão
        session = EtpSession.query.filter_by(session_id=session_id).first()
        if not session:
            return jsonify({'error': 'Sessão não encontrada'}), 404

        # Preparar dados da sessão
        session_data = {
            'session_id': session_id,
            'answers': session.get_answers(),
            'user_id': session.user_id
        }

        # Gerar preview usando sistema dinâmico
        preview_content = etp_generator.generate_complete_etp(
            session_data=session_data,
            context_data=None,
            is_preview=True
        )

        return jsonify({
            'success': True,
            'preview_content': preview_content,
            'message': 'Preview gerado com sucesso usando sistema dinâmico',
            'generation_method': 'dynamic_prompts',
            'is_preview': True
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e),
            'generation_method': 'dynamic_prompts'
        }), 500

@etp_dynamic_bp.route('/session/<session_id>', methods=['GET'])
@cross_origin()
def get_session(session_id):
    """Retorna dados da sessão"""
    try:
        _ensure_initialized()
        session = EtpSession.query.filter_by(session_id=session_id).first()
        if not session:
            return jsonify({'error': 'Sessão não encontrada'}), 404

        return jsonify({
            'success': True,
            'session': session.to_dict()
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@etp_dynamic_bp.route('/consultative-options', methods=['POST'])
@cross_origin()
def generate_consultative_options():
    return jsonify({'error': 'Rota removida. Use POST /api/etp-dynamic/conversation'}), 404

@etp_dynamic_bp.route('/option-conversation', methods=['POST'])
@cross_origin()
def handle_option_conversation():
    """Gerencia conversa sobre as opções apresentadas"""
    try:
        _ensure_initialized()
        if not etp_generator:
            return jsonify({'error': 'Gerador ETP não configurado'}), 500

        # PASSO 10: Reutilizar session_id
        data = request.get_json(force=True)
        sid = (data.get("session_id") or "").strip() or None
        session = EtpSession.query.filter_by(session_id=sid).first() if sid else None
        
        # PASSO 10: Base de resposta com session_id
        resp_base = {"success": True, "session_id": sid} if sid else {"success": True}
        
        user_message = data.get('message', '').strip()
        options = data.get('options', [])
        conversation_history = data.get('conversation_history', [])

        if not user_message:
            return jsonify({'error': 'Mensagem do usuário é obrigatória'}), 400

        # Analisar se o usuário fez uma escolha final
        choice_analysis_prompt = f"""
        Analise se o usuário fez uma escolha definitiva entre as opções apresentadas.

        Opções disponíveis: {[opt['name'] for opt in options]}
        Mensagem do usuário: "{user_message}"

        Retorne JSON:
        {{
            "made_choice": true/false,
            "chosen_option": "nome da opção escolhida" ou null,
            "needs_clarification": true/false,
            "response_type": "choice|question|clarification"
        }}
        """

        model_name = get_model_name()
        choice_response = etp_generator.client.chat.completions.create(
            model=model_name,
            messages=[{"role": "user", "content": choice_analysis_prompt}],
            max_tokens=200,
            temperature=0.3
        )

        choice_result = json.loads(choice_response.choices[0].message.content.strip())

        # Gerar resposta contextual
        context_prompt = f"""
        Você é um consultor especialista em contratações públicas conversando sobre opções de atendimento.

        Opções apresentadas: {json.dumps(options, indent=2)}
        Mensagem do usuário: "{user_message}"
        Análise da escolha: {json.dumps(choice_result)}

        Responda de forma natural e consultiva, ajudando o usuário a:
        - Esclarecer dúvidas sobre as opções
        - Tomar uma decisão informada
        - Entender implicações de cada escolha

        Se o usuário fez uma escolha, confirme e oriente próximos passos.
        Se ainda está decidindo, ajude com mais informações.
        """

        ai_response = etp_generator.client.chat.completions.create(
            model=model_name,
            messages=[
                {"role": "system", "content": "Você é um consultor especialista em contratações públicas."},
                {"role": "user", "content": context_prompt}
            ],
            max_tokens=800,
            temperature=0.7
        )

        ai_response_text = ai_response.choices[0].message.content.strip()
        
        return jsonify({
            **resp_base,
            'kind': 'option_conversation',
            'ai_response': ai_response_text,
            'message': ai_response_text,
            'choice_analysis': choice_result,
            'timestamp': datetime.now().isoformat()
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e),
            'timestamp': datetime.now().isoformat()
        }), 500

@etp_dynamic_bp.route('/chat', methods=['POST'])
@cross_origin()
def chat_endpoint():
    """
    Unified chat endpoint that combines Analyzer (Prompt 1) + Dialogue (Prompt 2).
    Maintains session state without resets.
    
    Request JSON:
    {
        "session_id": "<uuid from client>",
        "message": "<user message>",
        "conversation_history": [{"role":"user|assistant","content":"..."}],
        "need": "<current necessity or empty>",
        "requirements": [{"id":"R1","text":"..."}, ...],
        "version": <int>
    }
    
    Response JSON:
    {
        "reply_markdown": "<natural response in Markdown>",
        "requirements": [{"id":"R1","text":"..."}, ...],
        "meta": {
            "need": "<current necessity>",
            "version": <int>
        }
    }
    """
    try:
        _ensure_initialized()
        if not etp_generator:
            return jsonify({'error': 'Gerador ETP não configurado'}), 500

        data = request.get_json()
        
        # Extract request data
        session_id = data.get('session_id', '').strip()
        user_message = data.get('message', '').strip()
        conversation_history = data.get('conversation_history', [])
        client_need = data.get('need', '').strip()
        client_requirements = data.get('requirements', [])
        client_version = data.get('version', 0)

        if not user_message:
            return jsonify({'error': 'Mensagem é obrigatória'}), 400

        # STEP 1: Load or create session (DO NOT generate new session_id)
        session = None
        if session_id:
            session = EtpSession.query.filter_by(session_id=session_id).first()
        
        if not session:
            # Only create new session if client didn't provide one
            if not session_id:
                session_id = str(uuid.uuid4())
            
            session = EtpSession(
                user_id=1,  # Default user
                session_id=session_id,
                status='active',
                answers=json.dumps({}),
                conversation_stage='collect_need',
                necessity=client_need if client_need else None,
                requirements_json=json.dumps(client_requirements, ensure_ascii=False) if client_requirements else json.dumps([]),
                requirements_version=client_version,
                created_at=datetime.utcnow()
            )
            db.session.add(session)
            db.session.commit()
        
        # STEP 2: Call Analyzer (Prompt 1) to check if message contains new necessity
        analyzer_result = call_analyzer_prompt(
            user_message=user_message,
            conversation_history=conversation_history,
            current_need=session.necessity or ""
        )
        
        contains_need = analyzer_result.get('contains_need', False)
        need_description = analyzer_result.get('need_description', '')
        
        # NOTE: Automatic reset removed - only restart_necessity intent should reset
        # If analyzer detects new necessity, it will be handled by dialogue model
        # without forcing a reset. User must explicitly use restart_necessity intent.
        
        # STEP 3: Call Dialogue (Prompt 2) with exact system prompt
        dialogue_input = {
            'need': session.necessity or "",
            'requirements': session.get_requirements(),
            'version': session.requirements_version,
            'history': conversation_history,
            'message': user_message
        }
        
        dialogue_result = call_dialogue_model(dialogue_input)
        
        # STEP 4: Update session state
        session.necessity = dialogue_result['meta']['need']
        session.set_requirements(dialogue_result['requirements'])
        session.requirements_version = dialogue_result['meta']['version']
        session.updated_at = datetime.utcnow()
        
        db.session.commit()
        
        return jsonify(dialogue_result)

    except Exception as e:
        print(f"🔸 Erro no chat endpoint: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'error': f'Erro no chat: {str(e)}'
        }), 500


def call_analyzer_prompt(user_message, conversation_history, current_need):
    """
    Analyzer (Prompt 1): Detects if user message contains a new necessity.
    Returns: {"contains_need": bool, "need_description": str}
    """
    try:
        _ensure_initialized()
        
        messages = [
            {
                "role": "system",
                "content": """Você é um ANALISADOR de intenção. Recebe a última mensagem do usuário + histórico.
Tarefa: dizer se a mensagem DEFINE uma nova necessidade (tema do ETP) ou não.

REGRAS:
- Se já existe uma necessidade definida no histórico e a nova mensagem só pede ajustes nos requisitos
  (ex.: "inclua outros", "troque o R4", "remova o item 2", "mantenha esses"), ENTÃO NÃO crie nova necessidade.
  Retorne contains_need=false e need_description="".
- Só retorne contains_need=true quando o usuário de fato apresentar um NOVO tema/escopo para a contratação
  (ou disser explicitamente que mudou o tema).
- Saída OBRIGATÓRIA em JSON estrito:
  {
    "contains_need": true|false,
    "need_description": "<texto ou vazio>"
  }
- Não escreva nada além do JSON."""
            }
        ]
        
        # Add conversation history
        for msg in conversation_history:
            messages.append({
                "role": msg.get("role", "user"),
                "content": msg.get("content", "")
            })
        
        # Add current need context if exists
        if current_need:
            messages.append({
                "role": "system",
                "content": f"Necessidade atual já definida: {current_need}"
            })
        
        # Add user message
        messages.append({
            "role": "user",
            "content": user_message
        })

        model_name = get_model_name()
        response = etp_generator.client.chat.completions.create(
            model=model_name,
            messages=messages,
            max_tokens=300,
            temperature=0.1,
            response_format={"type": "json_object"}
        )

        result = json.loads(response.choices[0].message.content)
        return {
            'contains_need': result.get('contains_need', False),
            'need_description': result.get('need_description', '')
        }
    except Exception as e:
        print(f"🔸 Erro no Analyzer: {e}")
        return {'contains_need': False, 'need_description': ''}


def call_dialogue_model(dialogue_input):
    """
    Dialogue (Prompt 2): ETP Consultant with RAG-first approach and strict JSON output.
    Returns: {"necessidade": str, "requisitos": [], "estado": {}}
    """
    try:
        _ensure_initialized()
        
        need = dialogue_input['need']
        requirements = dialogue_input['requirements']
        version = dialogue_input['version']
        history = dialogue_input['history']
        user_message = dialogue_input['message']
        
        # RAG-first: Search knowledge base for relevant requirements
        kb_context = ""
        if need:
            try:
                rag_results = search_requirements("generic", need, k=5)
                if rag_results:
                    kb_context = "\n\nConteúdo recuperado da base de conhecimento:\n"
                    for idx, result in enumerate(rag_results[:3], 1):
                        kb_context += f"{idx}. {result.get('content', '')}\n"
            except Exception as e:
                print(f"⚠️ Erro ao buscar RAG: {e}")
        
        # New system prompt from issue description
        system_prompt = """Você é a orquestradora conversacional do AutoDoc Licitação.
Objetivo: conduzir uma conversa natural para montar um ETP.

Regras de interação:

1. Sem botões. Nunca peça ou mencione botões. Espere respostas digitadas.

2. Sem "comando não reconhecido". Quando a intenção não estiver clara, reformule o que entendeu e faça uma pergunta objetiva.

3. Nunca avance de etapa sem confirmação explícita do usuário (ex.: "ok", "seguir", "pode manter", "aceito", "fechou").

4. Quando sugerir requisitos, não repita a lista em modo robótico. Faça um parágrafo curto explicando a lógica ("Por que esses requisitos?") e pergunte se quer manter, ajustar ou trocar algo — sempre por texto.

5. Se o usuário escrever ajustes em linguagem natural (ex.: "tira manutenção", "troca 3 por X", "inclui treinamento"), você deve:
   - extrair as mudanças,
   - refazer a lista,
   - confirmar com o usuário,
   - só então seguir.

6. Se houver erro interno de geração/preview, não avance nem peça desculpas genéricas e seguir adiante. Diga que houve um erro técnico e que você vai permanecer na etapa atual até o usuário confirmar nova tentativa.

7. Tom de voz: direto, humano e objetivo. Nada de "Perfeito! Coletei todas as informações necessárias" antes da confirmação do usuário.

8. Saídas sempre em texto corrido.

9. Palavras-chave de confirmação a reconhecer: "ok", "ok pode seguir", "segue", "prosseguir", "manter", "aceito", "concordo", "fechou", "pode gerar", "gostei dos requisitos".

Formato de saída JSON:

{
  "necessidade": "<texto curto e claro>",
  "requisitos": ["R1 ...", "R2 ...", "R3 ...", "R4 ...", "R5 ..."],
  "estado": {
    "etapa_atual": "<nome_da_etapa>",
    "proxima_etapa": "<nome_da_proxima_etapa>|null",
    "origem_requisitos": "base|gerado",
    "requisitos_confirmados": true|false
  }
}

RAG-first: sempre consulte a base de conhecimento antes de escrever qualquer coisa. Só gere conteúdo novo quando a base não cobrir.

Idioma: português do Brasil."""
        
        messages = [{"role": "system", "content": system_prompt}]
        
        # Add RAG context if available
        if kb_context:
            messages.append({"role": "system", "content": kb_context})
        
        # Convert old format requirements to new format (list of strings)
        requisitos_str = []
        if requirements:
            for req in requirements:
                req_id = req.get('id', '')
                req_text = req.get('text', '')
                requisitos_str.append(f"{req_id} — {req_text}")
        
        # Determine current stage
        etapa_atual = "coleta_necessidade"
        if need and not requirements:
            etapa_atual = "sugestao_requisitos"
        elif need and requirements:
            etapa_atual = "ajustes_requisitos"
        
        # Add context about current state
        context_msg = f"""Contexto atual:
- Necessidade: {need if need else "ainda não definida"}
- Requisitos atuais: {json.dumps(requisitos_str, ensure_ascii=False) if requisitos_str else "[]"}
- Etapa atual: {etapa_atual}
- Versão: {version}"""
        
        messages.append({"role": "system", "content": context_msg})
        
        # Add conversation history
        for msg in history[-10:]:  # Last 10 messages
            messages.append({
                "role": msg.get("role", "user"),
                "content": msg.get("content", "")
            })
        
        # Add current user message
        messages.append({"role": "user", "content": user_message})

        model_name = get_model_name()
        response = etp_generator.client.chat.completions.create(
            model=model_name,
            messages=messages,
            max_tokens=1500,
            temperature=0.1,  # Lower temperature for more consistent JSON output
            response_format={"type": "json_object"}
        )

        result = json.loads(response.choices[0].message.content)
        
        # Convert new format to old format for backward compatibility
        # New format: {"necessidade": str, "requisitos": ["R1 ...", "R2 ..."], "estado": {...}}
        # Old format: {"reply_markdown": str, "requirements": [{"id": "R1", "text": "..."}], "meta": {...}}
        
        if 'necessidade' in result and 'requisitos' in result and 'estado' in result:
            # New format detected - convert to old format
            converted_requirements = []
            requisitos = result.get('requisitos', [])
            
            for idx, req_str in enumerate(requisitos):
                # Parse "R1 — text" or "R1 - text" format
                if '—' in req_str:
                    parts = req_str.split('—', 1)
                    req_id = parts[0].strip()
                    req_text = parts[1].strip()
                elif ' - ' in req_str:
                    parts = req_str.split(' - ', 1)
                    req_id = parts[0].strip()
                    req_text = parts[1].strip()
                else:
                    # Fallback if format is different
                    req_id = f"R{idx+1}"
                    req_text = req_str.strip()
                
                converted_requirements.append({
                    "id": req_id,
                    "text": req_text
                })
            
            estado = result.get('estado', {})
            necessidade = result.get('necessidade', need or '')
            
            # Generate a simple reply message based on stage
            etapa_atual = estado.get('etapa_atual', 'sugestao_requisitos')
            reply = ""
            if etapa_atual == 'coleta_necessidade':
                reply = "Por favor, descreva a necessidade da contratação."
            elif etapa_atual == 'sugestao_requisitos':
                reply = "Aqui estão os requisitos sugeridos com base na necessidade identificada."
            elif etapa_atual == 'ajustes_requisitos':
                reply = "Requisitos atualizados. Você pode solicitar mais ajustes ou confirmar."
            elif etapa_atual == 'confirmacao_requisitos':
                if estado.get('requisitos_confirmados'):
                    reply = "Requisitos confirmados! Podemos prosseguir para a próxima etapa."
                else:
                    reply = "Por favor, revise os requisitos e confirme se está de acordo."
            else:
                reply = "Como posso ajudá-lo com o ETP?"
            
            # Increment version if requirements changed
            new_version = version
            if converted_requirements != requirements:
                new_version = version + 1
            
            result = {
                'reply_markdown': reply,
                'requirements': converted_requirements,
                'meta': {
                    'need': necessidade,
                    'version': new_version,
                    'estado': estado  # Keep original estado for debugging
                }
            }
        else:
            # Old format or fallback - ensure no justification fields
            if 'requirements' in result:
                for req in result['requirements']:
                    if 'justification' in req:
                        del req['justification']
                    if 'justificativa' in req:
                        del req['justificativa']
            
            # Validate response structure
            if 'reply_markdown' not in result:
                result['reply_markdown'] = "Olá! Como posso ajudar com seu ETP?"
            if 'requirements' not in result:
                result['requirements'] = requirements
            if 'meta' not in result:
                result['meta'] = {'need': need, 'version': version}
        
        return result
        
    except Exception as e:
        print(f"🔸 Erro no Dialogue: {e}")
        # Return fallback response
        return {
            'reply_markdown': f"Entendi sua mensagem. Como posso ajudar?",
            'requirements': dialogue_input['requirements'],
            'meta': {
                'need': dialogue_input['need'],
                'version': dialogue_input['version']
            }
        }


@etp_dynamic_bp.route('/chat-conversational', methods=['POST'])
@cross_origin()
def chat_conversational():
    """
    Complete conversational ETP flow without buttons (13-state machine).
    Implements the full specification from issue description.
    
    Request JSON:
    {
        "session_id": "<uuid from client>",
        "message": "<user message>"
    }
    
    Response JSON:
    {
        "success": true,
        "ai_response": "<natural response text>",
        "conversation_stage": "<current state>",
        "session_id": "<uuid>",
        "requirements": [{"id": "R1", "text": "..."}, ...],  # when applicable
        "state_changed": true|false
    }
    """
    try:
        _ensure_initialized()
        
        data = request.get_json()
        session_id = data.get('session_id', '').strip()
        user_message = data.get('message', '').strip()
        
        if not user_message:
            return jsonify({'success': False, 'error': 'Mensagem é obrigatória'}), 400
        
        # Load or create session
        session = ensure_session(session_id)
        current_state = session.conversation_stage or 'collect_need'
        
        # Get session data
        answers = session.get_answers() or {}
        requirements = session.get_requirements() or []
        
        session_data = {
            'necessity': session.necessity,
            'requirements': requirements,
            'answers': answers,
            'solution_path': answers.get('solution_path'),
            'pca': answers.get('pca'),
            'legal_norms': answers.get('legal_norms'),
            'quant_value': answers.get('quant_value'),
            'parcelamento': answers.get('parcelamento')
        }
        
        logger.info(f"[CONVERSATIONAL] State={current_state}, Message={user_message[:50]}")
        
        # Parse user intent
        intent = csm.parse_user_intent(user_message, current_state, session_data)
        logger.info(f"[CONVERSATIONAL] Intent={intent['intent']}, Confidence={intent['confidence']}")
        
        # Handle unclear intent
        if intent['intent'] == 'ask_clarification':
            question = intent['slots'].get('question', 'Não entendi. Pode reformular?')
            return jsonify(csm.handle_unclear_intent(current_state, question))
        
        # Process intent and update session data
        state_changed = False
        ai_response = ""
        
        # ========== STATE-SPECIFIC PROCESSING ==========
        
        if current_state == 'collect_need':
            if intent['intent'] == 'answer':
                # Store necessity
                session.necessity = user_message
                session.updated_at = datetime.utcnow()
                
                # Move to suggest_requirements
                next_state = 'suggest_requirements'
                session.conversation_stage = next_state
                state_changed = True
                
                # Generate requirements using RAG (dynamic count based on complexity)
                try:
                    rag_results = search_requirements("generic", user_message, k=12)
                    if rag_results and len(rag_results) > 0:
                        # Extract requirements from RAG (use up to 12 results)
                        suggested_reqs = []
                        for idx, result in enumerate(rag_results[:12], 1):
                            req_text = result.get('content', f'Requisito {idx}')
                            suggested_reqs.append({
                                'id': f'R{idx}',
                                'text': req_text
                            })
                        session.set_requirements(suggested_reqs)
                    else:
                        # Fallback using FallbackGenerator (10 requirements)
                        fallback_gen = FallbackGenerator()
                        fallback_reqs = fallback_gen.suggest_requirements(user_message)
                        suggested_reqs = []
                        for idx, req_text in enumerate(fallback_reqs, 1):
                            suggested_reqs.append({
                                'id': f'R{idx}',
                                'text': req_text
                            })
                        session.set_requirements(suggested_reqs)
                except Exception as e:
                    logger.error(f"Error in RAG search: {e}")
                    # Use fallback requirements from FallbackGenerator
                    fallback_gen = FallbackGenerator()
                    fallback_reqs = fallback_gen.suggest_requirements(user_message)
                    suggested_reqs = []
                    for idx, req_text in enumerate(fallback_reqs, 1):
                        suggested_reqs.append({
                            'id': f'R{idx}',
                            'text': req_text
                        })
                    session.set_requirements(suggested_reqs)
                
                db.session.commit()
                
                # Generate response
                ai_response = csm.generate_state_response(next_state, session_data, intent)
        
        elif current_state == 'suggest_requirements':
            # User can confirm or adjust requirements naturally
            if intent['intent'] == 'confirm':
                # Move to solution_strategies
                next_state = 'solution_strategies'
                session.conversation_stage = next_state
                state_changed = True
                db.session.commit()
                
                ai_response = csm.generate_state_response(next_state, session_data, intent)
            else:
                # Stay and let AI respond to user's feedback
                ai_response = csm.generate_state_response(current_state, session_data, intent)
        
        elif current_state == 'solution_strategies':
            if intent['intent'] == 'choose_strategy' or intent['intent'] == 'confirm':
                # Store the strategy choice (natural language accepted)
                strategy = intent['slots'].get('strategy', user_message)
                answers['solution_strategy'] = strategy
                session.set_answers(answers)
                
                # Move to pca
                next_state = 'pca'
                session.conversation_stage = next_state
                state_changed = True
                db.session.commit()
                
                session_data['solution_strategy'] = strategy
                ai_response = csm.generate_state_response(next_state, session_data, intent)
            else:
                # Stay and let AI respond
                ai_response = csm.generate_state_response(current_state, session_data, intent)
        
        elif current_state == 'pca':
            if intent['intent'] == 'answer' or intent['intent'] == 'skip' or intent['intent'] == 'confirm':
                # Store PCA answer (or handle "não sei")
                pca_value = intent['slots'].get('value', user_message)
                pca_desc = intent['slots'].get('description', '')
                pca_text = f"{pca_value}" + (f" - {pca_desc}" if pca_desc else "")
                answers['pca'] = pca_text
                session.set_answers(answers)
                
                # Move to legal_norms
                next_state = 'legal_norms'
                session.conversation_stage = next_state
                state_changed = True
                db.session.commit()
                
                session_data['pca'] = pca_text
                ai_response = csm.generate_state_response(next_state, session_data, intent)
            else:
                # Handle "não sei" - AI should suggest options
                ai_response = csm.generate_state_response(current_state, session_data, intent)
        
        elif current_state == 'legal_norms':
            if intent['intent'] == 'answer' or intent['intent'] == 'skip' or intent['intent'] == 'confirm':
                # Store legal norms (or handle "não sei")
                legal_norms = intent['slots'].get('text', user_message)
                answers['legal_norms'] = legal_norms
                session.set_answers(answers)
                
                # Move to qty_value
                next_state = 'qty_value'
                session.conversation_stage = next_state
                state_changed = True
                db.session.commit()
                
                session_data['legal_norms'] = legal_norms
                ai_response = csm.generate_state_response(next_state, session_data, intent)
            else:
                # Handle "não sei" - AI should suggest norms
                ai_response = csm.generate_state_response(current_state, session_data, intent)
        
        elif current_state == 'qty_value':
            if intent['intent'] == 'answer' or intent['intent'] == 'skip' or intent['intent'] == 'confirm':
                # Store quantitative/value info (or handle "não sei")
                slots = intent['slots']
                if 'text' in slots:
                    answers['quant_value'] = slots['text']
                else:
                    # Use user message directly for natural input
                    answers['quant_value'] = user_message
                
                session.set_answers(answers)
                
                # Move to installment
                next_state = 'installment'
                session.conversation_stage = next_state
                state_changed = True
                db.session.commit()
                
                session_data['quant_value'] = answers['quant_value']
                ai_response = csm.generate_state_response(next_state, session_data, intent)
            else:
                # Handle "não sei" - AI should suggest value ranges
                ai_response = csm.generate_state_response(current_state, session_data, intent)
        
        elif current_state == 'installment':
            if intent['intent'] == 'answer' or intent['intent'] == 'skip' or intent['intent'] == 'confirm':
                # Store installment info (or handle "não sei")
                answers['installment'] = user_message
                session.set_answers(answers)
                
                # Move to summary
                next_state = 'summary'
                session.conversation_stage = next_state
                state_changed = True
                db.session.commit()
                
                # Update session_data for summary
                session_data['answers'] = answers
                session_data['installment'] = answers['installment']
                ai_response = csm.generate_state_response(next_state, session_data, intent)
            else:
                # Handle "não sei" - AI should explain pros/cons
                ai_response = csm.generate_state_response(current_state, session_data, intent)
        
        elif current_state == 'summary':
            if intent['intent'] == 'confirm' or intent['intent'] == 'confirm_generate':
                # Move to preview
                next_state = 'preview'
                session.conversation_stage = next_state
                state_changed = True
                
                # Generate ETP preview
                try:
                    ai_response = "Gerando a prévia do ETP com as informações coletadas..."
                    
                    # Build ETP data structure
                    etp_data = {
                        'necessidade': session.necessity,
                        'requisitos': session.get_requirements(),
                        'estrategia_solucao': answers.get('solution_strategy', 'não informado'),
                        'pca': answers.get('pca', 'não informado'),
                        'normas_legais': answers.get('legal_norms', 'não informado'),
                        'quantitativo_valor': answers.get('quant_value', 'não informado'),
                        'parcelamento': answers.get('installment', 'não informado')
                    }
                    
                    # Store in session for later use
                    session.generated_etp = json.dumps(etp_data, ensure_ascii=False)
                    
                    # Move to preview
                    next_state = 'preview'
                    session.conversation_stage = next_state
                    db.session.commit()
                    
                    ai_response += "\n\n" + csm.generate_state_response(next_state, session_data, intent)
                    
                except Exception as e:
                    logger.error(f"Error generating ETP: {e}")
                    # Stay in generate_etp state
                    return jsonify(csm.handle_service_error(current_state, str(e)))
            
            elif intent['intent'] == 'request_adjustment':
                # Stay in confirm_summary for adjustment
                ai_response = "Qual informação você gostaria de ajustar? Me diga e eu atualizo."
        
        elif current_state == 'preview':
            # Preview is the final state - user can download or start new session
            session.status = 'completed'
            db.session.commit()
            ai_response = csm.generate_state_response(current_state, session_data, intent)
        
        # If no response was generated, use state response
        if not ai_response:
            ai_response = csm.generate_state_response(session.conversation_stage, session_data, intent)
        
        # Build response
        response = {
            'success': True,
            'ai_response': ai_response,
            'conversation_stage': session.conversation_stage,
            'session_id': session.session_id,
            'state_changed': state_changed
        }
        
        # Include requirements if relevant
        if session.conversation_stage in ['suggest_requirements', 'solution_strategies']:
            response['requirements'] = session.get_requirements()
        
        return jsonify(response)
    
    except Exception as e:
        logger.error(f"Error in chat_conversational: {e}")
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': f'Erro no chat conversacional: {str(e)}'
        }), 500


@etp_dynamic_bp.route('/chat-persist', methods=['POST'])
@cross_origin()
def chat_persist():
    """Send a message in a conversation with full database persistence"""
    try:
        data = request.get_json(force=True)
        conversation_id = data.get('conversation_id')
        message = data.get('message', '').strip()
        
        if not conversation_id or not message:
            return jsonify({
                'success': False,
                'error': 'conversation_id e message são obrigatórios'
            }), 400
        
        # Verify conversation exists
        conv = ConversationRepo.get(conversation_id)
        if not conv:
            return jsonify({
                'success': False,
                'error': 'Conversa não encontrada'
            }), 404
        
        # Save user message
        user_msg = MessageRepo.add(
            conversation_id=conversation_id,
            role='user',
            content=message
        )
        
        # Get conversation history for context
        all_messages = MessageRepo.list_for_conversation(conversation_id)
        history = [
            {"role": msg.role, "content": msg.content}
            for msg in all_messages[:-1]  # Exclude the just-added user message
        ]
        
        # Generate AI response using OpenAI
        try:
            client = get_llm_client()
            model = get_model_name()
            
            # Build messages for LLM
            system_prompt = (
                "Você é um assistente especializado em ajudar a criar Estudos Técnicos Preliminares (ETP). "
                "Faça perguntas claras e objetivas para coletar informações sobre a necessidade de contratação."
            )
            
            llm_messages = [{"role": "system", "content": system_prompt}]
            llm_messages.extend(history)
            llm_messages.append({"role": "user", "content": message})
            
            response = client.chat.completions.create(
                model=model,
                messages=llm_messages,
                temperature=0.7,
                max_tokens=1000
            )
            
            ai_response = response.choices[0].message.content
            
        except Exception as e:
            logger.error(f"Error calling LLM: {e}")
            ai_response = "Desculpe, ocorreu um erro ao processar sua mensagem. Por favor, tente novamente."
        
        # Save assistant message
        assistant_msg = MessageRepo.add(
            conversation_id=conversation_id,
            role='assistant',
            content=ai_response
        )
        
        db.session.commit()
        
        logger.info(f"[CONVERSATION] Message exchange in {conversation_id}")
        
        return jsonify({
            'success': True,
            'user_message': {
                'id': user_msg.id,
                'content': user_msg.content,
                'created_at': user_msg.created_at.isoformat()
            },
            'assistant_message': {
                'id': assistant_msg.id,
                'content': assistant_msg.content,
                'created_at': assistant_msg.created_at.isoformat()
            }
        }), 200
        
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error in chat-persist: {e}")
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': f'Erro ao processar mensagem: {str(e)}'
        }), 500

@etp_dynamic_bp.route('/conversation', methods=['POST'])
@cross_origin()
def etp_conversation():
    """Conduz conversa natural usando o modelo fine-tuned para coleta de informações do ETP"""
    try:
        _ensure_initialized()
        # NOTE: We no longer return 500 if etp_generator is missing
        # The simple generator with fallback will handle requirements generation

        data = request.get_json()
        print("🔹 Recebi do frontend:", data)

        # Reset suave de conversa
        if data.get('reset') is True:
            sid = (data.get('session_id') or '').strip() or None
            session = ensure_session(sid)
            # limpa a sessão do ETP, preservando o session_id
            session.necessity = None
            session.set_requirements([])
            session.set_answers({})
            session.conversation_stage = 'collect_need'
            session.updated_at = datetime.utcnow()
            db.session.commit()
            return jsonify({
                'success': True,
                'message': 'Vamos começar do zero. Qual é a necessidade da contratação?',
                'conversation_stage': session.conversation_stage,
                'session_id': session.session_id
            })

        # Aceitar ambos os nomes para compatibilidade com o front atual
        user_message = (data.get('user_message') or data.get('message') or '').strip()
        session_id = data.get('session_id')
        conversation_history = data.get('conversation_history', [])
        answered_questions = data.get('answered_questions', [])

        if not user_message:
            return jsonify({'success': False, 'error': 'Mensagem é obrigatória'}), 400

        # REQUIREMENT: session_id is mandatory
        request_sid = (session_id or "").strip() or None
        if not request_sid:
            return jsonify({
                'success': False, 
                'error': 'session_id é obrigatório. Use POST /api/etp-dynamic/new para criar uma nova sessão.'
            }), 400

        # Get or create session with the provided session_id
        session = ensure_session(request_sid)
        
        # REQUIREMENT 1: Ensure session_id and client are never None
        # If session_id is None, create new and reinitialize
        if not session.session_id:
            session.session_id = str(uuid.uuid4())
            session.conversation_stage = 'collect_need'
            session.updated_at = datetime.utcnow()
            db.session.commit()
            logger.info(f"[SESSION] Created new session_id: {session.session_id}")
        
        # Ensure OpenAI client is initialized for etp_generator
        global _openai_client, _simple_generator
        try:
            if _openai_client is None:
                api_key = os.getenv('OPENAI_API_KEY')
                if api_key:
                    import openai
                    _openai_client = openai.OpenAI(api_key=api_key)
                    logger.info("[CLIENT] OpenAI client initialized")
            
            if _simple_generator is None:
                _simple_generator = get_etp_generator(openai_client=_openai_client)
                logger.info("[GENERATOR] ETP generator initialized")
        except Exception as e:
            logger.error(f"[CLIENT] Failed to initialize client: {e}")
            # Don't fail - use fallback generator
            if _simple_generator is None:
                _simple_generator = FallbackGenerator()

        # Base de resposta com session_id
        resp_base = {"success": True, "session_id": session.session_id}
        
        print(f"🔹 [ANTES] request_sid={request_sid} | session_id={session.session_id} | stage={session.conversation_stage} | necessity={bool(session.necessity)}")
        print(f"🔹 [INPUT] Mensagem: '{user_message[:50]}{'...' if len(user_message) > 50 else '}'}")

        # PASSO 3: Interpretador de comandos ANTES do LLM
        # GUARDRAIL: Handle confirm_requirements stage
        if session.conversation_stage == 'confirm_requirements':
            # Check if user confirmed
            user_confirmed = is_user_confirmed(user_message)
            
            if user_confirmed:
                # ORDEM CORRETA: CAMINHO → PCA → NORMAS → QTD/VALOR → PARCELAMENTO → RESUMO → GERAR
                # Gerar transição em prosa natural
                next_question = _generate_prose_transition('confirm_to_solution_path', {'necessity': session.necessity})
                if not next_question:
                    next_question = "Certo. Vamos identificar a melhor forma de atender essa necessidade."
                
                session.conversation_stage = 'solution_path'
                session.updated_at = datetime.utcnow()
                db.session.commit()
                
                return jsonify({
                    'success': True,
                    'kind': 'text',
                    'ai_response': next_question,
                    'message': next_question,
                    'conversation_stage': session.conversation_stage,
                    'session_id': session.session_id
                })
            else:
                # User did not confirm, return to refine_requirements for more adjustments
                session.conversation_stage = 'refine_requirements'
                session.updated_at = datetime.utcnow()
                db.session.commit()
                
                return jsonify({
                    **resp_base,
                    'kind': 'text',
                    'ai_response': 'Sem confirmação. Posso ajustar mais algo? (diga o que mudar ou "pode seguir" quando estiver pronto)',
                    'message': 'Sem confirmação. Posso ajustar mais algo? (diga o que mudar ou "pode seguir" quando estiver pronto)',
                    'conversation_stage': session.conversation_stage
                })
        
        # GUARDRAIL: Handle suggest_requirements and refine_requirements stages with state machine
        if session.conversation_stage in ['suggest_requirements', 'refine_requirements', 'review_requirements']:
            from domain.usecase.etp.requirements_interpreter import parse_update_command, aplicar_comando
            
            current_requirements = session.get_requirements()
            command_result = parse_update_command(user_message, current_requirements)
            
            print(f"🔹 Comando parseado: {command_result}")
            
            # Check if user confirmed using state machine validation
            user_confirmed = is_user_confirmed(user_message)
            print(f"🔹 [STATE_MACHINE] user_confirmed={user_confirmed}")
            
            # Ramo explícito de confirmação - GUARDRAIL: only transition with user_confirmed = true
            if command_result and command_result.get('intent') == 'confirm':
                if not user_confirmed:
                    # Should not happen as 'confirm' intent means confirmation detected, but be safe
                    return jsonify({
                        **resp_base,
                        'kind': 'text',
                        'ai_response': 'Para confirmar os requisitos, use palavras como: "ok", "confirmo", "pode seguir", "aceito", "concordo".',
                        'message': 'Para confirmar os requisitos, use palavras como: "ok", "confirmo", "pode seguir", "aceito", "concordo".',
                        'conversation_stage': session.conversation_stage
                    })
                
                # GUARDRAIL: Validate state transition to confirm_requirements
                next_state = 'confirm_requirements'
                is_valid, error_msg = validate_state_transition(session.conversation_stage, next_state, user_confirmed)
                
                if not is_valid:
                    print(f"🔸 [STATE_MACHINE] Transição inválida: {error_msg}")
                    return jsonify({
                        **resp_base,
                        'kind': 'text',
                        'ai_response': f"Não é possível confirmar agora: {error_msg}",
                        'message': f"Não é possível confirmar agora: {error_msg}",
                        'conversation_stage': session.conversation_stage
                    })
                
                # User confirmed - now transition to solution_path (not confirm_requirements)
                # The issue specifies: do not ask "Confirmo gerar o ETP" after requirements
                # Instead, move directly to solution_path
                session.conversation_stage = 'solution_path'
                session.updated_at = datetime.utcnow()
                db.session.commit()
                
                # Gerar transição em prosa natural
                next_question = _generate_prose_transition('confirm_to_solution_path', {'necessity': session.necessity})
                if not next_question:
                    next_question = "Certo. Vamos identificar a melhor forma de atender essa necessidade."
                
                return jsonify({
                    'success': True,
                    'kind': 'text',
                    'ai_response': next_question,
                    'message': next_question,
                    'conversation_stage': session.conversation_stage,
                    'session_id': session.session_id
                })
            
            if command_result['intent'] == 'ask':
                # User asked a question - provide helpful context about requirements
                return jsonify({
                    **resp_base,
                    'kind': 'text',
                    'ai_response': 'Estamos revisando os requisitos. Se algo não faz sentido ou falta algo importante, me avise que ajusto. Quando estiver adequado, confirme para seguirmos.',
                    'message': 'Estamos revisando os requisitos. Se algo não faz sentido ou falta algo importante, me avise que ajusto. Quando estiver adequado, confirme para seguirmos.',
                    'conversation_stage': session.conversation_stage
                })
            
            if command_result['intent'] == 'restart_necessity':
                # Reset session to collect new necessity
                session.necessity = None
                session.conversation_stage = 'collect_need'
                session.set_requirements([])
                session.updated_at = datetime.utcnow()
                db.session.commit()
                
                return jsonify({
                    **resp_base,
                    'kind': 'text',
                    'ai_response': 'Entendido. Vamos recomeçar. Qual é a nova necessidade da contratação?',
                    'message': 'Entendido. Vamos recomeçar. Qual é a nova necessidade da contratação?',
                    'conversation_stage': 'collect_need'
                })
            
            elif command_result['intent'] in ['remove', 'edit', 'add', 'keep_only', 'reorder']:
                # PASSO 3: Apply the command to requirements with stable renumbering
                aplicar_comando(command_result, session, session.necessity)
                
                # Handle combo commands (e.g., remove + add)
                if command_result.get('_combo_add'):
                    for add_item in command_result['_combo_add']:
                        add_cmd = {
                            'intent': 'add',
                            'items': [add_item.get('text', '')],
                            'message': f'Adicionando: {add_item.get("text", "")}'
                        }
                        aplicar_comando(add_cmd, session, session.necessity)
                
                # GUARDRAIL: Stay in refine_requirements stage (loop allowed)
                session.conversation_stage = 'refine_requirements'
                session.updated_at = datetime.utcnow()
                db.session.commit()
                
                # PASSO 6: Return with unified contract - natural messages per action type
                updated_requirements = session.get_requirements()
                intent = command_result['intent']
                if intent == 'remove':
                    msg = "Removi. Veja como ficou a lista atualizada e me diga se ajusto mais algo ou se posso seguir."
                elif intent == 'edit':
                    msg = "Perfeito. Atualizei e já deixei a lista abaixo com a mudança aplicada. Sigo ajustando?"
                elif intent == 'add':
                    msg = "Incluído. A lista abaixo já reflete o novo item. Ajusto mais algo?"
                else:
                    msg = command_result.get('message') or 'Atualização aplicada. Veja a lista atualizada abaixo.'
                    msg += " Ajusto mais algo ou posso seguir?"
                
                return jsonify({
                    **resp_base,
                    'kind': 'requirements_update',
                    'necessity': session.necessity,
                    'requirements': updated_requirements,
                    'message': msg,
                    'ai_response': msg,
                    'conversation_stage': session.conversation_stage
                })
            
            # Usuário pediu "adicionar mais" mas sem conteúdo → ofereça candidatos
            if re.search(r'\b(adicione|adicionar|incluir|mais requisito|mais requisitos)\b', user_message, flags=re.I) and _is_uncertain(user_message):
                candidates = _suggest_requirements(session.necessity)
                pool = [c for c in candidates if c not in [r.get('text', '') for r in current_requirements]][:6]
                if pool:
                    lines = [f"{i+1}) {it}" for i,it in enumerate(pool)]
                    msg = "Sugestões para incluir. Responda com os **números** (ex.: \"1 e 3\"):\n" + "\n".join(lines)
                    ans = session.get_answers() or {}
                    ans['req_candidates'] = pool
                    session.set_answers(ans)
                    db.session.commit()
                    return jsonify({**resp_base,'message': msg,'requirements': current_requirements,'conversation_stage': 'refine_requirements'})
            
            # Seleção numérica ("inclui 1 e 3")
            ans = session.get_answers() or {}
            pool = ans.get('req_candidates') or []
            picks = _pick_numbers(user_message, len(pool)) if pool else []
            if picks:
                for p in picks:
                    new_id = f'R{len(current_requirements) + 1}'
                    current_requirements.append({'id': new_id, 'text': pool[p-1]})
                ans['req_candidates'] = []
                session.set_answers(ans)
                session.set_requirements(current_requirements)
                db.session.commit()
                return jsonify({**resp_base,'message': "Incluí os itens selecionados. Posso seguir?",'requirements': current_requirements,'conversation_stage': 'refine_requirements'})
            
            # GUARDRAIL: If intent is 'other' or 'unclear', use handle_other_intent
            if command_result['intent'] in ['other', 'unclear']:
                clarification = handle_other_intent()
                return jsonify({
                    **resp_base,
                    'kind': 'text',
                    'ai_response': clarification['ai_response'],
                    'message': clarification['message'],
                    'conversation_stage': session.conversation_stage
                })
            
            # Fallback for any other case
            return jsonify({
                **resp_base,
                'kind': 'text',
                'ai_response': 'Entendi parcialmente. Quer remover, trocar ou incluir algo?',
                'message': 'Entendi parcialmente. Quer remover, trocar ou incluir algo?',
                'conversation_stage': session.conversation_stage
            })

        # ORDEM das ETAPAS (sem auto-avanço):
        # solution_path -> pca -> legal_norms -> qty_value -> installment -> summary -> gerar
        if session.conversation_stage == 'solution_path':
            ans = session.get_answers() or {}
            text = user_message.strip()
            
            # Check if user is asking for help
            if any(k in user_message.lower() for k in ['ajuda', 'me ajude', 'me ajuda', 'como', 'não sei', 'qual escolher']):
                help_msg = _consult_path_explainer(session.necessity)[0]
                return jsonify({**resp_base, 'message': help_msg, 'conversation_stage': 'solution_path'})
            
            # Check if already have a tentative choice waiting for approval
            if ans.get('solution_path_tentative'):
                # User is responding to confirmation request
                if _is_confirm(user_message) or _yes(user_message):
                    # Approved - advance to pca
                    chosen = ans['solution_path_tentative']
                    ans['solution_path'] = chosen
                    ans['solution_justification'] = f"Opção escolhida: {chosen}. Justificativa baseada na análise e aderência à necessidade."
                    session.set_answers(ans)
                    session.conversation_stage = 'pca'
                    session.updated_at = datetime.utcnow()
                    db.session.commit()
                    
                    pca_msg = _generate_prose_transition('solution_path_to_pca', {'chosen': chosen})
                    if not pca_msg:
                        pca_msg = f"Certo, vamos com {chosen}. Sobre o planejamento anual: já existe previsão para isso no PCA?"
                    
                    return jsonify({**resp_base,'message': pca_msg, 'conversation_stage': 'pca'})
                else:
                    # User wants to reconsider - clear tentative and show options again
                    ans.pop('solution_path_tentative', None)
                    session.set_answers(ans)
                    db.session.commit()
                    msg_ctx, _ = _consult_path_explainer(session.necessity)
                    return jsonify({**resp_base,'message': msg_ctx,'conversation_stage': 'solution_path'})
            
            # No tentative choice yet - parse user's selection
            msg_ctx, options = _consult_path_explainer(session.necessity)
            choice_kw = _any_in(user_message, options)
            picks = _pick_numbers(user_message, len(options))
            
            if _is_uncertain(user_message) or (not choice_kw and not picks):
                # User didn't pick anything - show options
                return jsonify({**resp_base,'message': msg_ctx,'conversation_stage': 'solution_path'})
            
            # User picked something - save tentatively and ask for confirmation
            chosen = choice_kw or options[picks[0]-1]
            
            # Guardar TODAS as alternativas (para o estudo final)
            analysis_lines = [l for l in msg_ctx.splitlines() if re.match(r'^\d+\)', l)]
            parsed = []
            for line in analysis_lines:
                try:
                    _, rest = line.split(')', 1)
                    name, desc = rest.strip().split('—', 1)
                    parsed.append({'opcao': name.strip(), 'analise': desc.strip()})
                except Exception:
                    pass
            ans['solution_options'] = parsed
            ans['solution_path_tentative'] = chosen
            session.set_answers(ans)
            db.session.commit()
            
            # Ask for confirmation
            confirm_msg = f"Entendi que você escolheu: {chosen}. Posso seguir com essa opção para o PCA?"
            return jsonify({**resp_base,'message': confirm_msg,'conversation_stage': 'solution_path'})

        if session.conversation_stage == 'pca':
            ans = session.get_answers() or {}
            
            # ENHANCED: Check if user is requesting PCA construction
            user_lower = user_message.lower()
            pca_build_keywords = ['monta um pca', 'faça o pca', 'criar pca', 'gerar pca', 'construir pca', 'monte um pca']
            is_pca_request = any(kw in user_lower for kw in pca_build_keywords)
            
            # Check for "what is PCA" questions
            if any(phrase in user_lower for phrase in ['o que é pca', 'o que e pca', 'como faz pca', 'como fazer pca']):
                explanation = """O PCA (Plano de Contratações Anual) é o documento que registra todas as contratações previstas do órgão para o ano. Ele deve conter:
- Objetivos e metas institucionais
- Funções de despesa
- Vínculo com PPA/LDO
- Justificativa das contratações planejadas

Se não tiver um PCA pronto, posso montar um rascunho básico para você. Quer que eu faça isso?"""
                return jsonify({**resp_base, 'message': explanation, 'conversation_stage': 'pca'})
            
            # If user doesn't have PCA and requests to build one
            if is_pca_request or (any(kw in user_lower for kw in ['não tenho', 'nao tenho', 'não possuo']) and 'pca' in user_lower):
                # Generate PCA draft
                necessity = session.necessity or "contratação descrita"
                solution_path = ans.get('solution_path', 'solução escolhida')
                
                pca_draft = f"""Rascunho de PCA para {necessity}:

**Objetivos e Metas:**
- Atender à necessidade de {necessity}
- Garantir qualidade e conformidade na prestação
- Assegurar continuidade operacional

**Função de Despesa:** 
- Classificar conforme natureza (investimento/custeio)

**Vínculo PPA/LDO:**
- Alinhar com programa/ação do Plano Plurianual
- Confirmar previsão orçamentária na LDO

**Justificativa:**
- Contratação necessária para {solution_path}
- Requisitos técnicos validados
- Estimativa de recursos a definir

Posso usar este rascunho de PCA para seguirmos?"""
                
                ans['pca_draft_offered'] = pca_draft
                session.set_answers(ans)
                db.session.commit()
                
                return jsonify({**resp_base, 'message': pca_draft, 'conversation_stage': 'pca'})
            
            # If we offered a draft and user is confirming
            if ans.get('pca_draft_offered') and (_is_confirm(user_message) or _yes(user_message)):
                pca_value = ans['pca_draft_offered']
                ans['pca'] = pca_value
                ans.pop('pca_draft_offered', None)
                session.set_answers(ans)
                session.conversation_stage = 'legal_norms'
                session.updated_at = datetime.utcnow()
                db.session.commit()
                
                legal_msg = _generate_prose_transition('pca_to_legal_norms', {'pca_response': 'PCA estruturado'})
                if not legal_msg:
                    legal_msg = "Ótimo! Sobre as normas legais: costumamos usar a Lei 14.133/21 como base, junto com o regulamento local aplicável. Se houver algo específico do setor ou proteção de dados envolvida, podemos incluir. O que faz sentido para você?"
                
                return jsonify({**resp_base, 'message': legal_msg, 'conversation_stage': 'legal_norms'})
            
            # Detectar se usuário precisa de ajuda ou está incerto
            if _is_uncertain(user_message) or any(k in user_lower for k in ['ajuda', 'como', 'não sei']):
                # Usuário não tem informação ou precisa de orientação
                help_msg = "O PCA (Plano de Contratações Anual) lista as compras previstas do órgão. Se ainda não constar, é bom registrar essa demanda junto ao setor de planejamento antes de seguir. Ou se preferir, posso montar um rascunho básico de PCA para você. O que prefere?"
                return jsonify({**resp_base, 'message': help_msg, 'conversation_stage': 'pca'})
            
            # Verificar confirmação explícita
            if not _is_confirm(user_message) and not _yes(user_message):
                # Usuário deu resposta mas não confirmou - salvar e pedir confirmação
                ans['pca_tentative'] = user_message
                session.set_answers(ans)
                db.session.commit()
                confirm_msg = f"Entendi que sobre o PCA: {user_message}. Posso registrar assim e seguir para as normas legais?"
                return jsonify({**resp_base, 'message': confirm_msg, 'conversation_stage': 'pca'})
            
            # Confirmação explícita - avançar
            pca_value = ans.get('pca_tentative', user_message)
            ans['pca'] = pca_value
            session.set_answers(ans)
            session.conversation_stage = 'legal_norms'
            session.updated_at = datetime.utcnow()
            db.session.commit()
            
            # Gerar transição em prosa natural
            legal_msg = _generate_prose_transition('pca_to_legal_norms', {'pca_response': pca_value})
            if not legal_msg:
                legal_msg = "Entendido. Sobre as normas legais: costumamos usar a Lei 14.133/21 como base, junto com o regulamento local aplicável. Se houver algo específico do setor ou proteção de dados envolvida, podemos incluir. O que faz sentido para você?"
            
            return jsonify({**resp_base, 'message': legal_msg, 'conversation_stage': 'legal_norms'})

        if session.conversation_stage == 'legal_norms':
            ans = session.get_answers() or {}
            # Detectar ajuda
            if any(k in user_message.lower() for k in ['ajuda', 'como', 'não sei', 'quais', 'me ajuda']):
                help_msg = "As principais são: 1) Lei 14.133/2021 (nova lei de licitações), 2) Regulamento/Decreto local do seu ente, 3) LGPD se houver dados pessoais, 4) Normas técnicas do setor (ex: ABNT, vigilância sanitária). Você pode escolher por número ou me dizer quais aplicam ao seu caso. Te parece bom assim?"
                return jsonify({**resp_base, 'message': help_msg, 'conversation_stage': 'legal_norms'})
            
            # Parse resposta
            picks = _pick_numbers(user_message, 4)
            if picks:
                mapping = {1:"Lei 14.133/2021", 2:"Regulamento/Decreto local", 3:"LGPD", 4:"Normas técnicas/regulatórias do setor"}
                tentative_norms = ', '.join(mapping[i] for i in picks)
            elif _is_uncertain(user_message):
                tentative_norms = "Lei 14.133/2021; Regulamento local; (demais a confirmar)"
            else:
                tentative_norms = user_message
            
            # Verificar confirmação
            if not _is_confirm(user_message) and not _yes(user_message):
                ans['legal_norms_tentative'] = tentative_norms
                session.set_answers(ans)
                db.session.commit()
                confirm_msg = f"Entendi: {tentative_norms}. Posso seguir com essas normas?"
                return jsonify({**resp_base, 'message': confirm_msg, 'conversation_stage': 'legal_norms'})
            
            # Confirmado - avançar
            ans['legal_norms'] = ans.get('legal_norms_tentative', tentative_norms)
            session.set_answers(ans)
            session.conversation_stage = 'qty_value'
            session.updated_at = datetime.utcnow()
            db.session.commit()
            
            # Gerar transição em prosa natural
            qty_msg = _generate_prose_transition('legal_norms_to_qty', {})
            if not qty_msg:
                qty_msg = "Certo. Sobre quantitativo e valor: dá para estimar por consumo recente, contratos similares ou pesquisa de mercado. Se preferir, pode me passar a unidade e ordem de grandeza que ajustamos depois."
            
            return jsonify({**resp_base, 'message': qty_msg, 'conversation_stage': 'qty_value'})

        if session.conversation_stage in ('qty_value', 'quantitativo_valor'):
            ans = session.get_answers() or {}
            # Detectar ajuda
            if any(k in user_message.lower() for k in ['ajuda', 'como', 'não sei', 'não tenho']):
                help_msg = "Você pode estimar baseado em: histórico de consumo, contratos anteriores, ou pesquisa de mercado. Se não tiver agora, pode indicar só a unidade (ex: 'por mês', 'por usuário') e ordem de grandeza. Quando tiver uma ideia, me passa e eu confirmo. Ok?"
                return jsonify({**resp_base, 'message': help_msg, 'conversation_stage': session.conversation_stage})
            
            # Parse resposta
            if _is_uncertain(user_message):
                tentative_qty = "A definir (estimativa pendente). Unidade/escopo a detalhar."
            else:
                tentative_qty = user_message
            
            # Verificar confirmação
            if not _is_confirm(user_message) and not _yes(user_message):
                ans['qty_value_tentative'] = tentative_qty
                session.set_answers(ans)
                db.session.commit()
                confirm_msg = f"Registrei: {tentative_qty}. Posso seguir para o parcelamento com essa informação?"
                return jsonify({**resp_base, 'message': confirm_msg, 'conversation_stage': session.conversation_stage})
            
            # Confirmado - avançar
            qty_value = ans.get('qty_value_tentative', tentative_qty)
            ans['qty_value'] = qty_value
            session.set_answers(ans)
            session.conversation_stage = 'installment'
            session.updated_at = datetime.utcnow()
            db.session.commit()
            
            # Gerar transição em prosa natural
            installment_msg = _generate_prose_transition('qty_to_installment', {'qty_response': qty_value})
            if not installment_msg:
                installment_msg = "Entendido. Sobre parcelamento: dividir em lotes pode ampliar a competição e escalonar entregas; por outro lado, aumenta a logística. Pelo que descrevemos, faz sentido parcelar ou manter como item único?"
            
            return jsonify({**resp_base, 'message': installment_msg, 'conversation_stage': 'installment'})

        if session.conversation_stage == 'installment':
            ans = session.get_answers() or {}
            # Detectar ajuda
            if any(k in user_message.lower() for k in ['ajuda', 'como', 'não sei', 'me ajuda']):
                help_msg = "Parcelar pode aumentar a competição (mais empresas conseguem participar de lotes menores) e escalonar entregas. Mas também aumenta a gestão contratual. Se não souber agora, pode dizer 'a definir' e ajustamos depois. Te parece razoável?"
                return jsonify({**resp_base, 'message': help_msg, 'conversation_stage': 'installment'})
            
            # Parse resposta
            tentative_installment = "não" if _no(user_message) else ("sim" if _yes(user_message) else user_message)
            
            # Verificar confirmação
            if not _is_confirm(user_message) and tentative_installment not in ['sim', 'não']:
                ans['installment_tentative'] = tentative_installment
                session.set_answers(ans)
                db.session.commit()
                confirm_msg = f"Entendi: {tentative_installment}. Posso fechar assim e montar o resumo?"
                return jsonify({**resp_base, 'message': confirm_msg, 'conversation_stage': 'installment'})
            
            # Confirmado - avançar para resumo
            ans['installment'] = ans.get('installment_tentative', tentative_installment)
            session.set_answers(ans)
            
            # Monta RESUMO e só gera sob comando explícito
            resumo_lines = [
                f"Necessidade: {session.necessity}",
                "",
                "Análise de alternativas para atendimento:",
            ]
            # Traz todas as opções com análise
            opts = ans.get('solution_options') or []
            if opts:
                for i,o in enumerate(opts, start=1):
                    resumo_lines.append(f"  {i}) {o.get('opcao')}: {o.get('analise')}")
            else:
                resumo_lines.append("  (Alternativas registradas não disponíveis)")
            resumo_lines += [
                f"Opção escolhida: {ans.get('solution_path')} — {ans.get('solution_justification')}",
                "",
                f"PCA: {ans.get('pca')}",
                f"Normas legais: {ans.get('legal_norms')}",
                f"Quantitativo/Valor: {ans.get('qty_value')}",
                f"Parcelamento: {ans.get('installment')}",
                "",
                'Este é o resumo do que coletamos. Quando estiver pronto para gerar o documento, só me avisar.'
            ]
            resumo = "\n".join(resumo_lines)
            session.conversation_stage = 'summary'
            session.updated_at = datetime.utcnow()
            db.session.commit()
            return jsonify({
                **resp_base,
                'message': resumo,
                'conversation_stage': 'summary'
            })

        # SUMMARY: Wait for explicit generation command
        # Geração do ETP (apenas quando o usuário pedir "pode gerar", "gerar etp", etc.)
        if session.conversation_stage == 'summary':
            lm = _normalize(user_message).lower()
            if any(k in lm for k in ["pode gerar","gerar etp","gera etp","ok gerar","gerar o etp"]):
                ans = session.get_answers() or {}
                prompt = (
                    f"Necessidade: {session.necessity}\n"
                    f"Análise de alternativas: {ans.get('solution_options')}\n"
                    f"Escolha: {ans.get('solution_path')} — {ans.get('solution_justification')}\n"
                    f"PCA: {ans.get('pca')}\n"
                    f"Normas legais: {ans.get('legal_norms')}\n"
                    f"Quantitativo/Valor: {ans.get('qty_value')}\n"
                    f"Parcelamento: {ans.get('installment')}\n"
                    "Redija uma seção de Estudo Técnico Preliminar clara e coerente, citando as alternativas "
                    "e justificando a decisão à luz da necessidade apresentada. Estruture em subtítulos."
                )
                try:
                    final_text = final_gen.generate(prompt)
                    return jsonify({**resp_base,
                                    'message': "Segue a versão para o ETP (prévia):\n\n" + final_text,
                                    'etp_preview': final_text,
                                    'conversation_stage': 'done'})
                except Exception:
                    return jsonify({**resp_base,
                                    'message': 'Não consegui gerar agora. Revise o resumo acima e tente "gerar etp" novamente. '
                                               'Seu progresso foi mantido.',
                                    'conversation_stage': 'summary'})

        # FASE: legal_norms (PCA)
        if session.conversation_stage == 'legal_norms':
            ans = session.get_answers() or {}
            pca = ans.get('pca') or {}
            out = parse_legal_norms(user_message, ans)

            if out['intent'] == 'pca_yes':
                pca.update({'present': True})
                ans['pca'] = pca
                session.set_answers(ans)
                session.updated_at = datetime.utcnow()
                db.session.commit()
                return jsonify(_text_payload(session,
                    'Ótimo. Há previsão no PCA. Informe número/ano/item se desejar. Quando terminar, diga "seguir".'))

            if out['intent'] == 'pca_details':
                pca.update({'details': out['payload']['raw']})
                ans['pca'] = pca
                session.set_answers(ans)
                session.updated_at = datetime.utcnow()
                db.session.commit()
                return jsonify(_text_payload(session,
                    'Detalhes do PCA registrados. Diga "seguir" para avançar para pesquisa de preços.'))

            if out['intent'] == 'pca_no':
                pca.update({'present': False})
                ans['pca'] = pca
                session.set_answers(ans)
                session.updated_at = datetime.utcnow()
                db.session.commit()
                return jsonify(_text_payload(session,
                    'Sem previsão no PCA. Registre uma justificativa, se houver, e diga "seguir" para avançar.'))

            if out['intent'] == 'pca_unknown':
                pca.update({'present': None})
                ans['pca'] = pca
                session.set_answers(ans)
                session.updated_at = datetime.utcnow()
                db.session.commit()
                return jsonify(_text_payload(session,
                    'Sem certeza sobre o PCA. Se descobrir depois, informe os detalhes. Diga "seguir" para avançar.'))

            if out['intent'] == 'proceed_next':
                session.conversation_stage = 'price_research'
                session.updated_at = datetime.utcnow()
                db.session.commit()
                return jsonify(_text_payload(session,
                    "Vamos à pesquisa de preços. Informe o método utilizado (ex.: painel de preços, cotações com fornecedores, histórico de contratos).",
                    {'conversation_stage': 'price_research'}))

            # unclear
            return jsonify(_text_payload(session,
                'Não captei a informação sobre o PCA. Diga "sim", "não", "não sei" ou forneça número/ano/item.'))

        # FASE: price_research
        if session.conversation_stage == 'price_research':
            ans = session.get_answers() or {}
            pr = ans.get('price_research') or {'method': None, 'supplier_count': None, 'evidence_links': []}
            out = parse_price_research(user_message, ans)

            if out['intent'] == 'method_select':
                pr['method'] = out['payload']['method']
                ans['price_research'] = pr
                session.set_answers(ans)
                session.updated_at = datetime.utcnow()
                db.session.commit()
                return jsonify(_text_payload(session, "Método registrado. Informe a quantidade de fornecedores consultados, se aplicável."))

            if out['intent'] == 'supplier_count':
                pr['supplier_count'] = out['payload']['count']
                ans['price_research'] = pr
                session.set_answers(ans)
                session.updated_at = datetime.utcnow()
                db.session.commit()
                return jsonify(_text_payload(session, 'Quantidade registrada. Caso tenha links de evidência, envie-os agora. Depois diga "concluído".'))

            if out['intent'] == 'link_evidence':
                pr['evidence_links'] = list(set((pr.get('evidence_links') or []) + out['payload']['urls']))
                ans['price_research'] = pr
                session.set_answers(ans)
                session.updated_at = datetime.utcnow()
                db.session.commit()
                return jsonify(_text_payload(session, 'Link registrado. Envie mais links se houver, ou diga "concluído".'))

            if out['intent'] == 'mark_done':
                ans['price_research'] = pr
                session.set_answers(ans)
                session.conversation_stage = 'legal_basis'
                session.updated_at = datetime.utcnow()
                db.session.commit()
                return jsonify(_text_payload(session, "Pesquisa de preços concluída. Agora, indique a base legal aplicável e eventuais observações.",
                                             {'conversation_stage': 'legal_basis'}))

            return jsonify(_text_payload(session, 'Você pode me passar o método usado, quantas fontes consultou, e os links. Quando tiver enviado tudo, é só avisar que concluiu.'))

        # FASE: legal_basis
        if session.conversation_stage == 'legal_basis':
            ans = session.get_answers() or {}
            lb = ans.get('legal_basis') or {'text': None, 'notes': []}
            out = parse_legal_basis(user_message, ans)

            if out['intent'] == 'legal_basis_set':
                lb['text'] = out['payload']['text']
                ans['legal_basis'] = lb
                session.set_answers(ans)
                session.updated_at = datetime.utcnow()
                db.session.commit()
                return jsonify(_text_payload(session, 'Base legal registrada. Se quiser, envie observações ou diga "finalizar".'))

            if out['intent'] == 'legal_basis_notes':
                lb['notes'] = list((lb.get('notes') or [])) + [out['payload']['text']]
                ans['legal_basis'] = lb
                session.set_answers(ans)
                session.updated_at = datetime.utcnow()
                db.session.commit()
                return jsonify(_text_payload(session, 'Observação registrada. Envie mais observações ou diga "finalizar".'))

            if out['intent'] == 'finalize':
                ans['legal_basis'] = lb
                session.set_answers(ans)
                session.conversation_stage = 'done'
                session.updated_at = datetime.utcnow()
                db.session.commit()
                return jsonify(_text_payload(session, "Coleta concluída. Podemos gerar o documento ou revisar se preferir.",
                                             {'conversation_stage': 'done'}))

            return jsonify(_text_payload(session, 'Envie a base legal, observações ou diga "finalizar".'))

        # PASSO 5: If no session necessity exists, we need to capture it
        if not session.necessity:
            # Try to analyze the need with etp_generator if available, otherwise just use the message
            need_description = None
            contains_need = False
            
            try:
                if etp_generator and hasattr(etp_generator, 'client'):
                    from domain.usecase.etp.utils_parser import analyze_need_safely
                    contains_need, need_description = analyze_need_safely(user_message, etp_generator.client)
                    print(f"🔹 [ANALYZER] contains_need={contains_need}, description='{need_description or 'None'}'")
                else:
                    # Fallback: assume the message is the necessity if it's long enough
                    if len(user_message.strip()) > 10:
                        contains_need = True
                        need_description = user_message.strip()
                        print(f"🔹 [ANALYZER FALLBACK] Usando mensagem como necessidade: {need_description}")
            except Exception as e:
                # If analysis fails, use the message as necessity
                print(f"🔸 [ANALYZER ERROR] {e} - usando fallback")
                if len(user_message.strip()) > 10:
                    contains_need = True
                    need_description = user_message.strip()

            if contains_need and need_description:
                print(f"🔹 [LOCK] Necessidade identificada e travada: {need_description}")
                
                # GUARDRAIL: LOCK NECESSITY and advance to suggest_requirements
                # After suggest_requirements, ALWAYS go to refine_requirements (mandatory state flow)
                session.necessity = need_description
                session.conversation_stage = 'suggest_requirements'
                session.updated_at = datetime.utcnow()
                db.session.commit()
                
                print(f"🔹 [DEPOIS] Sessão: {session.session_id}, estágio: {session.conversation_stage}")

                # GUARDRAIL: Validate generator exists before generating requirements
                gen = _get_simple_generator()
                is_valid, error_msg = validate_generator_exists(gen)
                
                if not is_valid:
                    print(f"🔸 [GENERATOR ERROR] {error_msg}")
                    # Return controlled error without changing state
                    error_response = handle_http_error(500, error_msg, session.conversation_stage)
                    return jsonify({
                        **resp_base,
                        **error_response
                    })
                
                # Generate requirements using the simple generator with FALLBACK (NEVER 500)
                try:
                    suggested = gen.suggest_requirements(need_description)
                    origin = gen.__class__.__name__
                    print(f"🔹 [GENERATOR] Usando {origin} para gerar requisitos")
                except Exception as e:
                    # GUARDRAIL: Handle HTTP error without changing state
                    print(f"🔸 [GENERATOR ERROR] {e} - usando fallback explícito")
                    traceback.print_exc()
                    # Use FallbackGenerator
                    try:
                        origin = "FallbackGenerator"
                        suggested = FallbackGenerator().suggest_requirements(need_description)
                    except Exception as fallback_error:
                        # If even fallback fails, return controlled error
                        error_response = handle_http_error(500, str(fallback_error), session.conversation_stage)
                        return jsonify({
                            **resp_base,
                            **error_response
                        })
                
                # Convert simple list to structured format
                structured_requirements = []
                for i, req_text in enumerate(suggested, 1):
                    structured_requirements.append({
                        'id': f'R{i}',
                        'text': req_text,
                        'justification': 'Requisito sugerido pelo sistema'
                    })
                
                # Store requirements in session
                session.set_requirements(structured_requirements)
                
                # GUARDRAIL: After suggest_requirements, ALWAYS go to refine_requirements
                session.conversation_stage = get_next_state_after_suggestion(session.conversation_stage)
                session.updated_at = datetime.utcnow()
                db.session.commit()
                
                print(f"🔹 [STATE_MACHINE] Transição: suggest_requirements → {session.conversation_stage}")
                
                # REQUIREMENT 4: Generate unified single-message response with contextual intro and rationale
                # Build contextual intro based on necessity
                necessity_lower = need_description.lower()
                if 'seguranç' in necessity_lower or 'vigilân' in necessity_lower:
                    intro_paragraph = "Faz sentido o que você descreveu. Essa necessidade tem impacto direto na segurança e conformidade, então vou organizar os requisitos de forma objetiva."
                elif 'manutençã' in necessity_lower or 'conservaçã' in necessity_lower:
                    intro_paragraph = "Entendi. A manutenção preventiva e o acompanhamento técnico são essenciais para garantir a continuidade operacional. Organizei os requisitos pensando nisso."
                elif 'tecnologi' in necessity_lower or 'sistema' in necessity_lower or 'software' in necessity_lower:
                    intro_paragraph = "Perfeito. Para uma contratação de tecnologia, precisamos requisitos que garantam entrega técnica, suporte adequado e segurança da informação. Veja o que pensei."
                elif 'consultor' in necessity_lower or 'assessor' in necessity_lower:
                    intro_paragraph = "Entendi sua demanda. Para consultoria especializada, os requisitos precisam focar em qualificação técnica, experiência comprovada e metodologia clara. Aqui estão."
                else:
                    intro_paragraph = "Faz sentido o que você descreveu. Essa necessidade precisa de requisitos que garantam qualidade, conformidade e continuidade da prestação. Vou detalhar."
                
                # Build numbered list
                req_list_text = "\n".join([f"{i}. {req['text']}" for i, req in enumerate(structured_requirements, 1)])
                
                # Build contextual rationale
                if 'seguranç' in necessity_lower or 'vigilân' in necessity_lower:
                    rationale_paragraph = "Escolhi esses pontos porque estruturam capacidade técnica, conformidade regulatória e operação contínua. Isso reduz risco de segurança e garante aderência legal."
                elif 'manutençã' in necessity_lower or 'conservaçã' in necessity_lower:
                    rationale_paragraph = "Esses requisitos cobrem experiência técnica, disponibilidade de equipe e insumos, e continuidade do serviço. Isso minimiza paradas não planejadas e mantém a infraestrutura funcionando."
                elif 'tecnologi' in necessity_lower or 'sistema' in necessity_lower or 'software' in necessity_lower:
                    rationale_paragraph = "Os requisitos focam em competência técnica comprovada, suporte ativo, segurança da informação e escalabilidade. Isso garante entrega funcional e reduz riscos de descontinuidade tecnológica."
                else:
                    rationale_paragraph = "Escolhi esses pontos porque estruturam competência comprovada, capacidade técnica e operação contínua com aderência regulatória. Isso amarra a prestação do serviço sem lacunas e reduz risco de entrega e de conformidade."
                
                # Combine all parts
                ai_response = (
                    f"{intro_paragraph}\n\n"
                    f"{req_list_text}\n\n"
                    f"{rationale_paragraph}\n\n"
                    f"Se quiser ajustar algo, diga no seu jeito mesmo e eu já refaço."
                )
                
                # PASSO 6: Return with unified contract (ALWAYS SUCCESS)
                # Return refine_requirements stage (not suggest_requirements)
                return jsonify({
                    **resp_base,
                    "kind": "requirements_suggestion", 
                    "necessity": session.necessity,
                    "requirements": structured_requirements,
                    "intro_paragraph": intro_paragraph,
                    "rationale_paragraph": rationale_paragraph,
                    "ai_response": ai_response,
                    "message": ai_response,
                    "conversation_stage": session.conversation_stage  # Will be refine_requirements
                })
            
            # PASSO 5: If necessity not detected, ask for it
            print(f"🔹 [NO_LOCK] contains_need=False ou parse falhou → pedindo necessidade")
            return jsonify({
                **resp_base,
                'kind': 'text',
                'ai_response': 'Olá! Para começar o ETP, preciso entender a necessidade. Qual é a descrição da necessidade da contratação?',
                'message': 'Olá! Para começar o ETP, preciso entender a necessidade. Qual é a descrição da necessidade da contratação?',
                'conversation_stage': 'collect_need'
            })

        # If we have a necessity but no clear command was processed, use LLM with updated prompts
        # This handles cases where the command parser returned 'unclear'
        if session.conversation_stage in ['suggest_requirements', 'review_requirements']:
            system_content = f"""Você é um especialista em licitações que está ajudando a revisar requisitos para um Estudo Técnico Preliminar (ETP).

IMPORTANTE: Você está na fase de revisão de requisitos. A necessidade já está definida: {session.necessity}

JAMAIS trate mensagens como "ajuste o último", "remover 2", "trocar 3" como nova necessidade. Apenas atualize a lista existente.

Foque apenas em:
- Confirmar requisitos apresentados
- Processar ajustes específicos solicitados
- Adicionar novos requisitos
- Remover requisitos específicos

A necessidade está TRAVADA e não deve ser alterada."""
        
        # Continue with normal LLM processing for other stages
        system_content = f"""Você é um especialista em licitações que está ajudando a coletar informações para um Estudo Técnico Preliminar (ETP).

Necessidade já capturada: {session.necessity}

Seu objetivo é conduzir uma conversa natural e fluida para coletar as seguintes informações obrigatórias:
1. Descrição da necessidade da contratação ✓ (já coletada)
2. Confirmação dos requisitos sugeridos ✓ (já processada) 
3. Se há previsão no PCA (Plano de Contratações Anual)
4. Quais normas legais pretende utilizar
5. Qual o quantitativo e valor estimado
6. Se haverá parcelamento da contratação

Mantenha o tom conversacional e profissional. Faça uma pergunta por vez.
Não repita informações já coletadas."""

        # Build conversation context
        messages = [{"role": "system", "content": system_content}]
        
        # Add conversation history
        for msg in conversation_history[-5:]:  # Last 5 messages for context
            role = "user" if msg.get('sender') == 'user' else "assistant"
            messages.append({"role": role, "content": msg.get('text', '')})
        
        # Add current user message
        messages.append({"role": "user", "content": user_message})

        # Generate response - ensure client is initialized
        if not etp_generator or not hasattr(etp_generator, 'client') or not etp_generator.client:
            # Fallback: use get_llm_client() helper
            client = get_llm_client()
            if not client:
                return jsonify({
                    **resp_base,
                    'kind': 'text',
                    'ai_response': 'Erro: cliente OpenAI não inicializado. Verifique a configuração.',
                    'message': 'Erro: cliente OpenAI não inicializado. Verifique a configuração.',
                    'conversation_stage': session.conversation_stage
                })
        else:
            client = etp_generator.client
        
        model_name = get_model_name()
        response = client.chat.completions.create(
            model=model_name,
            messages=messages,
            max_tokens=800,
            temperature=0.7
        )

        ai_response = response.choices[0].message.content.strip()

        # PASSO 6: Return with unified contract
        return jsonify({
            **resp_base,
            'kind': 'text',
            'ai_response': ai_response,
            'message': ai_response,
            'conversation_stage': session.conversation_stage
        })

    except Exception as e:
        print(f"🔸 Erro na conversa: {e}")
        return jsonify({
            'success': False,
            'error': f'Erro na conversa: {str(e)}'
        }), 500

@etp_dynamic_bp.route('/confirm-requirements', methods=['POST'])
@cross_origin()
def confirm_requirements():
    """Processa a confirmação, ajuste ou rejeição de requisitos pelo usuário"""
    try:
        _ensure_initialized()
        if not etp_generator:
            return jsonify({'error': 'Gerador ETP não configurado'}), 500

        # PASSO 8: Reutilizar session_id corretamente
        data = request.get_json(force=True)
        sid = (data.get("session_id") or "").strip() or None
        session = EtpSession.query.filter_by(session_id=sid).first() if sid else None
        
        if not session:
            return jsonify({'error': 'Sessão não encontrada'}), 404
            
        # PASSO 8: Base de resposta com session_id
        resp_base = {"success": True, "session_id": session.session_id}
        
        user_action = data.get('action', '')  # 'accept', 'modify', 'add', 'remove'
        requirements = data.get('requirements', [])
        user_message = data.get('message', '')

        # Standardize the next question for consistent flow
        next_question = "👉 **Agora, há previsão no PCA (Plano de Contratações Anual)?**"
        
        # Processar a ação do usuário
        if user_action == 'accept':
            # Usuário aceitou todos os requisitos
            confirmed_requirements = requirements
            ai_response = f"**Perfeito! Requisitos confirmados.**\n\n{next_question}"
            print(f"🔹 Usuário aceitou os requisitos: {len(confirmed_requirements)} requisitos confirmados")

        elif user_action == 'modify':
            # Usuário quer modificar alguns requisitos
            print(f"🔹 Usuário solicitou modificação nos requisitos: {user_message}")
            modify_prompt = f"""
            O usuário quer modificar os requisitos sugeridos. Processe a solicitação:

            Requisitos originais: {requirements}
            Solicitação do usuário: "{user_message}"

            Retorne APENAS um JSON com:
            - "updated_requirements": array com requisitos atualizados
            - "explanation": breve explicação das mudanças feitas
            """

            # Ensure client is initialized
            if not etp_generator or not hasattr(etp_generator, 'client') or not etp_generator.client:
                client = get_llm_client()
                if not client:
                    # Fallback: keep original requirements
                    confirmed_requirements = requirements
                    ai_response = f"**Mantive os requisitos originais.** Cliente não inicializado.\n\n{next_question}"
                    # Skip the LLM call
                    response = None
            else:
                client = etp_generator.client
            
            if client:
                model_name = get_model_name()
                response = client.chat.completions.create(
                    model=model_name,
                    messages=[
                        {"role": "system", "content": "Você é um especialista em requisitos técnicos para licitações."},
                        {"role": "user", "content": modify_prompt}
                    ],
                    max_tokens=800,
                    temperature=0.7
                )

            # PASSO 5: Parse response safely
            if response:
                from domain.usecase.etp.utils_parser import parse_json_relaxed
                modification_result = parse_json_relaxed(response.choices[0].message.content.strip())
                
                if modification_result and 'updated_requirements' in modification_result:
                    confirmed_requirements = modification_result['updated_requirements']
                    explanation = modification_result.get('explanation', 'Requisitos atualizados conforme solicitado.')
                    ai_response = f"**Requisitos atualizados!**\n\n{explanation}\n\n{next_question}"
                else:
                    # Fallback if parsing fails
                    confirmed_requirements = requirements
                    ai_response = f"**Mantive os requisitos originais.** Não consegui processar a modificação solicitada.\n\n{next_question}"
            # else: confirmed_requirements and ai_response already set in the fallback above

        else:
            # Ação não reconhecida - manter requisitos originais
            confirmed_requirements = requirements
            ai_response = f"**Requisitos mantidos.**\n\n{next_question}"

        # Armazenar requisitos confirmados na sessão
        answers = session.get_answers()
        answers['confirmed_requirements'] = confirmed_requirements
        session.set_answers(answers)
        session.updated_at = datetime.utcnow()

        db.session.commit()

        # PASSO 8: Não retroceder estágio - manter em review_requirements ou avançar
        if user_action == 'accept':
            session.conversation_stage = 'legal_norms'  # Avançar para próxima fase
        else:
            session.conversation_stage = 'review_requirements'  # Manter em revisão
        
        session.updated_at = datetime.utcnow()
        db.session.commit()

        return jsonify({
            **resp_base,
            'kind': 'requirements_confirmed',
            'confirmed_requirements': confirmed_requirements,
            'ai_response': ai_response,
            'message': ai_response,
            'conversation_stage': session.conversation_stage
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'Erro ao confirmar requisitos: {str(e)}'
        }), 500

@etp_dynamic_bp.route('/suggest-requirements', methods=['POST'])
@cross_origin()
def suggest_requirements():
    """Sugere requisitos baseados na necessidade identificada"""
    try:
        _ensure_initialized()
        if not etp_generator:
            return jsonify({'error': 'Gerador ETP não configurado'}), 500

        data = request.get_json()
        necessity = data.get('necessity', '').strip()

        if not necessity:
            return jsonify({'error': 'Necessidade é obrigatória'}), 400

        # Usar RAG para encontrar requisitos similares
        rag_results = search_requirements("generic", necessity, k=5)
        
        # Gerar requisitos no formato R# — descrição (sem justificativas)
        requirements_prompt = f"""
        Baseado na necessidade: "{necessity}"
        
        E nos seguintes exemplos de requisitos similares:
        {json.dumps(rag_results, indent=2)}
        
        Gere uma lista de 3-5 requisitos específicos e objetivos para esta contratação.
        
        FORMATO OBRIGATÓRIO:
        Retorne APENAS uma lista de requisitos no formato:
        R1 — <descrição do requisito em uma única linha>
        R2 — <descrição do requisito em uma única linha>
        R3 — <descrição do requisito em uma única linha>
        
        REGRAS ESTRITAS:
        - NÃO inclua justificativas, explicações ou qualquer texto adicional
        - Cada linha deve começar com R seguido de número, espaço, travessão — e espaço
        - Cada requisito em uma única linha
        - Sem bullets, asteriscos, numeração dupla, tabelas ou JSON
        - Requisitos devem ser específicos e verificáveis
        
        Retorne SOMENTE as linhas no formato R# — descrição, nada mais.
        """
        
        model_name = get_model_name()
        response = etp_generator.client.chat.completions.create(
            model=model_name,
            messages=[
                {"role": "system", "content": "Você é um especialista em licitações que gera requisitos técnicos precisos no formato R# — descrição, sem justificativas."},
                {"role": "user", "content": requirements_prompt}
            ],
            max_tokens=1000,
            temperature=0.7
        )
        
        # Parse response in R# — format
        result_raw = response.choices[0].message.content.strip()
        
        # Parse requirements from R# — format
        requirements = []
        for line in result_raw.split('\n'):
            line = line.strip()
            if line and (line.startswith('R') or (line and line[0].isdigit())):
                requirements.append(line)
        
        # Fallback if no requirements were parsed
        if not requirements:
            requirements = [
                f"R1 — Requisitos técnicos específicos para {necessity.lower()}",
                "R2 — Especificações técnicas adequadas ao objeto da contratação",
                "R3 — Comprovação de capacidade técnica adequada"
            ]

        return jsonify({
            'success': True,
            'requirements': requirements,
            'message': 'Requisitos sugeridos baseados na necessidade identificada e na base de conhecimento.',
            'necessity': necessity,
            'rag_sources': len(rag_results),
            'timestamp': datetime.now().isoformat()
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'Erro ao sugerir requisitos: {str(e)}'
        }), 500

@etp_dynamic_bp.route('/analyze-response', methods=['POST'])
@cross_origin()
def analyze_response_stub():
    return jsonify({'error': 'Rota removida. Use POST /api/etp-dynamic/conversation'}), 404

@etp_dynamic_bp.route('/generate-document', methods=['POST'])
@cross_origin()
def generate_document():
    data = request.get_json(silent=True) or {}
    sid = (data.get("session_id") or "").strip() or None
    if not sid:
        return jsonify({'success': False, 'error': 'session_id ausente'}), 400
    session = EtpSession.query.filter_by(session_id=sid).first()
    if not session:
        return jsonify({'success': False, 'error': 'Sessão não encontrada'}), 404
    
    # Check if user has confirmed requirements
    answers = session.get_answers() or {}
    user_confirmed = answers.get('user_confirmed_requirements', False)
    
    if not user_confirmed and session.conversation_stage in ['suggest_requirements', 'review_requirements']:
        return jsonify({
            'success': False,
            'error': 'Os requisitos precisam ser confirmados antes de gerar o documento.',
            'message': 'Por favor, confirme os requisitos antes de gerar o ETP.',
            'kind': 'confirmation_required',
            'conversation_stage': session.conversation_stage
        }), 400
    
    # Compor e renderizar
    doc_json = compose_etp_document(session)
    html = render_etp_html(doc_json)
    etp_doc = EtpDocument(session_id=sid, doc_json=doc_json, html=html)
    db.session.add(etp_doc); db.session.commit()
    return jsonify(_text_payload(session, f"Perfeito. Requisitos confirmados. Vou gerar o ETP com base neles.\n\nDocumento gerado com sucesso (ID: {etp_doc.id}).",
                                 {'doc_id': etp_doc.id, 'kind': 'doc_generated'}))

@etp_dynamic_bp.route('/document/<int:doc_id>/html', methods=['GET'])
@cross_origin()
def get_document_html(doc_id: int):
    etp_doc = EtpDocument.query.get(doc_id)
    if not etp_doc:
        return jsonify({'success': False, 'error': 'Documento não encontrado'}), 404
    return jsonify({'success': True, 'doc_id': etp_doc.id, 'html': etp_doc.html})

@etp_dynamic_bp.route('/document/<int:doc_id>/download-docx', methods=['GET'])
@cross_origin()
def download_docx(doc_id: int):
    etp_doc = EtpDocument.query.get(doc_id)
    if not etp_doc:
        return jsonify({'success': False, 'error': 'Documento não encontrado'}), 404
    d = DocxDocument()
    d.add_heading(etp_doc.doc_json.get('title', 'Documento'), level=1)
    for s in etp_doc.doc_json.get('sections', []):
        d.add_heading(s.get('title', ''), level=2)
        if s.get('content'):
            d.add_paragraph(str(s['content']))
        if s.get('items'):
            for it in s['items']:
                d.add_paragraph(f"- {it}")
        for key in ['details', 'method', 'supplier_count']:
            if s.get(key) is not None:
                d.add_paragraph(f"{key}: {s.get(key)}")
        if s.get('evidence_links'):
            d.add_paragraph("Evidências:")
            for link in s['evidence_links']:
                d.add_paragraph(f"* {link}")
        if s.get('notes'):
            d.add_paragraph("Observações:")
            for note in s['notes']:
                d.add_paragraph(f"* {note}")
    buf = BytesIO()
    d.save(buf); buf.seek(0)
    return send_file(
        buf,
        as_attachment=True,
        download_name=f"ETP_{etp_doc.session_id}_{etp_doc.id}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@etp_dynamic_bp.route('/regen-requirements', methods=['POST'])
@cross_origin()
def regen_requirements():
    """
    Regenerate requirements when frontend detects empty requirements list.
    Uses corrective prompt to ensure valid output.
    
    Request JSON:
    {
        "conversation_id": "<uuid>",
        "stage": "collect_need" | "refine"
    }
    
    Response JSON:
    {
        "success": true,
        "intro": "...",
        "requirements": ["1. ...", "2. ..."],
        "justification": "..."
    }
    """
    try:
        data = request.get_json()
        conversation_id = data.get('conversation_id')
        stage = data.get('stage', 'collect_need')
        
        if not conversation_id:
            return jsonify({'success': False, 'error': 'conversation_id é obrigatório'}), 400
        
        # Load conversation
        conv = ConversationRepo.get(conversation_id)
        if not conv:
            return jsonify({'success': False, 'error': 'Conversa não encontrada'}), 404
        
        # Get session
        session = EtpSession.query.filter_by(session_id=str(conversation_id)).first()
        if not session:
            return jsonify({'success': False, 'error': 'Sessão não encontrada'}), 404
        
        necessity = session.necessity or "contratação"
        
        # Build history
        history = []
        try:
            messages = MessageRepo.list_for_conversation(conversation_id) or []
        except Exception:
            messages = []
        for msg in messages[-10:]:
            history.append({'role': msg.role, 'content': msg.content})
        
        # Get RAG context
        from rag.retrieval import retrieve_for_stage
        from application.ai.generator import generate_answer
        
        context_chunks = retrieve_for_stage(necessity, stage, k=12)
        
        rag_context = {
            'chunks': context_chunks,
            'necessity': necessity,
            'requirements': session.get_requirements() or []
        }
        
        # Regenerate with corrective prompt embedded in generate_answer
        result = generate_answer(stage, history, necessity, rag_context)
        
        # Update session with regenerated requirements
        requirements = result.get('requirements', [])
        if requirements:
            session.set_requirements(requirements)
            db.session.commit()
        
        return jsonify({
            'success': True,
            'intro': result.get('intro', ''),
            'requirements': requirements,
            'justification': result.get('justification', '')
        })
        
    except Exception as e:
        logger.error(f"[REGEN_REQ] Error: {e}")
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': f'Erro ao regenerar requisitos: {str(e)}'
        }), 500

@etp_dynamic_bp.route('/admin/repair-answers', methods=['POST'])
@cross_origin()
def repair_answers():
    """
    Endpoint de manutenção para normalizar sessões antigas com answers como string.
    Percorre todas as sessões e força a normalização via get_answers().
    """
    try:
        fixed_count = 0
        for session in EtpSession.query.all():
            before = session.answers
            # Normaliza em memória usando get_answers() que atualiza self.answers
            _ = session.get_answers()
            # Verifica se houve mudança
            if before != session.answers:
                fixed_count += 1
        
        # Salva todas as mudanças de uma vez
        db.session.commit()
        
        return jsonify({
            'success': True,
            'fixed': fixed_count,
            'message': f'{fixed_count} sessão(ões) normalizada(s) com sucesso.'
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({
            'success': False,
            'error': f'Erro ao reparar sessões: {str(e)}'
        }), 500
