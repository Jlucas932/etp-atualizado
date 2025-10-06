import os
import tempfile
import re
from datetime import datetime
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

class WordFormatterWithBorders:
    """Formatador Word com bordas baseado no modelo da concorrência"""
    
    def __init__(self):
        self.blue_color = RGBColor(31, 78, 121)  # Azul escuro #1f4e79
        self.white_color = RGBColor(255, 255, 255)  # Branco
        self.black_color = RGBColor(0, 0, 0)  # Preto
        
    def create_document_with_borders(self, content: str, session_data: Dict = None) -> str:
        """Cria documento Word com bordas e formatação baseada no modelo da concorrência"""
        try:
            doc = Document()
            
            # Configurar página e margens
            self._configure_page_with_borders(doc)
            
            # Criar estilos personalizados
            self._create_custom_styles(doc)
            
            # Adicionar cabeçalho institucional
            self._add_institutional_header(doc, session_data)
            
            # Adicionar título principal com fundo azul
            self._add_main_title_with_background(doc)
            
            # Adicionar introdução em caixa
            self._add_introduction_box(doc)
            
            # Processar e adicionar conteúdo com formatação
            self._process_content_with_formatting(doc, content)
            
            # Adicionar rodapé
            self._add_footer_with_info(doc)
            
            # Aplicar bordas ao documento inteiro
            self._apply_document_borders(doc)
            
            # Salvar documento
            doc_path = self._save_document_with_timestamp(doc)
            
            return doc_path
            
        except Exception as e:
            raise Exception(f"Erro ao criar documento Word com bordas: {str(e)}")
    
    def _configure_page_with_borders(self, doc: Document):
        """Configura página com margens para acomodar bordas"""
        sections = doc.sections
        for section in sections:
            # Margens maiores para acomodar bordas
            section.top_margin = Inches(1.2)     # 3cm
            section.bottom_margin = Inches(1.2)  # 3cm
            section.left_margin = Inches(1.4)    # 3.5cm
            section.right_margin = Inches(1.4)   # 3.5cm
            
            # Configurar cabeçalho e rodapé
            section.header_distance = Inches(0.8)
            section.footer_distance = Inches(0.8)
    
    def _create_custom_styles(self, doc: Document):
        """Cria estilos personalizados baseados no modelo"""
        styles = doc.styles
        
        # Estilo para títulos de seção (fundo azul)
        if 'Secao Azul ETP' not in [s.name for s in styles]:
            section_style = styles.add_style('Secao Azul ETP', WD_STYLE_TYPE.PARAGRAPH)
            
            section_format = section_style.paragraph_format
            section_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            section_format.space_before = Pt(18)
            section_format.space_after = Pt(12)
            section_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            
            section_font = section_style.font
            section_font.name = 'Times New Roman'
            section_font.size = Pt(14)
            section_font.bold = True
            section_font.color.rgb = self.white_color
        
        # Estilo para subtítulos
        if 'Subtitulo ETP' not in [s.name for s in styles]:
            subtitle_style = styles.add_style('Subtitulo ETP', WD_STYLE_TYPE.PARAGRAPH)
            
            subtitle_format = subtitle_style.paragraph_format
            subtitle_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            subtitle_format.space_before = Pt(12)
            subtitle_format.space_after = Pt(6)
            subtitle_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            
            subtitle_font = subtitle_style.font
            subtitle_font.name = 'Times New Roman'
            subtitle_font.size = Pt(12)
            subtitle_font.bold = True
            subtitle_font.color.rgb = self.black_color
        
        # Estilo para corpo do texto
        if 'Corpo ETP' not in [s.name for s in styles]:
            body_style = styles.add_style('Corpo ETP', WD_STYLE_TYPE.PARAGRAPH)
            
            body_format = body_style.paragraph_format
            body_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            body_format.line_spacing = 1.4
            body_format.space_after = Pt(6)
            body_format.first_line_indent = Inches(0.5)  # Recuo primeira linha
            
            body_font = body_style.font
            body_font.name = 'Times New Roman'
            body_font.size = Pt(12)
            body_font.color.rgb = self.black_color
        
        # Estilo para cabeçalho
        if 'Cabecalho ETP' not in [s.name for s in styles]:
            header_style = styles.add_style('Cabecalho ETP', WD_STYLE_TYPE.PARAGRAPH)
            
            header_format = header_style.paragraph_format
            header_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_format.space_after = Pt(12)
            
            header_font = header_style.font
            header_font.name = 'Times New Roman'
            header_font.size = Pt(11)
            header_font.bold = True
            header_font.color.rgb = self.black_color
    
    def _add_institutional_header(self, doc: Document, session_data: Dict = None):
        """Adiciona cabeçalho institucional baseado no modelo"""
        # Criar tabela para layout do cabeçalho
        header_table = doc.add_table(rows=1, cols=2)
        header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Célula esquerda - Informações do órgão
        left_cell = header_table.cell(0, 0)
        left_para = left_cell.paragraphs[0]
        left_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Adicionar brasão (emoji como placeholder)
        brasao_run = left_para.add_run("🏛️\n")
        brasao_run.font.size = Pt(24)
        
        # Informações do governo
        gov_run = left_para.add_run("Governo do Estado do Maranhão\n")
        gov_run.font.name = 'Times New Roman'
        gov_run.font.size = Pt(11)
        gov_run.font.bold = True
        
        secretaria_run = left_para.add_run("SECRETARIA DE ESTADO DA ADMINISTRAÇÃO - SEAD\n")
        secretaria_run.font.name = 'Times New Roman'
        secretaria_run.font.size = Pt(10)
        secretaria_run.font.bold = True
        
        salic_run = left_para.add_run("SECRETARIA ADJUNTA DE LICITAÇÕES E COMPRAS ESTRATÉGICAS - SALIC")
        salic_run.font.name = 'Times New Roman'
        salic_run.font.size = Pt(10)
        salic_run.font.bold = True
        
        # Célula direita - Caixa de protocolo
        right_cell = header_table.cell(0, 1)
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Adicionar caixa de protocolo
        protocolo_text = "SEAD/SALIC\nNº\nProc.:\nRub.:"
        protocolo_run = right_para.add_run(protocolo_text)
        protocolo_run.font.name = 'Times New Roman'
        protocolo_run.font.size = Pt(9)
        protocolo_run.font.bold = True
        
        # Aplicar borda à célula de protocolo
        self._add_cell_border(right_cell)
        
        # Remover bordas da tabela principal
        self._remove_table_borders(header_table)
        
        # Adicionar linha separadora
        separator_para = doc.add_paragraph()
        separator_run = separator_para.add_run("_" * 100)
        separator_run.font.size = Pt(8)
        separator_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_main_title_with_background(self, doc: Document):
        """Adiciona título principal com fundo azul"""
        title_para = doc.add_paragraph("ESTUDO TÉCNICO PRELIMINAR", style='Secao Azul ETP')
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Aplicar fundo azul
        self._add_blue_background(title_para)
        
        # Adicionar espaço
        doc.add_paragraph()
    
    def _add_introduction_box(self, doc: Document):
        """Adiciona introdução em caixa como no modelo"""
        intro_para = doc.add_paragraph()
        intro_run = intro_para.add_run(
            "O presente documento caracteriza a primeira etapa da fase de planejamento e apresenta "
            "os devidos estudos para a contratação de solução que melhor atenderá à necessidade "
            "descrita abaixo."
        )
        intro_run.font.name = 'Times New Roman'
        intro_run.font.size = Pt(12)
        intro_run.font.bold = True
        intro_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Adicionar segunda parte
        intro_para2 = doc.add_paragraph()
        intro_run2 = intro_para2.add_run(
            "O objetivo principal é identificar a necessidade e identificar a melhor solução para "
            "supri-la, em observância às normas vigentes e aos princípios que regem a Administração Pública."
        )
        intro_run2.font.name = 'Times New Roman'
        intro_run2.font.size = Pt(12)
        intro_run2.font.bold = True
        intro_para2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Aplicar borda à introdução
        self._add_paragraph_border(intro_para)
        self._add_paragraph_border(intro_para2)
        
        # Espaço após introdução
        doc.add_paragraph()
    
    def _process_content_with_formatting(self, doc: Document, content: str):
        """Processa conteúdo aplicando formatação baseada no modelo"""
        lines = content.split('\n')
        current_paragraph = []
        in_table = False
        table_data = []
        
        for line in lines:
            line = line.strip()
            
            if not line:
                # Linha vazia
                if current_paragraph:
                    self._add_formatted_paragraph(doc, '\n'.join(current_paragraph))
                    current_paragraph = []
                if in_table and table_data:
                    self._create_formatted_table(doc, table_data)
                    table_data = []
                    in_table = False
                continue
            
            # Verificar se é título de seção principal
            if self._is_main_section_title(line):
                # Finalizar conteúdo anterior
                if current_paragraph:
                    self._add_formatted_paragraph(doc, '\n'.join(current_paragraph))
                    current_paragraph = []
                if in_table and table_data:
                    self._create_formatted_table(doc, table_data)
                    table_data = []
                    in_table = False
                
                # Adicionar título de seção com fundo azul
                self._add_section_title_with_background(doc, line)
                
            elif self._is_subsection_title(line):
                # Finalizar conteúdo anterior
                if current_paragraph:
                    self._add_formatted_paragraph(doc, '\n'.join(current_paragraph))
                    current_paragraph = []
                if in_table and table_data:
                    self._create_formatted_table(doc, table_data)
                    table_data = []
                    in_table = False
                
                # Adicionar subtítulo
                self._add_subsection_title(doc, line)
                
            elif self._is_table_line(line):
                # Finalizar parágrafo anterior
                if current_paragraph:
                    self._add_formatted_paragraph(doc, '\n'.join(current_paragraph))
                    current_paragraph = []
                
                # Processar linha de tabela
                in_table = True
                table_data.append(line)
                
            else:
                # Finalizar tabela se necessário
                if in_table and table_data:
                    self._create_formatted_table(doc, table_data)
                    table_data = []
                    in_table = False
                
                # Adicionar ao parágrafo atual
                current_paragraph.append(line)
        
        # Finalizar conteúdo restante
        if current_paragraph:
            self._add_formatted_paragraph(doc, '\n'.join(current_paragraph))
        if in_table and table_data:
            self._create_formatted_table(doc, table_data)
    
    def _is_main_section_title(self, line: str) -> bool:
        """Verifica se é título de seção principal (1., 2., etc.)"""
        pattern = r'^\d+\.\s+[A-ZÁÀÂÃÉÊÍÓÔÕÚÇ\s]+$'
        return bool(re.match(pattern, line.upper()))
    
    def _is_subsection_title(self, line: str) -> bool:
        """Verifica se é título de subseção (1.1, 1.2, etc.)"""
        pattern = r'^\d+\.\d+\.?\s+[A-Za-záàâãéêíóôõúç]+'
        return bool(re.match(pattern, line))
    
    def _is_table_line(self, line: str) -> bool:
        """Verifica se é linha de tabela"""
        return '|' in line and line.count('|') >= 2
    
    def _add_section_title_with_background(self, doc: Document, title: str):
        """Adiciona título de seção com fundo azul"""
        title_para = doc.add_paragraph(title.upper(), style='Secao Azul ETP')
        self._add_blue_background(title_para)
    
    def _add_subsection_title(self, doc: Document, title: str):
        """Adiciona subtítulo"""
        doc.add_paragraph(title, style='Subtitulo ETP')
    
    def _add_formatted_paragraph(self, doc: Document, text: str):
        """Adiciona parágrafo formatado"""
        if text.strip():
            doc.add_paragraph(text, style='Corpo ETP')
    
    def _create_formatted_table(self, doc: Document, table_data: List[str]):
        """Cria tabela formatada baseada no modelo"""
        if not table_data:
            return
        
        # Processar dados da tabela
        processed_data = []
        for line in table_data:
            cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            if cells:
                processed_data.append(cells)
        
        if not processed_data:
            return
        
        # Determinar número de colunas
        max_cols = max(len(row) for row in processed_data)
        
        # Criar tabela
        table = doc.add_table(rows=len(processed_data), cols=max_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Preencher dados
        for row_idx, row_data in enumerate(processed_data):
            table_row = table.rows[row_idx]
            
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < len(table_row.cells):
                    cell = table_row.cells[col_idx]
                    cell.text = cell_data
                    
                    # Formatação da célula
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(11)
                            
                            # Primeira linha como cabeçalho
                            if row_idx == 0:
                                run.font.bold = True
                                # Aplicar fundo azul ao cabeçalho
                                self._add_cell_blue_background(cell)
                            
                            # Destacar valores monetários
                            if 'R$' in cell_data or 'TOTAL' in cell_data.upper():
                                run.font.bold = True
        
        # Adicionar espaço após tabela
        doc.add_paragraph()
    
    def _add_blue_background(self, paragraph):
        """Adiciona fundo azul escuro ao parágrafo"""
        try:
            shading_elm = parse_xml(
                f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="1f4e79"/>'
            )
            paragraph._element.get_or_add_pPr().append(shading_elm)
        except Exception:
            pass
    
    def _add_cell_blue_background(self, cell):
        """Adiciona fundo azul à célula da tabela"""
        try:
            shading_elm = parse_xml(
                f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="1f4e79"/>'
            )
            cell._element.get_or_add_tcPr().append(shading_elm)
            
            # Alterar cor da fonte para branco
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = self.white_color
        except Exception:
            pass
    
    def _add_cell_border(self, cell):
        """Adiciona borda à célula"""
        try:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            
            # Criar elemento de bordas
            tcBorders = OxmlElement('w:tcBorders')
            
            # Definir bordas
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '12')  # Espessura
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')  # Cor preta
                tcBorders.append(border)
            
            tcPr.append(tcBorders)
        except Exception:
            pass
    
    def _add_paragraph_border(self, paragraph):
        """Adiciona borda ao parágrafo"""
        try:
            pPr = paragraph._element.get_or_add_pPr()
            
            # Criar elemento de bordas
            pBdr = OxmlElement('w:pBdr')
            
            # Definir bordas
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '8')
                border.set(qn('w:space'), '4')
                border.set(qn('w:color'), '1f4e79')  # Cor azul
                pBdr.append(border)
            
            pPr.append(pBdr)
        except Exception:
            pass
    
    def _remove_table_borders(self, table):
        """Remove bordas da tabela"""
        try:
            tbl = table._element
            tblPr = tbl.tblPr
            
            # Criar elemento de bordas vazias
            tblBorders = OxmlElement('w:tblBorders')
            
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'nil')
                tblBorders.append(border)
            
            tblPr.append(tblBorders)
        except Exception:
            pass
    
    def _apply_document_borders(self, doc: Document):
        """Aplica bordas ao documento inteiro"""
        try:
            # Acessar as configurações da seção
            sectPr = doc.sections[0]._sectPr
            
            # Criar elemento de bordas da página
            pgBorders = OxmlElement('w:pgBorders')
            pgBorders.set(qn('w:offsetFrom'), 'page')
            
            # Definir bordas da página
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '24')  # Espessura da borda (3pt)
                border.set(qn('w:space'), '24')  # Espaço da borda
                border.set(qn('w:color'), '000000')  # Cor preta
                pgBorders.append(border)
            
            sectPr.append(pgBorders)
        except Exception as e:
            print(f"Aviso: Não foi possível aplicar bordas ao documento: {e}")
    
    def _add_footer_with_info(self, doc: Document):
        """Adiciona rodapé com informações"""
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.clear()
        
        # Linha separadora
        separator_run = footer_para.add_run("_" * 80 + "\n")
        separator_run.font.size = Pt(8)
        
        # Informações do documento
        info_run = footer_para.add_run(
            f"Documento elaborado em conformidade com a Lei nº 14.133/2021\n"
            f"Data de elaboração: {datetime.now().strftime('%d/%m/%Y')}"
        )
        info_run.font.name = 'Times New Roman'
        info_run.font.size = Pt(10)
        info_run.font.bold = True
        
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Adicionar numeração de páginas
        page_para = footer.add_paragraph()
        page_run = page_para.add_run()
        
        # Campo de página atual
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        page_run._element.append(fldChar1)
        
        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"
        page_run._element.append(instrText)
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        page_run._element.append(fldChar2)
        
        page_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _save_document_with_timestamp(self, doc: Document) -> str:
        """Salva documento com timestamp em arquivo temporário"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"ETP_Com_Bordas_{timestamp}.docx"
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx", prefix="etp_", dir=None) as tmp:
            doc.save(tmp.name)
            return tmp.name
        
    def create_etp_with_borders(self, etp_content: str, session_data: Dict = None) -> str:
        """Método principal para criar ETP com bordas"""
        return self.create_document_with_borders(etp_content, session_data)

