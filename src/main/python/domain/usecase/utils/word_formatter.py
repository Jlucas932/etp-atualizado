import logging
import os
import tempfile
import re
from datetime import datetime
from typing import Dict, List, Optional
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


logger = logging.getLogger(__name__)

class ProfessionalWordFormatter:
    """Formatador profissional de documentos Word para ETP"""
    
    def __init__(self):
        self.blue_color = RGBColor(31, 78, 121)  # Azul escuro #1f4e79
        self.white_color = RGBColor(255, 255, 255)  # Branco
        
    def create_professional_document(self, content: str, session_data: Dict = None) -> str:
        """Cria documento Word com formatação profissional"""
        try:
            doc = Document()
            
            # Configurar margens e layout
            self._configure_page_layout(doc)
            
            # Criar estilos personalizados
            self._create_custom_styles(doc)
            
            # Verificar se o conteúdo já contém cabeçalho institucional
            content_has_header = self._content_has_institutional_header(content)
            content_has_main_title = self._content_has_main_title(content)
            
            # Adicionar cabeçalho apenas se não existir no conteúdo
            if not content_has_header:
                self._add_header(doc)
            
            # Adicionar título principal apenas se não existir no conteúdo
            if not content_has_main_title:
                self._add_main_title(doc)
            
            # Processar e adicionar conteúdo (com limpeza de duplicações)
            self._process_and_add_content(doc, content)
            
            # Adicionar rodapé
            self._add_footer(doc)
            
            # Adicionar seção de assinatura
            self._add_signature_section(doc, session_data)
            
            # Salvar documento
            doc_path = self._save_document(doc)
            
            return doc_path
            
        except Exception as e:
            raise Exception(f"Erro ao criar documento Word: {str(e)}")
    
    def _configure_page_layout(self, doc: Document):
        """Configura layout da página conforme especificações"""
        section = doc.sections[0]
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        
        # Configurar cabeçalho e rodapé
        section.header_distance = Inches(0.5)
        section.footer_distance = Inches(0.5)
        
        # Adicionar bordas da página (moldura)
        self._add_page_borders(doc)
    
    def _create_custom_styles(self, doc: Document):
        """Cria estilos personalizados para o documento"""
        styles = doc.styles
        
        # Estilo para títulos principais (com fundo azul)
        if 'Titulo Principal ETP' not in [s.name for s in styles]:
            title_style = styles.add_style('Titulo Principal ETP', WD_STYLE_TYPE.PARAGRAPH)
            
            # Formatação do parágrafo
            title_format = title_style.paragraph_format
            title_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            title_format.space_before = Pt(12)
            title_format.space_after = Pt(12)
            title_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            
            # Formatação da fonte
            title_font = title_style.font
            title_font.name = 'Calibri'
            title_font.size = Pt(12)
            title_font.bold = True
            title_font.color.rgb = self.white_color
        
        # Estilo para subtítulos
        if 'Subtitulo ETP' not in [s.name for s in styles]:
            subtitle_style = styles.add_style('Subtitulo ETP', WD_STYLE_TYPE.PARAGRAPH)
            
            subtitle_format = subtitle_style.paragraph_format
            subtitle_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            subtitle_format.space_before = Pt(8)
            subtitle_format.space_after = Pt(6)
            subtitle_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            
            subtitle_font = subtitle_style.font
            subtitle_font.name = 'Calibri'
            subtitle_font.size = Pt(12)
            subtitle_font.bold = True
            subtitle_font.color.rgb = RGBColor(0, 0, 0)
        
        # Estilo para corpo do texto
        if 'Corpo Texto ETP' not in [s.name for s in styles]:
            body_style = styles.add_style('Corpo Texto ETP', WD_STYLE_TYPE.PARAGRAPH)
            
            body_format = body_style.paragraph_format
            body_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            body_format.line_spacing = 1.5
            body_format.first_line_indent = Inches(0.49)  # 1,25cm
            body_format.space_after = Pt(6)
            
            body_font = body_style.font
            body_font.name = 'Calibri'
            body_font.size = Pt(12)
            body_font.color.rgb = RGBColor(0, 0, 0)
    
    def _add_header(self, doc: Document):
        """Adiciona cabeçalho institucional"""
        header = doc.sections[0].header
        header_para = header.paragraphs[0]
        
        # Limpar parágrafo existente
        header_para.clear()
        
        # Adicionar texto do cabeçalho
        run = header_para.add_run("GOVERNO DO ESTADO")
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        header_para.add_run("\n")
        
        run = header_para.add_run("SECRETARIA DE ADMINISTRAÇÃO")
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        header_para.add_run("\n")
        
        run = header_para.add_run("ESTUDO TÉCNICO PRELIMINAR")
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Adicionar linha separadora
        header.add_paragraph("_" * 80).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_main_title(self, doc: Document):
        """Adiciona título principal do documento com fundo azul"""
        # Cabeçalho azul com fundo preenchido
        blue_header_para = doc.add_paragraph()
        blue_header_run = blue_header_para.add_run("ESTUDO TÉCNICO PRELIMINAR")
        blue_header_run.font.name = 'Calibri'
        blue_header_run.font.size = Pt(12)
        blue_header_run.font.bold = True
        blue_header_run.font.color.rgb = self.white_color
        blue_header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Adicionar fundo azul
        self._add_blue_background(blue_header_para)
        
        # Espaço
        doc.add_paragraph()
        
        # Caixa de apresentação inicial com borda cinza
        self._add_introductory_box(doc)
    
    def _add_introductory_box(self, doc: Document):
        """Adiciona caixa de apresentação inicial com borda cinza"""
        # Primeiro parágrafo introdutório
        intro_para1 = doc.add_paragraph(
            "O presente documento caracteriza a primeira etapa da fase de planejamento e apresenta "
            "os devidos estudos para a contratação de solução que melhor atenderá à necessidade "
            "descrita abaixo, em observância às normas vigentes e aos princípios que regem a "
            "Administração Pública, especialmente a Lei nº 14.133/2021."
        )
        intro_para1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        intro_para1.paragraph_format.line_spacing = 1.15  # Leve espaçamento entre linhas
        
        # Formatação do primeiro parágrafo
        for run in intro_para1.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Segundo parágrafo com objetivo
        intro_para2 = doc.add_paragraph(
            "O objetivo principal é identificar a necessidade de contratação e definir os "
            "requisitos técnicos e especificações que orientarão o processo licitatório, "
            "garantindo a melhor solução para a Administração Pública."
        )
        intro_para2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        intro_para2.paragraph_format.line_spacing = 1.15  # Leve espaçamento entre linhas
        
        # Formatação do segundo parágrafo
        for run in intro_para2.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Adicionar borda cinza fina aos parágrafos
        self._add_gray_border(intro_para1)
        self._add_gray_border(intro_para2)
        
        # Espaço após a caixa
        doc.add_paragraph()
    
    def _add_gray_border(self, paragraph):
        """Adiciona borda cinza discreta ao parágrafo"""
        try:
            # Criar elemento de borda
            border_elm = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:top w:val="single" w:sz="4" w:space="1" w:color="808080"/>'
                f'<w:left w:val="single" w:sz="4" w:space="4" w:color="808080"/>'
                f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="808080"/>'
                f'<w:right w:val="single" w:sz="4" w:space="4" w:color="808080"/>'
                f'</w:pBdr>'
            )
            paragraph._element.get_or_add_pPr().append(border_elm)
        except Exception:
            # Se não conseguir aplicar a borda, continuar sem ela
            pass
    
    def _process_and_add_content(self, doc: Document, content: str):
        """Processa e adiciona o conteúdo principal do ETP (com limpeza de duplicações)"""
        lines = content.split('\n')
        current_paragraph = []
        seen_sections = set()  # Para evitar seções duplicadas
        
        # Filtrar linhas de cabeçalho institucional se existirem
        filtered_lines = self._filter_institutional_headers(lines)
        
        for line in filtered_lines:
            line = line.strip()
            
            if not line:
                # Linha vazia - finalizar parágrafo atual se houver
                if current_paragraph:
                    self._add_paragraph_to_doc(doc, '\n'.join(current_paragraph))
                    current_paragraph = []
                continue
            
            # Identificar tipo de linha
            if self._is_main_section_title(line):
                # Verificar se já vimos esta seção
                section_key = line.upper().strip()
                if section_key in seen_sections:
                    continue  # Pular seção duplicada
                
                seen_sections.add(section_key)
                
                # Finalizar parágrafo anterior
                if current_paragraph:
                    self._add_paragraph_to_doc(doc, '\n'.join(current_paragraph))
                    current_paragraph = []
                
                # Adicionar título principal
                self._add_main_section_title(doc, line)
                
            elif self._is_subsection_title(line):
                # Finalizar parágrafo anterior
                if current_paragraph:
                    self._add_paragraph_to_doc(doc, '\n'.join(current_paragraph))
                    current_paragraph = []
                
                # Adicionar subtítulo
                self._add_subsection_title(doc, line)
                
            elif self._is_table_content(line):
                # Finalizar parágrafo anterior
                if current_paragraph:
                    self._add_paragraph_to_doc(doc, '\n'.join(current_paragraph))
                    current_paragraph = []
                
                # Processar tabela
                self._add_table_content(doc, line)
                
            else:
                # Conteúdo normal - adicionar ao parágrafo atual
                current_paragraph.append(line)
        
        # Finalizar último parágrafo
        if current_paragraph:
            self._add_paragraph_to_doc(doc, '\n'.join(current_paragraph))
    
    def _is_main_section_title(self, line: str) -> bool:
        """Verifica se a linha é um título de seção principal"""
        pattern = r'^\d+\.\s+[A-ZÁÀÂÃÉÊÍÓÔÕÚÇ\s]+$'
        return bool(re.match(pattern, line.upper()))
    
    def _is_subsection_title(self, line: str) -> bool:
        """Verifica se a linha é um título de subseção"""
        pattern = r'^\d+\.\d+\s+[A-Za-záàâãéêíóôõúç\s]+'
        return bool(re.match(pattern, line))
    
    def _is_table_content(self, line: str) -> bool:
        """Verifica se a linha contém conteúdo de tabela"""
        return '|' in line and line.count('|') >= 2
    
    def _add_main_section_title(self, doc: Document, title: str):
        """Adiciona título de seção principal com fundo azul"""
        para = doc.add_paragraph(title.upper(), style='Titulo Principal ETP')
        
        # Adicionar fundo azul escuro
        self._add_blue_background(para)
    
    def _add_subsection_title(self, doc: Document, title: str):
        """Adiciona título de subseção"""
        doc.add_paragraph(title, style='Subtitulo ETP')
    
    def _add_paragraph_to_doc(self, doc: Document, text: str):
        """Adiciona parágrafo de texto normal"""
        if text.strip():
            doc.add_paragraph(text, style='Corpo Texto ETP')
    
    def _add_table_content(self, doc: Document, table_line: str):
        """Adiciona conteúdo de tabela formatada"""
        # Dividir linha por |
        cells = [cell.strip() for cell in table_line.split('|') if cell.strip()]
        
        if len(cells) >= 2:
            # Criar tabela simples
            table = doc.add_table(rows=1, cols=len(cells))
            table.style = 'Table Grid'
            
            # Adicionar dados
            row = table.rows[0]
            for i, cell_text in enumerate(cells):
                if i < len(row.cells):
                    cell = row.cells[i]
                    cell.text = cell_text
                    
                    # Formatação da célula
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.name = 'Calibri'
                            run.font.size = Pt(10)
    
    def _add_blue_background(self, paragraph):
        """Adiciona fundo azul escuro ao parágrafo"""
        try:
            # Criar elemento de sombreamento
            shading_elm = parse_xml(
                f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="1f4e79"/>'
            )
            paragraph._element.get_or_add_pPr().append(shading_elm)
        except Exception:
            # Se não conseguir aplicar o fundo, continuar sem ele
            pass
    
    def _add_page_borders(self, doc: Document):
        """Adiciona bordas da página (moldura) com linha cinza discreta"""
        try:
            # Acessar as configurações da seção
            sectPr = doc.sections[0]._sectPr
            
            # Criar elemento de bordas da página
            pgBorders = OxmlElement('w:pgBorders')
            pgBorders.set(qn('w:offsetFrom'), 'page')
            
            # Definir bordas da página com cor cinza discreta
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')  # Linha reta
                border.set(qn('w:sz'), '12')  # Espessura média (1.5pt)
                border.set(qn('w:space'), '24')  # Espaço da borda
                border.set(qn('w:color'), '808080')  # Cor cinza discreta
                pgBorders.append(border)
            
            sectPr.append(pgBorders)
        except Exception as e:
            # Se não conseguir aplicar bordas, continuar sem elas
            logger.warning("Aviso: Não foi possível aplicar bordas da página: %s", e, exc_info=True)
    
    def _add_footer(self, doc: Document):
        """Adiciona rodapé com numeração de páginas"""
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Adicionar numeração de página
        run = footer_para.runs[0] if footer_para.runs else footer_para.add_run()
        
        # Criar campo de numeração
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._element.append(fldChar1)
        
        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"
        run._element.append(instrText)
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._element.append(fldChar2)
        
        # Adicionar texto adicional
        footer_para.add_run(" de ")
        
        # Campo para total de páginas
        run2 = footer_para.add_run()
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'begin')
        run2._element.append(fldChar3)
        
        instrText2 = OxmlElement('w:instrText')
        instrText2.text = "NUMPAGES"
        run2._element.append(instrText2)
        
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')
        run2._element.append(fldChar4)
    
    def _add_signature_section(self, doc: Document, session_data: Dict = None):
        """Adiciona seção de assinatura"""
        # Espaço antes da assinatura
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Linha para assinatura
        sig_para = doc.add_paragraph("_" * 50)
        sig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Nome do responsável
        name_para = doc.add_paragraph("Nome do Responsável pela Elaboração")
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.runs[0]
        name_run.font.name = 'Calibri'
        name_run.font.size = Pt(12)
        name_run.font.bold = True
        
        # Cargo
        cargo_para = doc.add_paragraph("Cargo/Função")
        cargo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cargo_run = cargo_para.runs[0]
        cargo_run.font.name = 'Calibri'
        cargo_run.font.size = Pt(11)
        
        # Data
        date_para = doc.add_paragraph(f"Data: {datetime.now().strftime('%d/%m/%Y')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.runs[0]
        date_run.font.name = 'Calibri'
        date_run.font.size = Pt(11)
    
    def _save_document(self, doc: Document) -> str:
        """Salva o documento e retorna o caminho"""
        temp_dir = tempfile.gettempdir()
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"ETP_Profissional_{timestamp}.docx"
        doc_path = os.path.join(temp_dir, filename)
        
        doc.save(doc_path)
        return doc_path
    
    def _content_has_institutional_header(self, content: str) -> bool:
        """Verifica se o conteúdo já possui cabeçalho institucional"""
        lines = content.upper().split('\n')
        for i, line in enumerate(lines[:10]):  # Verificar primeiras 10 linhas
            if 'GOVERNO DO ESTADO' in line:
                return True
            if 'SECRETARIA DE ADMINISTRAÇÃO' in line:
                return True
        return False
    
    def _content_has_main_title(self, content: str) -> bool:
        """Verifica se o conteúdo já possui título principal do ETP"""
        lines = content.upper().split('\n')
        for line in lines[:15]:  # Verificar primeiras 15 linhas
            if 'ESTUDO TÉCNICO PRELIMINAR' in line and 'ETP' in line:
                return True
        return False
    
    def _filter_institutional_headers(self, lines: List[str]) -> List[str]:
        """Remove linhas de cabeçalho institucional duplicadas do conteúdo"""
        filtered_lines = []
        institutional_keywords = [
            'GOVERNO DO ESTADO',
            'SECRETARIA DE ADMINISTRAÇÃO',
            'ESTUDO TÉCNICO PRELIMINAR'
        ]
        
        skip_next_empty_lines = 0
        
        for line in lines:
            line_upper = line.strip().upper()
            
            # Verificar se é linha de cabeçalho institucional
            is_institutional = any(keyword in line_upper for keyword in institutional_keywords)
            
            if is_institutional:
                # Pular esta linha e as próximas linhas vazias
                skip_next_empty_lines = 3
                continue
            
            # Pular linhas vazias após cabeçalho institucional
            if skip_next_empty_lines > 0 and not line.strip():
                skip_next_empty_lines -= 1
                continue
            
            # Resetar contador se encontrou conteúdo
            if line.strip():
                skip_next_empty_lines = 0
            
            filtered_lines.append(line)
        
        return filtered_lines
    
    def create_table_from_data(self, doc: Document, headers: List[str], data: List[List[str]], title: str = None):
        """Cria tabela formatada com dados"""
        if title:
            title_para = doc.add_paragraph(title, style='Subtitulo ETP')
        
        # Criar tabela
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        # Adicionar cabeçalhos
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = header
            
            # Formatação do cabeçalho
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
                    run.font.bold = True
            
            # Fundo cinza claro para cabeçalho
            try:
                shading_elm = parse_xml(
                    f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="f2f2f2"/>'
                )
                cell._element.get_or_add_tcPr().append(shading_elm)
            except:
                pass
        
        # Adicionar dados
        for row_data in data:
            row = table.add_row()
            for i, cell_data in enumerate(row_data):
                if i < len(row.cells):
                    cell = row.cells[i]
                    cell.text = str(cell_data)
                    
                    # Formatação da célula
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.name = 'Calibri'
                            run.font.size = Pt(10)
        
        # Espaço após tabela
        doc.add_paragraph()


class WordFormatter:
    """Simple wrapper for Word document creation compatible with EtpController"""
    
    def __init__(self):
        self.formatter = ProfessionalWordFormatter()
    
    def create_word_document(self, content: str, title: str = None) -> BytesIO:
        """Creates a Word document and returns a BytesIO buffer"""
        try:
            # Create document using professional formatter
            doc = Document()
            
            # Configure page layout
            self.formatter._configure_page_layout(doc)
            
            # Create custom styles
            self.formatter._create_custom_styles(doc)
            
            # Add header
            self.formatter._add_header(doc)
            
            # Add main title
            self.formatter._add_main_title(doc)
            
            # Process and add content
            self.formatter._process_and_add_content(doc, content)
            
            # Add footer
            self.formatter._add_footer(doc)
            
            # Add signature section
            self.formatter._add_signature_section(doc, None)
            
            # Create BytesIO buffer
            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            return doc_buffer
            
        except Exception as e:
            raise Exception(f"Erro ao criar documento Word: {str(e)}")

