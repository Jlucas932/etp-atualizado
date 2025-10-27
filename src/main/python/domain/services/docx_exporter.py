"""
DOCX Exporter Service
Generates ETP documents in .docx format using a template with proper styles.
"""

import os
import logging
from io import BytesIO
from typing import Dict, Any, List, Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

logger = logging.getLogger(__name__)


class DocxExporter:
    """
    Service to export ETP data to .docx format.
    Uses a template file when available, otherwise creates styled document.
    """
    
    def __init__(self, template_path: Optional[str] = None):
        """
        Initialize the exporter.
        
        Args:
            template_path: Path to template .docx file with styles and placeholders
        """
        self.template_path = template_path
        
        # Check if template exists
        if template_path and os.path.exists(template_path):
            self.use_template = True
            logger.info(f"[DOCX_EXPORTER] Using template: {template_path}")
        else:
            self.use_template = False
            logger.warning(f"[DOCX_EXPORTER] Template not found, using styled generation")
    
    def export_etp(self, etp_data: Dict[str, Any]) -> BytesIO:
        """
        Export ETP data to .docx format.
        
        Args:
            etp_data: Dictionary containing ETP information:
                - title: Document title
                - organ: Contracting organization
                - object: Procurement object
                - necessity: Description of need
                - requirements: List of requirements
                - solution_strategy: Solution approach
                - pca: PCA information
                - legal_norms: Legal basis
                - quant_value: Quantity and value
                - parcelamento: Installment plan
                - justifications: Justifications
                - signatures: Signature block
        
        Returns:
            BytesIO: Buffer containing the generated .docx file
        """
        if self.use_template:
            return self._export_with_template(etp_data)
        else:
            return self._export_with_styles(etp_data)
    
    def _export_with_template(self, etp_data: Dict[str, Any]) -> BytesIO:
        """Export using template file with placeholder replacement."""
        try:
            doc = Document(self.template_path)
            
            # Replace placeholders in paragraphs
            placeholders = self._build_placeholder_map(etp_data)
            
            for paragraph in doc.paragraphs:
                for key, value in placeholders.items():
                    if key in paragraph.text:
                        paragraph.text = paragraph.text.replace(key, str(value))
            
            # Replace placeholders in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for key, value in placeholders.items():
                            if key in cell.text:
                                cell.text = cell.text.replace(key, str(value))
            
            # Save to buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer
            
        except Exception as e:
            logger.error(f"[DOCX_EXPORTER] Error using template: {e}")
            # Fallback to styled generation
            return self._export_with_styles(etp_data)
    
    def _export_with_styles(self, etp_data: Dict[str, Any]) -> BytesIO:
        """Export by creating a styled document from scratch."""
        doc = Document()
        
        # Configure document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add title
        title = etp_data.get('title', 'Estudo Técnico Preliminar')
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata section
        self._add_section(doc, '1. IDENTIFICAÇÃO', level=1)
        organ = etp_data.get('organ', 'Não informado')
        self._add_paragraph(doc, f"Órgão: {organ}")
        
        object_desc = etp_data.get('object', 'Não informado')
        self._add_paragraph(doc, f"Objeto: {object_desc}")
        
        # Add necessity section
        necessity = etp_data.get('necessity', 'Não informada')
        self._add_section(doc, '2. NECESSIDADE DA CONTRATAÇÃO', level=1)
        self._add_paragraph(doc, necessity)
        
        # Add requirements section
        requirements = etp_data.get('requirements', [])
        if requirements:
            self._add_section(doc, '3. REQUISITOS DA CONTRATAÇÃO', level=1)
            for i, req in enumerate(requirements, 1):
                if isinstance(req, dict):
                    req_text = req.get('text', req.get('content', str(req)))
                else:
                    req_text = str(req)
                doc.add_paragraph(f"{i}. {req_text}", style='List Number')
        
        # Add solution strategy section
        solution = etp_data.get('solution_strategy', '')
        if solution:
            self._add_section(doc, '4. ESTRATÉGIA DE SOLUÇÃO', level=1)
            self._add_paragraph(doc, solution)
        
        # Add PCA section
        pca = etp_data.get('pca', '')
        if pca:
            self._add_section(doc, '5. PESQUISA DE PREÇOS E ESTIMATIVA DE CUSTOS', level=1)
            self._add_paragraph(doc, pca)
        
        # Add legal norms section
        legal_norms = etp_data.get('legal_norms', '')
        if legal_norms:
            self._add_section(doc, '6. FUNDAMENTAÇÃO LEGAL', level=1)
            self._add_paragraph(doc, legal_norms)
        
        # Add quantity and value section
        quant_value = etp_data.get('quant_value', '')
        if quant_value:
            self._add_section(doc, '7. QUANTITATIVO E VALOR ESTIMADO', level=1)
            self._add_paragraph(doc, quant_value)
        
        # Add installment section
        parcelamento = etp_data.get('parcelamento', '')
        if parcelamento:
            self._add_section(doc, '8. PARCELAMENTO', level=1)
            self._add_paragraph(doc, parcelamento)
        
        # Add justifications section
        justifications = etp_data.get('justifications', '')
        if justifications:
            self._add_section(doc, '9. JUSTIFICATIVAS', level=1)
            self._add_paragraph(doc, justifications)
        
        # Add signatures section
        signatures = etp_data.get('signatures', '')
        if signatures:
            doc.add_page_break()
            self._add_section(doc, 'ASSINATURAS', level=1)
            self._add_paragraph(doc, signatures)
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    def _build_placeholder_map(self, etp_data: Dict[str, Any]) -> Dict[str, str]:
        """Build a map of placeholders to their values."""
        return {
            '{{titulo}}': etp_data.get('title', 'Estudo Técnico Preliminar'),
            '{{orgao}}': etp_data.get('organ', 'Não informado'),
            '{{objeto}}': etp_data.get('object', 'Não informado'),
            '{{necessidade}}': etp_data.get('necessity', 'Não informada'),
            '{{requisitos}}': self._format_list(etp_data.get('requirements', [])),
            '{{estrategia}}': etp_data.get('solution_strategy', ''),
            '{{pca}}': etp_data.get('pca', ''),
            '{{normas}}': etp_data.get('legal_norms', ''),
            '{{quantitativo_valor}}': etp_data.get('quant_value', ''),
            '{{parcelamento}}': etp_data.get('parcelamento', ''),
            '{{justificativas}}': etp_data.get('justifications', ''),
            '{{assinaturas}}': etp_data.get('signatures', '')
        }
    
    def _format_list(self, items: List) -> str:
        """Format a list of items as numbered text."""
        if not items:
            return 'Nenhum requisito especificado'
        
        formatted = []
        for i, item in enumerate(items, 1):
            if isinstance(item, dict):
                text = item.get('text', item.get('content', str(item)))
            else:
                text = str(item)
            formatted.append(f"{i}. {text}")
        
        return '\n'.join(formatted)
    
    def _add_section(self, doc: Document, title: str, level: int = 1):
        """Add a section heading."""
        doc.add_heading(title, level=level)
    
    def _add_paragraph(self, doc: Document, text: str):
        """Add a paragraph with proper formatting."""
        if not text:
            return
        
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Set font
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(11)


def create_etp_template(output_path: str):
    """
    Create a template .docx file with proper styles and placeholders.
    This template can be customized and reused.
    
    Args:
        output_path: Path where to save the template file
    """
    doc = Document()
    
    # Configure margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    # Add title
    title = doc.add_heading('{{titulo}}', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add sections with placeholders
    doc.add_heading('1. IDENTIFICAÇÃO', level=1)
    doc.add_paragraph('Órgão: {{orgao}}')
    doc.add_paragraph('Objeto: {{objeto}}')
    
    doc.add_heading('2. NECESSIDADE DA CONTRATAÇÃO', level=1)
    doc.add_paragraph('{{necessidade}}')
    
    doc.add_heading('3. REQUISITOS DA CONTRATAÇÃO', level=1)
    doc.add_paragraph('{{requisitos}}')
    
    doc.add_heading('4. ESTRATÉGIA DE SOLUÇÃO', level=1)
    doc.add_paragraph('{{estrategia}}')
    
    doc.add_heading('5. PESQUISA DE PREÇOS E ESTIMATIVA DE CUSTOS', level=1)
    doc.add_paragraph('{{pca}}')
    
    doc.add_heading('6. FUNDAMENTAÇÃO LEGAL', level=1)
    doc.add_paragraph('{{normas}}')
    
    doc.add_heading('7. QUANTITATIVO E VALOR ESTIMADO', level=1)
    doc.add_paragraph('{{quantitativo_valor}}')
    
    doc.add_heading('8. PARCELAMENTO', level=1)
    doc.add_paragraph('{{parcelamento}}')
    
    doc.add_heading('9. JUSTIFICATIVAS', level=1)
    doc.add_paragraph('{{justificativas}}')
    
    doc.add_page_break()
    doc.add_heading('ASSINATURAS', level=1)
    doc.add_paragraph('{{assinaturas}}')
    
    # Save template
    doc.save(output_path)
    logger.info(f"[DOCX_EXPORTER] Template created at: {output_path}")
